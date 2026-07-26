"""
session_manager.py  —  Multi-Part Session Manager for One-Way ANOVA
Handles saving / loading / exporting of multiple ANOVA parts
(Part 1-A, 1-B, 2-A …) with DOCX and PDF batch export.
"""

import re
import numpy as np
import customtkinter as ctk
from tkinter import messagebox
from datetime import datetime

# ── Palette (matches main app) ────────────────────────────────────────────────
from ui_theme import (
    BG_DEEP, BG_CARD, BG_PANEL, BG_INPUT,
    ACCENT, ACCENT2, DANGER, WARN, SUCCESS, PURPLE,
    TEXT_PRI, TEXT_SEC, BORDER,
)

FONT_HEAD = ("Segoe UI", 14, "bold")
FONT_CARD = ("Segoe UI", 12, "bold")
FONT_BODY = ("Segoe UI", 11)
FONT_BTN  = ("Segoe UI", 11, "bold")
FONT_MONO = ("Consolas", 9)
FONT_TINY = ("Segoe UI", 9)


# ── Label generator ───────────────────────────────────────────────────────────

def _next_part_label(existing_labels: list) -> str:
    """Generate next Part label: Part 1-A, Part 1-B … Part 1-Z, Part 2-A …"""
    if not existing_labels:
        return "Part 1-A"
    max_num, max_letter = 1, "A"
    for lbl in existing_labels:
        m = re.match(r"Part (\d+)-([A-Z])", lbl)
        if m:
            n, l = int(m.group(1)), m.group(2)
            if n > max_num or (n == max_num and l > max_letter):
                max_num, max_letter = n, l
    next_letter = chr(ord(max_letter) + 1)
    if next_letter > "Z":
        return f"Part {max_num + 1}-A"
    return f"Part {max_num}-{next_letter}"


def _fmt(v, d=2):
    try:
        return f"{float(v):.{d}f}"
    except (TypeError, ValueError):
        return "N/A"


def _sig_color(is_significant: bool) -> str:
    return SUCCESS if is_significant else DANGER


# ── Part Card ─────────────────────────────────────────────────────────────────

class PartCard(ctk.CTkFrame):
    """One row in the session panel."""

    def __init__(self, master, part_data: dict, index: int,
                 on_delete, on_view, **kw):
        super().__init__(master, fg_color=BG_PANEL, corner_radius=8,
                         border_width=1, border_color=BORDER, **kw)
        self.part_data = part_data
        self._build(index, on_delete, on_view)

    def _build(self, index, on_delete, on_view):
        top = ctk.CTkFrame(self, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=(8, 2))

        # Coloured part-label badge
        badge_color = [ACCENT, ACCENT2, WARN, PURPLE, SUCCESS][index % 5]
        ctk.CTkLabel(top, text=self.part_data["label"],
                     font=("Segoe UI", 10, "bold"),
                     fg_color=badge_color, text_color="#0d1117",
                     corner_radius=5, padx=8, pady=2).pack(side="left")

        # Significant / Not Significant badge
        is_sig = self.part_data.get("is_significant", False)
        sig_text  = "✓ Significant" if is_sig else "✗ Not Significant"
        sig_color = SUCCESS if is_sig else DANGER
        ctk.CTkLabel(top, text=sig_text,
                     font=FONT_TINY,
                     fg_color=sig_color, text_color="#fff",
                     corner_radius=4, padx=6, pady=2).pack(side="left", padx=6)

        # Delete button
        ctk.CTkButton(top, text="✕", width=24, height=24,
                      fg_color="transparent", hover_color=DANGER,
                      text_color=TEXT_SEC, corner_radius=4,
                      font=("Segoe UI", 10, "bold"),
                      command=lambda: on_delete(self.part_data["label"])
                      ).pack(side="right")

        # View button
        ctk.CTkButton(top, text="👁", width=28, height=24,
                      fg_color=BG_INPUT, hover_color=ACCENT2,
                      text_color=TEXT_PRI, corner_radius=4,
                      font=("Segoe UI", 10),
                      command=lambda: on_view(self.part_data)
                      ).pack(side="right", padx=4)

        # Key stats row
        r = self.part_data
        df_b = r.get("df_between", "?")
        df_w = r.get("df_within",  "?")
        F    = _fmt(r.get("F_statistic"))
        p    = _fmt(r.get("p_value"))
        k    = len(r.get("groups", []))
        stats = (f"F({df_b}, {df_w}) = {F}   p = {p}   "
                 f"k = {k} groups   {r.get('decision', '')}")
        ctk.CTkLabel(self, text=stats, font=FONT_TINY,
                     text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 4))

        # Subtitle / title
        sub = r.get("report_subtitle") or r.get("report_title", "")
        if sub:
            ctk.CTkLabel(self, text=sub, font=("Segoe UI", 9, "italic"),
                         text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 6))


# ── Session Manager Panel ─────────────────────────────────────────────────────

class SessionManagerPanel(ctk.CTkToplevel):
    """Floating panel listing all saved ANOVA parts."""

    def __init__(self, master, saved_parts: list,
                 on_delete_part, on_export_all_docx, on_export_all_pdf):
        super().__init__(master)
        self.title("📁  Session Manager — Saved ANOVA Parts")
        self.geometry("520x660")
        self.minsize(420, 480)
        self.configure(fg_color=BG_DEEP)
        self.resizable(True, True)

        self.saved_parts    = saved_parts
        self.on_delete_part = on_delete_part
        self.on_export_docx = on_export_all_docx
        self.on_export_pdf  = on_export_all_pdf

        self.transient(master)
        self.lift()
        self.focus_force()

        self._build()

    def _build(self):
        # Header
        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=52)
        hdr.pack(fill="x"); hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="📁  Session Manager — ANOVA Parts",
                     font=FONT_HEAD, text_color=TEXT_PRI).pack(side="left", padx=16)
        self.count_label = ctk.CTkLabel(hdr,
                     text=f"{len(self.saved_parts)} part(s) saved",
                     font=FONT_TINY, text_color=TEXT_SEC)
        self.count_label.pack(side="right", padx=16)

        # Scrollable list
        self.scroll = ctk.CTkScrollableFrame(self, fg_color=BG_DEEP,
                                              scrollbar_button_color=BORDER)
        self.scroll.pack(fill="both", expand=True, padx=12, pady=12)
        self._populate_cards()

        # Export / close buttons
        btn_frame = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=68)
        btn_frame.pack(fill="x"); btn_frame.pack_propagate(False)

        ctk.CTkButton(btn_frame, text="📄  Export All → DOCX",
                      fg_color="#1d4ed8", hover_color="#1e3a8a",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_docx).pack(side="left", padx=(14, 6), pady=16)

        ctk.CTkButton(btn_frame, text="🖨  Export All → PDF",
                      fg_color=DANGER, hover_color="#b91c1c",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_pdf).pack(side="left", padx=6, pady=16)

        ctk.CTkButton(btn_frame, text="🗑  Clear All",
                      fg_color="#7f1d1d", hover_color="#450a0a",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self._clear_all).pack(side="left", padx=6, pady=16)

        ctk.CTkButton(btn_frame, text="✕ Close",
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.destroy).pack(side="right", padx=14, pady=16)

    def _populate_cards(self):
        for w in self.scroll.winfo_children():
            w.destroy()
        if not self.saved_parts:
            ctk.CTkLabel(self.scroll,
                         text="No parts saved yet.\nCompute ANOVA and click  '🗂 Save Part'.",
                         font=FONT_BODY, text_color=TEXT_SEC,
                         justify="center").pack(expand=True, pady=40)
            return
        for i, part in enumerate(self.saved_parts):
            c = PartCard(self.scroll, part, i,
                         on_delete=self._handle_delete,
                         on_view=self._handle_view)
            c.pack(fill="x", pady=4)

    def _handle_delete(self, label: str):
        if messagebox.askyesno("Delete Part",
                                f"Remove '{label}' from this session?"):
            self.on_delete_part(label)
            self.refresh()

    def _clear_all(self):
        if not self.saved_parts:
            messagebox.showinfo("Nothing to Clear", "There are no saved parts to clear.")
            return
        n = len(self.saved_parts)
        if messagebox.askyesno(
                "Clear All Parts",
                f"This will permanently remove all {n} saved part(s).\n\nAre you sure?"):
            for part in list(self.saved_parts):
                self.on_delete_part(part["label"])
            self.refresh()

    def _handle_view(self, part_data: dict):
        ViewPartWindow(self, part_data)

    def refresh(self):
        self._populate_cards()
        self.count_label.configure(text=f"{len(self.saved_parts)} part(s) saved")


# ── View Part Window ──────────────────────────────────────────────────────────

class ViewPartWindow(ctk.CTkToplevel):
    def __init__(self, master, part_data: dict):
        super().__init__(master)
        self.title(f"View — {part_data['label']}")
        self.geometry("560x520")
        self.configure(fg_color=BG_DEEP)

        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=48)
        hdr.pack(fill="x"); hdr.pack_propagate(False)
        ctk.CTkLabel(hdr,
                     text=f"  {part_data['label']}  —  {part_data.get('report_title', '')}",
                     font=FONT_CARD, text_color=TEXT_PRI).pack(side="left", padx=12)

        txt = ctk.CTkTextbox(self, fg_color=BG_INPUT, text_color=TEXT_PRI,
                              font=FONT_MONO, corner_radius=0)
        txt.pack(fill="both", expand=True, padx=12, pady=12)
        txt.insert("1.0", _build_part_summary(part_data))
        txt.configure(state="disabled")

        ctk.CTkButton(self, text="Close", command=self.destroy,
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=32).pack(pady=(0, 12))


# ── Text summary helper ───────────────────────────────────────────────────────

def _build_part_summary(r: dict) -> str:
    line = "─" * 52
    out  = f"{r.get('label', '')}\n{line}\n"
    out += f"{r.get('report_title', '')}\n"
    if r.get("report_subtitle"): out += f"{r['report_subtitle']}\n"
    if r.get("researcher_name"): out += f"by: {r['researcher_name']}\n"
    out += f"{line}\n\n"

    out += "DESCRIPTIVE STATISTICS\n" + "─" * 30 + "\n"
    for name, g in zip(r.get("group_names", []), r.get("groups", [])):
        out += (f"  {name}:  n={len(g)}  "
                f"M={_fmt(np.mean(g))}  "
                f"SD={_fmt(np.std(g, ddof=1))}\n")

    out += f"\n{'Source':<16}{'SS':>10}{'df':>6}{'MS':>12}{'F':>10}\n"
    out += "─" * 55 + "\n"
    out += (f"{'Between':<16}{_fmt(r.get('SS_between')):>10}"
            f"{str(r.get('df_between', '')):>6}"
            f"{_fmt(r.get('MS_between')):>12}"
            f"{_fmt(r.get('F_statistic')):>10}\n")
    out += (f"{'Within':<16}{_fmt(r.get('SS_within')):>10}"
            f"{str(r.get('df_within', '')):>6}"
            f"{_fmt(r.get('MS_within')):>12}\n")
    df_total = (r.get("df_between", 0) or 0) + (r.get("df_within", 0) or 0)
    out += (f"{'Total':<16}{_fmt(r.get('SS_total')):>10}"
            f"{str(df_total):>6}\n")

    out += f"\n{'═'*52}\nTEST RESULTS\n{'═'*52}\n"
    out += (f"F({r.get('df_between')}, {r.get('df_within')}) = "
            f"{_fmt(r.get('F_statistic'))},  "
            f"p = {_fmt(r.get('p_value'))},  "
            f"α = {r.get('alpha', 0.05)}\n\n")
    out += f"Decision:   {r.get('decision', '')}\n\n"
    out += f"Conclusion:\n{r.get('conclusion', '')}\n\n"

    if r.get("is_significant") and r.get("tukey"):
        out += f"\n{'─'*52}\nPOST HOC — Tukey HSD\n{'─'*52}\n"
        out += str(r["tukey"]) + "\n"

    out += f"\nSaved: {r.get('saved_at', '—')}\n"
    return out


# ── DOCX Export (all parts) ───────────────────────────────────────────────────

def export_all_to_docx(parts: list, filepath: str):
    """Write all saved ANOVA parts into a single APA-style DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_section_margins(sect):
        sect.top_margin    = Inches(0.60)
        sect.bottom_margin = Inches(0.60)
        sect.left_margin   = Inches(0.65)
        sect.right_margin  = Inches(0.65)

    def _apply_2col(sectPr):
        old = sectPr.find(qn("w:cols"))
        if old is not None:
            sectPr.remove(old)
        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"),   "2")
        cols_el.set(qn("w:space"), "720")
        sectPr.append(cols_el)

    def _insert_section_break(doc, break_type="nextPage"):
        p_el = OxmlElement("w:p")
        pPr  = OxmlElement("w:pPr")
        sPr  = OxmlElement("w:sectPr")
        pg_mar = OxmlElement("w:pgMar")
        for attr, val in [("w:top","864"),("w:bottom","864"),
                           ("w:left","936"),("w:right","936"),
                           ("w:header","720"),("w:footer","720"),("w:gutter","0")]:
            pg_mar.set(qn(attr), val)
        sPr.append(pg_mar)
        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"), "2"); cols_el.set(qn("w:space"), "720")
        sPr.append(cols_el)
        pg_type = OxmlElement("w:type")
        pg_type.set(qn("w:val"), break_type)
        sPr.append(pg_type)
        pPr.append(sPr); p_el.append(pPr)
        doc.element.body.append(p_el)

    def apa_borders(table):
        tbl = table._tbl; tblPr = tbl.tblPr
        tb  = OxmlElement("w:tblBorders")
        for bn in ["top","left","bottom","right","insideH","insideV"]:
            b = OxmlElement(f"w:{bn}"); b.set(qn("w:val"), "none"); tb.append(b)
        for bn, sz in [("top","12"),("bottom","12")]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz); tb.append(b)
        tblPr.append(tb)

    def add_header_sep(table):
        for cell in table.rows[0].cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
            tcBorders.append(bot); tcPr.append(tcBorders)

    def cf(cell, text, bold=False, italic=False, size=9, align="left", color=None):
        para = cell.paragraphs[0]
        para.alignment = (WD_PARAGRAPH_ALIGNMENT.CENTER
                          if align == "center" else WD_PARAGRAPH_ALIGNMENT.LEFT)
        run = para.add_run(str(text))
        run.font.size = Pt(size); run.bold = bold; run.italic = italic
        if color: run.font.color.rgb = RGBColor(*color)

    doc = Document()
    _set_section_margins(doc.sections[0])
    _apply_2col(doc.sections[0]._sectPr)

    for part_idx, r in enumerate(parts):
        if part_idx > 0:
            _insert_section_break(doc, "nextPage")

        # Part label heading
        lp = doc.add_heading(r.get("label", f"Part {part_idx + 1}"), level=1)
        lp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in lp.runs:
            run.font.size = Pt(13); run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Title / subtitle / author
        tp = doc.add_paragraph(r.get("report_title", "ANOVA ANALYSIS RESULTS"))
        tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in tp.runs: run.bold = True; run.font.size = Pt(14)

        if r.get("report_subtitle"):
            sp = doc.add_paragraph(r["report_subtitle"])
            sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sp.runs: run.italic = True; run.font.size = Pt(11)

        if r.get("researcher_name"):
            np_ = doc.add_paragraph(r["researcher_name"])
            np_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in np_.runs: run.italic = True; run.font.size = Pt(10)

        doc.add_paragraph()

        # ── Table 1: Descriptive Statistics ──────────────────────────────────
        p = doc.add_paragraph(); p.add_run("Table 1\n").bold = True
        p.runs[0].font.size = Pt(10)
        tp2 = doc.add_paragraph(r.get("desc_table_title", "Descriptive Statistics for Groups"))
        if tp2.runs: tp2.runs[0].italic = True; tp2.runs[0].font.size = Pt(9)

        groups      = r.get("groups", [])
        group_names = r.get("group_names", [])

        dt = doc.add_table(rows=len(groups) + 1, cols=4)
        apa_borders(dt)
        for i, h in enumerate(["Group", "n", "M", "SD"]):
            cf(dt.rows[0].cells[i], h, bold=True, size=9, align="center")
        add_header_sep(dt)
        for i, (name, g) in enumerate(zip(group_names, groups), 1):
            cf(dt.rows[i].cells[0], name, size=9)
            cf(dt.rows[i].cells[1], str(len(g)),             size=9, align="center")
            cf(dt.rows[i].cells[2], _fmt(np.mean(g)),        size=9, align="center")
            cf(dt.rows[i].cells[3], _fmt(np.std(g, ddof=1)), size=9, align="center")

        doc.add_paragraph()

        # ── Table 2: ANOVA Summary ────────────────────────────────────────────
        tp3 = doc.add_paragraph(
            r.get("anova_table_title", "Analysis of Variance Summary Table"))
        if tp3.runs: tp3.runs[0].italic = True; tp3.runs[0].font.size = Pt(9)

        at = doc.add_table(rows=4, cols=6)
        apa_borders(at)
        for i, h in enumerate(["Source", "SS", "df", "MS", "F", "p"]):
            cf(at.rows[0].cells[i], h, bold=True, size=9, align="center")
        add_header_sep(at)

        row1 = at.rows[1]
        cf(row1.cells[0], r.get("anova_between_label", "Between Groups"), size=9)
        for col, val in enumerate([_fmt(r.get("SS_between")), str(r.get("df_between", "")),
                                    _fmt(r.get("MS_between")), _fmt(r.get("F_statistic")),
                                    _fmt(r.get("p_value"))], 1):
            cf(row1.cells[col], val, size=9, align="center")

        row2 = at.rows[2]
        cf(row2.cells[0], r.get("anova_within_label", "Within Groups"), size=9)
        for col, val in enumerate([_fmt(r.get("SS_within")), str(r.get("df_within", "")),
                                    _fmt(r.get("MS_within"))], 1):
            cf(row2.cells[col], val, size=9, align="center")

        row3 = at.rows[3]
        cf(row3.cells[0], r.get("anova_total_label", "Total"), size=9)
        df_total = (r.get("df_between", 0) or 0) + (r.get("df_within", 0) or 0)
        for col, val in enumerate([_fmt(r.get("SS_total")), str(df_total)], 1):
            cf(row3.cells[col], val, size=9, align="center")

        doc.add_paragraph()

        # ── Test Results ──────────────────────────────────────────────────────
        rh = doc.add_paragraph(); rh.add_run("Test Results\n").bold = True
        rh.runs[0].font.size = Pt(10)

        if r.get("edited") and r.get("conclusion_text"):
            cp = doc.add_paragraph(r["conclusion_text"])
            for run in cp.runs: run.font.size = Pt(9)
        else:
            cp = doc.add_paragraph()
            cp.add_run(f"F({r.get('df_between')}, {r.get('df_within')}) = ")
            fr = cp.add_run(_fmt(r.get("F_statistic"))); fr.bold = True
            cp.add_run(", p = ")
            pr_ = cp.add_run(_fmt(r.get("p_value"))); pr_.bold = True
            cp.add_run("\n\nDecision: ")
            dr = cp.add_run(r.get("decision", "")); dr.bold = True
            cp.add_run(f"\n\n{r.get('conclusion', '')}")
            for run in cp.runs: run.font.size = Pt(9)

        # ── Post Hoc (Tukey) ──────────────────────────────────────────────────
        if r.get("is_significant") and r.get("tukey"):
            doc.add_paragraph()
            php = doc.add_paragraph(r.get("posthoc_title", "Post Hoc Comparisons (Tukey HSD)"))
            if php.runs: php.runs[0].italic = True; php.runs[0].font.size = Pt(9)
            phtext = r.get("posthoc_text", str(r["tukey"]))
            pph = doc.add_paragraph(phtext)
            for run in pph.runs: run.font.name = "Courier New"; run.font.size = Pt(7)

        doc.add_paragraph()

        # ── Raw Data Table ────────────────────────────────────────────────────
        rdt_title = doc.add_paragraph(r.get("rawdata_table_title", "Raw Data by Group"))
        if rdt_title.runs: rdt_title.runs[0].italic = True; rdt_title.runs[0].font.size = Pt(9)

        rdt = doc.add_table(rows=len(groups) + 1, cols=3)
        apa_borders(rdt)
        for i, h in enumerate(["Group", "n", "Values"]):
            cf(rdt.rows[0].cells[i], h, bold=True, size=8, align="center")
        add_header_sep(rdt)

        if r.get("raw_data_edits"):
            for i, edit in enumerate(r["raw_data_edits"], 1):
                cf(rdt.rows[i].cells[0], edit["group_name"], size=7)
                n_c = len(edit["values_text"].split(","))
                cf(rdt.rows[i].cells[1], str(n_c), size=7, align="center")
                vt = edit["values_text"]
                if len(vt) > 60: vt = vt[:57] + "…"
                cf(rdt.rows[i].cells[2], vt, size=7)
        else:
            for i, (name, g) in enumerate(zip(group_names, groups), 1):
                cf(rdt.rows[i].cells[0], name, size=7)
                cf(rdt.rows[i].cells[1], str(len(g)), size=7, align="center")
                vs = ", ".join(_fmt(v) for v in g)
                if len(vs) > 60: vs = vs[:57] + "…"
                cf(rdt.rows[i].cells[2], vs, size=7)

        doc.add_paragraph()

        # Per-part footer
        fp = doc.add_paragraph()
        _f = fp.add_run(
            f"Part saved: {r.get('saved_at', '')}   |   "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}"
        )
        _f.font.size = Pt(7)
        _f.font.color.rgb = RGBColor(128, 128, 128)
        _f.italic = True

    # Apply 2-col + margins to the final body sectPr
    body_sectPr = doc.element.body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        pg_mar = body_sectPr.find(qn("w:pgMar"))
        if pg_mar is None:
            pg_mar = OxmlElement("w:pgMar"); body_sectPr.append(pg_mar)
        for attr, val in [("w:top","864"),("w:bottom","864"),
                           ("w:left","936"),("w:right","936")]:
            pg_mar.set(qn(attr), val)
        _apply_2col(body_sectPr)

    doc.save(filepath)


# ── PDF Export (all parts) ────────────────────────────────────────────────────

def export_all_to_pdf(parts: list, filepath: str):
    """
    Write all saved ANOVA parts into a single APA-style PDF.
    Single-column layout — maximised, readable font sizes, one page per part.

    Font size summary vs old two-column version:
        Part label   13 pt  (was 8.5 pt)
        Title        12 pt  (was 7.5 pt)
        Subtitle     10 pt  (was 6.5 pt)
        Section hdr   9.5 pt  (was 6.5 pt)
        Body / result 9 pt  (was 6 pt)
        Table cells   8.5 pt  (was 5.5 pt)
        Mono / Tukey  7.5 pt  (was 4.8 pt)
        Footer        7.5 pt  (was 5 pt)
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, Table, TableStyle,
        HRFlowable, NextPageTemplate, PageBreak,
    )
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # ── Page geometry ─────────────────────────────────────────────────────────
    PAGE_W, PAGE_H = letter          # 612 × 792 pt
    ML = MR = 0.55 * inch
    MT = MB = 0.45 * inch
    COL_W = PAGE_W - ML - MR        # full usable width ≈ 7.17 inch
    COL_H = PAGE_H - MT - MB        # usable height   ≈ 7.02 inch

    # ── Colours ───────────────────────────────────────────────────────────────
    ACCENT_CL = colors.HexColor("#00796b")
    GREY_TEXT = colors.HexColor("#555555")
    BLACK     = colors.black

    styles = getSampleStyleSheet()

    # ── Paragraph styles ──────────────────────────────────────────────────────
    def PS(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    PartLbl = PS("PartLbl",
                 fontSize=13, fontName="Helvetica-Bold",
                 spaceAfter=2, spaceBefore=0,
                 textColor=colors.HexColor("#003366"))

    TitleSt = PS("TitleSt",
                 fontSize=12, fontName="Helvetica-Bold",
                 alignment=TA_CENTER, spaceAfter=2, spaceBefore=1)

    SubSt   = PS("SubSt",
                 fontSize=10, fontName="Helvetica-Oblique",
                 alignment=TA_CENTER, spaceAfter=1)

    SecHdr  = PS("SecHdr",
                 fontSize=9.5, fontName="Helvetica-Bold",
                 spaceAfter=2, spaceBefore=5,
                 textColor=colors.HexColor("#1a1a2e"))

    BodySt  = PS("BodySt",
                 fontSize=9, spaceAfter=2, leading=12)

    ResultSt = PS("ResultSt",
                  fontSize=11, spaceAfter=3, leading=15)

    SmallSt = PS("SmallSt",
                 fontSize=7.5, textColor=GREY_TEXT,
                 spaceAfter=0, leading=9)

    MonoSt  = PS("MonoSt",
                 fontSize=7.5, fontName="Courier",
                 spaceAfter=0, leading=9.5)

    # ── Shared APA table style (no colour — clean black & white) ─────────────
    def tbl_style(extra=None):
        base = [
            ("FONTNAME",      (0, 0), (-1,  0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1),   8.5),
            ("LEADING",       (0, 0), (-1, -1),   11),
            ("ALIGN",         (1, 0), (-1, -1),   "CENTER"),
            ("ALIGN",         (0, 0), (0,  -1),   "LEFT"),
            ("LINEABOVE",     (0, 0), (-1,  0),   1.2, BLACK),
            ("LINEBELOW",     (0, 0), (-1,  0),   0.6, BLACK),
            ("LINEBELOW",     (0,-1), (-1, -1),   1.2, BLACK),
            ("TOPPADDING",    (0, 0), (-1, -1),   2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1),   2.5),
            ("LEFTPADDING",   (0, 0), (-1, -1),   4),
            ("RIGHTPADDING",  (0, 0), (-1, -1),   4),
        ]
        if extra:
            base.extend(extra)
        return TableStyle(base)

    # ── One PageTemplate per part ─────────────────────────────────────────────
    doc = BaseDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=ML, rightMargin=MR,
        topMargin=MT,  bottomMargin=MB,
    )

    templates = []
    for i in range(len(parts)):
        frame = Frame(
            ML, MB, COL_W, COL_H,
            leftPadding=0, rightPadding=0,
            topPadding=0,  bottomPadding=0,
            id=f"p{i}_main",
        )
        templates.append(PageTemplate(id=f"part{i}", frames=[frame]))
    doc.addPageTemplates(templates)

    # ── Story ─────────────────────────────────────────────────────────────────
    story = []

    for part_idx, r in enumerate(parts):
        groups      = r.get("groups",      [])
        group_names = r.get("group_names", [])
        df_total    = (r.get("df_between", 0) or 0) + (r.get("df_within", 0) or 0)

        story.append(NextPageTemplate(f"part{part_idx}"))
        if part_idx > 0:
            story.append(PageBreak())

        # ── Part label + accent rule ──────────────────────────────────────────
        story.append(Paragraph(r.get("label", f"Part {part_idx + 1}"), PartLbl))
        story.append(HRFlowable(width=COL_W, thickness=1.5,
                                 color=ACCENT_CL, spaceAfter=3))

        # ── Title / subtitle / author ─────────────────────────────────────────
        story.append(Paragraph(
            r.get("report_title", "ANOVA ANALYSIS RESULTS"), TitleSt))
        if r.get("report_subtitle"):
            story.append(Paragraph(r["report_subtitle"], SubSt))
        if r.get("researcher_name"):
            story.append(Paragraph(r["researcher_name"], SubSt))
        story.append(Spacer(1, 4))

        # ── Table 1 — Descriptive Statistics ─────────────────────────────────
        story.append(Paragraph("Table 1. Descriptive Statistics for Groups", SecHdr))
        t1_data = [["Group", "n", "M", "SD"]]
        for name, g in zip(group_names, groups):
            t1_data.append([
                name,
                str(len(g)),
                _fmt(np.mean(g)),
                _fmt(np.std(g, ddof=1)),
            ])
        cw1 = [COL_W * 0.46, COL_W * 0.14, COL_W * 0.20, COL_W * 0.20]
        t1  = Table(t1_data, colWidths=cw1)
        t1.setStyle(tbl_style())
        story.append(t1)
        story.append(Spacer(1, 5))

        # ── Table 2 — ANOVA Summary ───────────────────────────────────────────
        story.append(Paragraph(
            "Table 2. " + r.get("anova_table_title", "Analysis of Variance Summary"), SecHdr))
        t2_data = [
            ["Source", "SS", "df", "MS", "F", "p"],
            [
                r.get("anova_between_label", "Between Groups"),
                _fmt(r.get("SS_between")),
                str(r.get("df_between", "")),
                _fmt(r.get("MS_between")),
                _fmt(r.get("F_statistic")),
                _fmt(r.get("p_value")),
            ],
            [
                r.get("anova_within_label", "Within Groups"),
                _fmt(r.get("SS_within")),
                str(r.get("df_within", "")),
                _fmt(r.get("MS_within")),
                "", "",
            ],
            [
                r.get("anova_total_label", "Total"),
                _fmt(r.get("SS_total")),
                str(df_total),
                "", "", "",
            ],
        ]
        cw2 = [COL_W*0.28, COL_W*0.16, COL_W*0.10,
               COL_W*0.16, COL_W*0.15, COL_W*0.15]
        t2  = Table(t2_data, colWidths=cw2)
        t2.setStyle(tbl_style())
        story.append(t2)
        story.append(Spacer(1, 5))

        # ── Test Results ──────────────────────────────────────────────────────
        story.append(Paragraph("Test Results", SecHdr))
        if r.get("edited") and r.get("conclusion_text"):
            decision_text = r["conclusion_text"].replace("\n", "<br/>")
            story.append(Paragraph(decision_text, BodySt))
        else:
            # F and p line — larger bold values
            story.append(Paragraph(
                f'<font size="12"><b>F</b></font>({r.get("df_between")}, {r.get("df_within")}) = '
                f'<font size="13"><b>{_fmt(r.get("F_statistic"))}</b></font>,&nbsp;&nbsp;'
                f'<font size="12"><b>p</b></font> = '
                f'<font size="13"><b>{_fmt(r.get("p_value"))}</b></font>,&nbsp;&nbsp;'
                f'&#945; = {r.get("alpha", 0.05)}',
                ResultSt))
            # Decision line — bold and prominent
            story.append(Paragraph(
                f'<font size="12"><b>Decision:&nbsp;</b></font>'
                f'<font size="12"><b>{r.get("decision", "")}</b></font>',
                ResultSt))
            # Conclusion sentence — normal body size
            story.append(Paragraph(r.get("conclusion", ""), BodySt))
        story.append(Spacer(1, 5))

        # ── Post Hoc (Tukey HSD) ──────────────────────────────────────────────
        if r.get("is_significant") and r.get("tukey"):
            story.append(Paragraph(
                r.get("posthoc_title", "Post Hoc Comparisons — Tukey HSD"), SecHdr))
            tukey_text = r.get("posthoc_text", str(r["tukey"]))
            for line in tukey_text.split("\n"):
                if line.strip():
                    story.append(Paragraph(line, MonoSt))
            story.append(Spacer(1, 4))

        # ── Raw Data Table ────────────────────────────────────────────────────
        story.append(Paragraph(
            r.get("rawdata_table_title", "Raw Data by Group"), SecHdr))

        MAX_VAL = 55
        t3_data = [["Group", "n", "Values"]]
        if r.get("raw_data_edits"):
            for edit in r["raw_data_edits"]:
                vt = edit["values_text"]
                if len(vt) > MAX_VAL:
                    vt = vt[:MAX_VAL - 1] + "…"
                t3_data.append([
                    edit["group_name"],
                    str(len(edit["values_text"].split(","))),
                    vt,
                ])
        else:
            for name, g in zip(group_names, groups):
                vs = ", ".join(_fmt(v) for v in g)
                if len(vs) > MAX_VAL:
                    vs = vs[:MAX_VAL - 1] + "…"
                t3_data.append([name, str(len(g)), vs])

        cw3 = [COL_W * 0.22, COL_W * 0.08, COL_W * 0.70]
        t3  = Table(t3_data, colWidths=cw3)
        t3.setStyle(tbl_style())
        story.append(t3)
        story.append(Spacer(1, 5))

        # ── Footer ────────────────────────────────────────────────────────────
        story.append(HRFlowable(width=COL_W, thickness=0.5,
                                 color=colors.grey, spaceAfter=2))
        story.append(Paragraph(
            f"Part saved: {r.get('saved_at', '—')}  &nbsp;|&nbsp;  "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}  "
            f"&nbsp;|&nbsp;  Part {part_idx + 1} of {len(parts)}",
            SmallSt,
        ))

    doc.build(story)