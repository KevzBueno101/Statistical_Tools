#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cohen's Kappa Desktop Calculator - Modern CustomTkinter Edition
APA Format with full statistical output, SE/CI, and interpretation panel.

Design: Refined dark-professional aesthetic with teal accents
- Sidebar navigation with icon buttons
- Card-based layout with elevation shadows
- Color-coded interpretation badge
- Smooth tab switching
"""

from __future__ import annotations

import os
import sys
from datetime import datetime

import numpy as np
import pandas as pd

try:
    import customtkinter as ctk
    from tkinter import filedialog, messagebox
    import tkinter as tk
except Exception as e:
    raise SystemExit(f"CustomTkinter is required: pip install customtkinter\n{e}")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as PLTable, TableStyle
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    from reportlab.lib import colors
    from reportlab.lib.units import inch
except Exception:
    raise SystemExit("ReportLab required: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn as QN
except Exception:
    raise SystemExit("python-docx required: pip install python-docx")


# ─────────────────────────────────────────────
#  THEME CONSTANTS
# ─────────────────────────────────────────────

TEAL      = "#0D9488"
TEAL_DK   = "#0F766E"
TEAL_LT   = "#14B8A6"
SLATE_900 = "#0F172A"
SLATE_800 = "#1E293B"
SLATE_700 = "#334155"
SLATE_600 = "#475569"
SLATE_400 = "#94A3B8"
SLATE_200 = "#E2E8F0"
WHITE     = "#F8FAFC"

KAPPA_COLORS = {
    "Poor":           ("#EF4444", "#FEE2E2"),
    "Slight":         ("#F97316", "#FFF7ED"),
    "Fair":           ("#EAB308", "#FEFCE8"),
    "Moderate":       ("#3B82F6", "#EFF6FF"),
    "Substantial":    ("#8B5CF6", "#F5F3FF"),
    "Almost Perfect": ("#10B981", "#ECFDF5"),
    "Perfect":        ("#10B981", "#ECFDF5"),
    "Undefined":      ("#6B7280", "#F9FAFB"),
}


# ─────────────────────────────────────────────
#  STATISTICS ENGINE
# ─────────────────────────────────────────────

class KappaStats:
    @staticmethod
    def compute_kappa(rater_a, rater_b, categories=None):
        ra = pd.Series(rater_a).astype(object)
        rb = pd.Series(rater_b).astype(object)

        if len(ra) != len(rb):
            raise ValueError("Rater A and Rater B must have the same number of ratings.")

        N = len(ra)
        if N < 2:
            raise ValueError("At least 2 ratings are required.")

        if categories is None:
            cats = sorted(set(ra.dropna().tolist()) | set(rb.dropna().tolist()))
            if not cats:
                raise ValueError("No valid rating categories found.")
        else:
            cats = list(categories)

        conf = pd.crosstab(ra, rb, rownames=["Rater A"], colnames=["Rater B"], dropna=False)
        for c in cats:
            if c not in conf.index:   conf.loc[c] = 0
            if c not in conf.columns: conf[c] = 0
        conf = conf.loc[cats, cats]

        Po  = float(np.trace(conf.values)) / float(N)
        rt  = conf.sum(axis=1).astype(float)
        ct  = conf.sum(axis=0).astype(float)
        Pe  = float(rt.values @ ct.values) / float(N * N)

        denom = 1.0 - Pe
        kappa = (Po - Pe) / denom if denom != 0 else 1.0

        se    = np.sqrt(Po * (1 - Po) / N)
        lower = max(-1.0, kappa - 1.96 * se)
        upper = min( 1.0, kappa + 1.96 * se)

        interp = KappaStats._interpret(kappa)
        details = {"confusion_matrix": conf.values, "categories": cats,
                   "Po": Po, "Pe": Pe, "N": N}

        return kappa, Po, Pe, N, interp, details, se, lower, upper

    @staticmethod
    def _interpret(k):
        if np.isnan(k): return "Undefined"
        k = np.clip(k, -1, 1)
        if k < 0.00: return "Poor"
        if k < 0.20: return "Slight"
        if k < 0.40: return "Fair"
        if k < 0.60: return "Moderate"
        if k < 0.80: return "Substantial"
        return "Almost Perfect"


# ─────────────────────────────────────────────
#  PDF EXPORT (APA)
# ─────────────────────────────────────────────

class APAPDFExporter:
    def __init__(self, path): self.path = path

    def generate(self, results, rater_name):
        doc = SimpleDocTemplate(self.path, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=54)
        story = []
        styles = getSampleStyleSheet()

        title_s = ParagraphStyle("T", parent=styles["Heading1"],
                                 fontSize=13, alignment=TA_CENTER,
                                 fontName="Helvetica-Bold", spaceAfter=4)
        sub_s   = ParagraphStyle("S", parent=styles["Normal"],
                                 fontSize=10, alignment=TA_CENTER,
                                 fontName="Helvetica", spaceAfter=20,
                                 textColor=colors.HexColor("#475569"))
        h2_s    = ParagraphStyle("H2", parent=styles["Heading2"],
                                 fontSize=12, fontName="Helvetica-Bold",
                                 spaceAfter=8, spaceBefore=14)
        body_s  = ParagraphStyle("B", parent=styles["Normal"],
                                 fontSize=8, fontName="Helvetica",
                                 spaceAfter=6, leading=16)
        note_s  = ParagraphStyle("N", parent=styles["Normal"],
                                 fontSize=7, fontName="Helvetica-Oblique",
                                 spaceAfter=6, textColor=colors.HexColor("#475569"))

        story.append(Paragraph("Cohen's Kappa Analysis", title_s))
        story.append(Paragraph(f"<b>Rater:</b> {rater_name}", h2_s))

        kappa = results["kappa"]; se = results["se"]
        lower = results["ci_lower"]; upper = results["ci_upper"]
        Po = results["Po"]; Pe = results["Pe"]; N = results["N"]
        interp = results["interpretation"]

        story.append(Paragraph("<i>Cohen's Kappa Reliability Statistics</i>",
                               ParagraphStyle("TT", parent=styles["Normal"],
                                              fontSize=11, fontName="Helvetica-Oblique",
                                              spaceAfter=6)))

        td = [
            ["", "", "", "95% CI", ""],
            ["", "κ", "SE", "Lower", "Upper"],
            ["Average kappa", f"{kappa:.3f}", f"{se:.3f}", f"{lower:.3f}", f"{upper:.3f}"],
            ["Pre-test – Post-test", f"{kappa:.3f}", f"{se:.3f}", f"{lower:.3f}", f"{upper:.3f}"],
        ]

        t = PLTable(td, colWidths=[2.2*inch, 0.8*inch, 0.7*inch, 0.7*inch, 0.7*inch])
        t.setStyle(TableStyle([
            ("SPAN",        (3, 0), (4, 0)),
            ("LINEABOVE",   (0, 0), (-1, 0),  1.5, colors.black),
            ("LINEBELOW",   (0, 1), (-1, 1),  1.0, colors.black),
            ("LINEBELOW",   (0,-1), (-1,-1),  1.5, colors.black),
            ("FONTNAME",    (0, 0), (-1, 1),  "Helvetica-Bold"),
            ("FONTNAME",    (1, 1), (1, 1),   "Helvetica-BoldOblique"),
            ("ALIGN",       (0, 0), (0, -1),  "LEFT"),
            ("ALIGN",       (1, 0), (-1, -1), "CENTER"),
            ("FONTSIZE",    (0, 0), (-1, -1), 10),
            ("TOPPADDING",  (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING",(0,0), (-1, -1), 5),
            ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.15*inch))
        story.append(Paragraph(
            f"<i>Note.</i> κ = Cohen's Kappa; SE = Standard Error; 95% CI = Confidence Interval. "
            f"N = {N}.", note_s))

        story.append(Spacer(1, 0.06*inch))
        story.append(Paragraph("Summary Statistics", h2_s))

        summary = [
            ["Statistic", "Value"],
            ["Cohen's Kappa (κ)", f"{kappa:.4f}"],
            ["Standard Error", f"{se:.4f}"],
            ["95% CI Lower", f"{lower:.4f}"],
            ["95% CI Upper", f"{upper:.4f}"],
            ["Observed Agreement (Po)", f"{Po:.4f}"],
            ["Expected Agreement (Pe)", f"{Pe:.4f}"],
            ["N (Observations)", str(N)],
            ["Interpretation", interp],
        ]
        st = PLTable(summary, colWidths=[2.8*inch, 2.0*inch])
        st.setStyle(TableStyle([
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("LINEABOVE",    (0, 0), (-1, 0),  1.2, colors.black),
            ("LINEBELOW",    (0, 0), (-1, 0),  1.0, colors.black),
            ("LINEBELOW",    (0,-1), (-1,-1),  1.2, colors.black),
            ("FONTSIZE",     (0, 0), (-1, -1), 10),
            ("TOPPADDING",   (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ]))
        story.append(st)

        story.append(Spacer(1, 0.25*inch))
        story.append(Paragraph("Interpretation Guide", h2_s))
        guide = (
            "<b>Kappa Interpretation Scale (Landis & Koch, 1977):</b><br/>"
            "• κ &lt; 0.00 : Poor agreement<br/>"
            "• 0.00 – 0.20 : Slight agreement<br/>"
            "• 0.21 – 0.40 : Fair agreement<br/>"
            "• 0.41 – 0.60 : Moderate agreement<br/>"
            "• 0.61 – 0.80 : Substantial agreement<br/>"
            "• 0.81 – 1.00 : Almost perfect agreement"
        )
        story.append(Paragraph(guide, body_s))

        ts = datetime.now().strftime("%B %d, %Y at %H:%M:%S")
        story.append(Spacer(0.5, 0.04*inch))
        story.append(Paragraph(
            f"Generated: {ts} | File: {os.path.abspath(self.path)}",
            ParagraphStyle("F", parent=styles["Italic"], fontSize=5,
                           fontName="Helvetica-Oblique",
                           textColor=colors.HexColor("#94A3B8"))))
        doc.build(story)


# ─────────────────────────────────────────────
#  DOCX EXPORT (APA)
# ─────────────────────────────────────────────

class APADOCXExporter:
    def __init__(self, path): self.path = path

    def generate(self, results, rater_name):
        doc = Document()

        def set_border(cell, top=None, bottom=None):
            tc   = cell._element
            tcPr = tc.get_or_add_tcPr()
            for b in ["top","left","bottom","right"]:
                ex = tcPr.find(QN(f"w:{b}"))
                if ex is not None: tcPr.remove(ex)
            tcB = OxmlElement("w:tcBorders")
            for name, val in [("top", top), ("bottom", bottom),
                               ("left", "nil"), ("right", "nil")]:
                el = OxmlElement(f"w:{name}")
                if val == "nil" or val is None:
                    el.set(QN("w:val"), "nil")
                else:
                    el.set(QN("w:val"), val)
                    el.set(QN("w:sz"), "18" if val == "thick" else "12")
                    el.set(QN("w:color"), "000000")
                tcB.append(el)
            tcPr.append(tcB)

        def cell_text(cell, text, right=False, bold=False, italic=False):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic

        # Title
        t = doc.add_heading("Cohen's Kappa Analysis", 0)
        doc.add_paragraph("Inter-Rater Reliability — APA 7th Edition Format")
        doc.add_heading(f"{rater_name}", level=2)

        kappa = results["kappa"]; se = results["se"]
        lower = results["ci_lower"]; upper = results["ci_upper"]
        Po = results["Po"]; Pe = results["Pe"]; N = results["N"]
        interp = results["interpretation"]

        doc.add_paragraph("Cohen's Kappa Reliability Statistics").runs[0].italic = True

        tbl = doc.add_table(rows=4, cols=5)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Row 0 - header
        r0 = tbl.rows[0].cells
        cell_text(r0[0], "Cohen's κ", bold=True, italic=True)
        cell_text(r0[3], "95% CI", bold=True)
        r0[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in tbl.rows[0].cells: set_border(c, top="thick", bottom="nil")

        # Row 1 - subheader
        r1 = tbl.rows[1].cells
        cell_text(r1[1], "κ",     bold=True, italic=True, right=True)
        cell_text(r1[2], "SE",    bold=True, right=True)
        cell_text(r1[3], "Lower", bold=True, right=True)
        cell_text(r1[4], "Upper", bold=True, right=True)
        for c in tbl.rows[1].cells: set_border(c, top="nil", bottom="single")

        # Row 2
        r2 = tbl.rows[2].cells
        cell_text(r2[0], "Average kappa")
        cell_text(r2[1], f"{kappa:.3f}", right=True)
        cell_text(r2[2], f"{se:.3f}",   right=True)
        cell_text(r2[3], f"{lower:.3f}", right=True)
        cell_text(r2[4], f"{upper:.3f}", right=True)
        for c in r2: set_border(c, top="nil", bottom="nil")

        # Row 3
        r3 = tbl.rows[3].cells
        cell_text(r3[0], "Pre-test – Post-test")
        cell_text(r3[1], f"{kappa:.3f}", right=True)
        cell_text(r3[2], f"{se:.3f}",   right=True)
        cell_text(r3[3], f"{lower:.3f}", right=True)
        cell_text(r3[4], f"{upper:.3f}", right=True)
        for c in r3: set_border(c, top="nil", bottom="thick")

        doc.add_paragraph(
            f"Note. κ = Cohen's Kappa; SE = Standard Error; 95% CI = Confidence Interval. N = {N}."
        ).runs[0].italic = True

        doc.add_heading("Summary Statistics", level=2)
        rows = [
            ("Cohen's Kappa (κ)",        f"{kappa:.4f}"),
            ("Standard Error",           f"{se:.4f}"),
            ("95% CI",                   f"[{lower:.4f}, {upper:.4f}]"),
            ("Observed Agreement (Po)",  f"{Po:.4f}"),
            ("Expected Agreement (Pe)",  f"{Pe:.4f}"),
            ("N (Observations)",         str(N)),
            ("Interpretation",           interp),
        ]
        st = doc.add_table(rows=len(rows)+1, cols=2)
        st.style = "Table Grid"
        for i, (label, val) in enumerate(rows):
            st.rows[i+1].cells[0].text = label
            st.rows[i+1].cells[1].text = val
        st.rows[0].cells[0].text = "Statistic"
        st.rows[0].cells[1].text = "Value"

        doc.add_heading("Interpretation Guide", level=2)
        doc.add_paragraph(
            "κ < 0.00: Poor  |  0.00–0.20: Slight  |  0.21–0.40: Fair  |  "
            "0.41–0.60: Moderate  |  0.61–0.80: Substantial  |  0.81–1.00: Almost Perfect"
        )
        doc.save(self.path)


# ─────────────────────────────────────────────
#  MODERN UI COMPONENTS
# ─────────────────────────────────────────────

class SectionCard(ctk.CTkFrame):
    """Elevated card container."""
    def __init__(self, master, title="", **kwargs):
        kwargs.setdefault("corner_radius", 12)
        kwargs.setdefault("fg_color", ("gray95", SLATE_800))
        super().__init__(master, **kwargs)
        if title:
            ctk.CTkLabel(self, text=title,
                         font=ctk.CTkFont("Georgia", 13, "bold"),
                         text_color=(SLATE_700, SLATE_200)
                         ).pack(anchor="w", padx=16, pady=(12, 4))


class BadgeLabel(ctk.CTkFrame):
    """Colored interpretation badge."""
    def __init__(self, master, text="—", interpretation="Undefined", **kwargs):
        fg, bg = KAPPA_COLORS.get(interpretation, ("#6B7280", "#F9FAFB"))
        super().__init__(master, fg_color=bg, corner_radius=20, **kwargs)
        self._lbl = ctk.CTkLabel(self, text=text,
                                  font=ctk.CTkFont("Georgia", 14, "bold"),
                                  text_color=fg)
        self._lbl.pack(padx=20, pady=6)

    def update_badge(self, text, interpretation):
        fg, bg = KAPPA_COLORS.get(interpretation, ("#6B7280", "#F9FAFB"))
        self.configure(fg_color=bg)
        self._lbl.configure(text=text, text_color=fg)


class StatCard(ctk.CTkFrame):
    """Mini stat display card."""
    def __init__(self, master, label="", value="—", **kwargs):
        kwargs.setdefault("corner_radius", 10)
        kwargs.setdefault("fg_color", ("white", SLATE_700))
        super().__init__(master, **kwargs)
        ctk.CTkLabel(self, text=label,
                     font=ctk.CTkFont("Helvetica", 10),
                     text_color=(SLATE_600, SLATE_400)
                     ).pack(pady=(10, 0))
        self.val_lbl = ctk.CTkLabel(self, text=value,
                                     font=ctk.CTkFont("Georgia", 18, "bold"),
                                     text_color=(SLATE_900, WHITE))
        self.val_lbl.pack(pady=(0, 10))

    def set_value(self, v): self.val_lbl.configure(text=v)


# ─────────────────────────────────────────────
#  SCROLLABLE DATA GRID
# ─────────────────────────────────────────────

class DataGrid(ctk.CTkScrollableFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)
        self.entries: list[tuple[ctk.CTkEntry, ctk.CTkEntry]] = []
        self._build_header()
        self.add_rows(8)

    def _build_header(self):
        hdr = ctk.CTkFrame(self, fg_color=(SLATE_200, SLATE_700), corner_radius=8)
        hdr.pack(fill="x", pady=(0, 4))
        ctk.CTkLabel(hdr, text="#",        width=40,  font=ctk.CTkFont("Helvetica", 11, "bold")).pack(side="left", padx=4, pady=6)
        ctk.CTkLabel(hdr, text="Rater A",  width=180, font=ctk.CTkFont("Helvetica", 11, "bold")).pack(side="left", padx=4)
        ctk.CTkLabel(hdr, text="Rater B",  width=180, font=ctk.CTkFont("Helvetica", 11, "bold")).pack(side="left", padx=4)

    def add_rows(self, n=1):
        for _ in range(n):
            idx = len(self.entries) + 1
            row = ctk.CTkFrame(self, fg_color="transparent")
            row.pack(fill="x", pady=1)
            ctk.CTkLabel(row, text=str(idx), width=40,
                         font=ctk.CTkFont("Helvetica", 10),
                         text_color=(SLATE_400, SLATE_400)).pack(side="left", padx=4)
            e1 = ctk.CTkEntry(row, width=180, height=30,
                               font=ctk.CTkFont("Helvetica", 11),
                               corner_radius=6)
            e1.pack(side="left", padx=4, pady=2)
            e2 = ctk.CTkEntry(row, width=180, height=30,
                               font=ctk.CTkFont("Helvetica", 11),
                               corner_radius=6)
            e2.pack(side="left", padx=4, pady=2)
            self.entries.append((e1, e2))

    def get_data(self) -> pd.DataFrame:
        rows = []
        for e1, e2 in self.entries:
            a, b = e1.get().strip(), e2.get().strip()
            if a or b:
                rows.append((a, b))
        if not rows:
            raise ValueError("No data entered.")
        return pd.DataFrame(rows, columns=["Rater A", "Rater B"])

    def populate(self, df: pd.DataFrame):
        for e1, e2 in self.entries:
            e1.delete(0, "end"); e2.delete(0, "end")
        needed = max(0, len(df) - len(self.entries))
        if needed: self.add_rows(needed)
        for i, (_, row) in enumerate(df.iterrows()):
            self.entries[i][0].delete(0, "end"); self.entries[i][0].insert(0, str(row.iloc[0]))
            self.entries[i][1].delete(0, "end"); self.entries[i][1].insert(0, str(row.iloc[1]))

    def clear(self):
        for e1, e2 in self.entries:
            e1.delete(0, "end"); e2.delete(0, "end")


# ─────────────────────────────────────────────
#  MAIN APPLICATION
# ─────────────────────────────────────────────

class KappaApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Cohen's Kappa Calculator")
        self.geometry("1100x780")
        self.minsize(900, 650)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("green")

        self.results: dict = {}
        self._build_ui()

    # ── UI CONSTRUCTION ──────────────────────

    def _build_ui(self):
        # Root grid
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self._build_sidebar()
        self._build_main()

    def _build_sidebar(self):
        sb = ctk.CTkFrame(self, width=220, corner_radius=0,
                           fg_color=(SLATE_200, SLATE_900))
        sb.grid(row=0, column=0, sticky="nsew")
        sb.grid_propagate(False)

        # Logo area
        logo = ctk.CTkFrame(sb, fg_color=(TEAL, TEAL_DK), corner_radius=0, height=80)
        logo.pack(fill="x")
        logo.pack_propagate(False)
        ctk.CTkLabel(logo, text="κ",
                     font=ctk.CTkFont("Georgia", 42, "bold"),
                     text_color="white").pack(expand=True)

        ctk.CTkLabel(sb, text="Cohen's Kappa",
                     font=ctk.CTkFont("Georgia", 15, "bold"),
                     text_color=(SLATE_800, WHITE)).pack(pady=(16, 2))
        ctk.CTkLabel(sb, text="APA Format Calculator",
                     font=ctk.CTkFont("Helvetica", 10),
                     text_color=(SLATE_600, SLATE_400)).pack()

        ctk.CTkFrame(sb, height=1, fg_color=(SLATE_400, SLATE_700)).pack(fill="x", padx=16, pady=16)

        # Rater name
        ctk.CTkLabel(sb, text="RATER NAME",
                     font=ctk.CTkFont("Helvetica", 10, "bold"),
                     text_color=(SLATE_600, SLATE_400)).pack(anchor="w", padx=16)
        self.rater_entry = ctk.CTkEntry(sb, height=36,
                                         font=ctk.CTkFont("Helvetica", 12),
                                         placeholder_text="e.g. Maria Surban",
                                         corner_radius=8)
        self.rater_entry.pack(fill="x", padx=16, pady=(4, 16))
        self.rater_entry.insert(0, "Name")

        # Action buttons
        btn_cfg = dict(height=40, font=ctk.CTkFont("Helvetica", 12, "bold"),
                       corner_radius=8)

        ctk.CTkButton(sb, text="📁  Import CSV/XLSX",
                       command=self._import_file,
                       fg_color=(SLATE_700, SLATE_700),
                       hover_color=(SLATE_600, SLATE_600),
                       **btn_cfg).pack(fill="x", padx=16, pady=3)

        ctk.CTkButton(sb, text="➕  Add Row",
                       command=lambda: self.grid.add_rows(1),
                       fg_color=(SLATE_700, SLATE_700),
                       hover_color=(SLATE_600, SLATE_600),
                       **btn_cfg).pack(fill="x", padx=16, pady=3)

        ctk.CTkButton(sb, text="🗑  Clear / Reset",
                       command=self._reset,
                       fg_color=(SLATE_700, SLATE_700),
                       hover_color=(SLATE_600, SLATE_600),
                       **btn_cfg).pack(fill="x", padx=16, pady=3)

        ctk.CTkFrame(sb, height=1, fg_color=(SLATE_400, SLATE_700)).pack(fill="x", padx=16, pady=16)

        ctk.CTkButton(sb, text="📊  Compute Kappa",
                       command=self._compute,
                       fg_color=TEAL, hover_color=TEAL_DK,
                       **btn_cfg).pack(fill="x", padx=16, pady=3)

        self.pdf_btn = ctk.CTkButton(sb, text="📄  Export PDF",
                                      command=self._export_pdf,
                                      state="disabled",
                                      fg_color=(SLATE_700, SLATE_700),
                                      hover_color=(SLATE_600, SLATE_600),
                                      **btn_cfg)
        self.pdf_btn.pack(fill="x", padx=16, pady=3)

        self.docx_btn = ctk.CTkButton(sb, text="📝  Export DOCX",
                                       command=self._export_docx,
                                       state="disabled",
                                       fg_color=(SLATE_700, SLATE_700),
                                       hover_color=(SLATE_600, SLATE_600),
                                       **btn_cfg)
        self.docx_btn.pack(fill="x", padx=16, pady=3)

        # Theme toggle
        ctk.CTkFrame(sb, height=1, fg_color=(SLATE_400, SLATE_700)).pack(fill="x", padx=16, pady=16)
        self._theme_mode = "dark"
        self._theme_btn = ctk.CTkButton(sb, text="☀  Light Mode",
                                         command=self._toggle_theme,
                                         fg_color="transparent",
                                         hover_color=(SLATE_300, SLATE_700) if False else (SLATE_700, SLATE_700),
                                         border_width=1,
                                         **btn_cfg)
        self._theme_btn.pack(fill="x", padx=16, pady=3)

        # Status bar
        self.status_lbl = ctk.CTkLabel(sb, text="Ready",
                                        font=ctk.CTkFont("Helvetica", 10),
                                        text_color=(SLATE_600, SLATE_400),
                                        wraplength=190)
        self.status_lbl.pack(side="bottom", padx=16, pady=12)

    def _build_main(self):
        main = ctk.CTkFrame(self, fg_color=("gray88", SLATE_800), corner_radius=0)
        main.grid(row=0, column=1, sticky="nsew")
        main.grid_columnconfigure(0, weight=3)
        main.grid_columnconfigure(1, weight=2)
        main.grid_rowconfigure(0, weight=1)

        # ── LEFT: Data Input ──────────────────
        left = ctk.CTkFrame(main, fg_color="transparent")
        left.grid(row=0, column=0, sticky="nsew", padx=(16,8), pady=16)
        left.grid_rowconfigure(1, weight=1)
        left.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(left, text="Data Input",
                     font=ctk.CTkFont("Georgia", 20, "bold"),
                     text_color=(SLATE_900, WHITE)).grid(row=0, column=0, sticky="w", pady=(0,10))

        grid_card = SectionCard(left, title="Ratings Table")
        grid_card.grid(row=1, column=0, sticky="nsew")

        self.grid = DataGrid(grid_card,
                              fg_color=("white", SLATE_800),
                              corner_radius=8)
        self.grid.pack(fill="both", expand=True, padx=12, pady=(0,12))

        # ── RIGHT: Results (pure pack layout) ─────────────────────
        right = ctk.CTkFrame(main, fg_color="transparent")
        right.grid(row=0, column=1, sticky="nsew", padx=(8,16), pady=16)

        ctk.CTkLabel(right, text="Results",
                     font=ctk.CTkFont("Georgia", 20, "bold"),
                     text_color=(SLATE_900, WHITE)).pack(anchor="w", pady=(0,10))

        # Kappa badge card
        badge_card = SectionCard(right, title="Kappa Coefficient")
        badge_card.pack(fill="x", pady=(0,8))

        self.badge = BadgeLabel(badge_card, text="—", interpretation="Undefined")
        self.badge.pack(pady=(4,4))

        interp_row = ctk.CTkFrame(badge_card, fg_color="transparent")
        interp_row.pack(fill="x", padx=16, pady=(0,12))
        ctk.CTkLabel(interp_row, text="Interpretation:",
                     font=ctk.CTkFont("Helvetica", 11),
                     text_color=(SLATE_600, SLATE_400)).pack(side="left")
        self.interp_lbl = ctk.CTkLabel(interp_row, text="—",
                                        font=ctk.CTkFont("Georgia", 11, "bold"),
                                        text_color=(TEAL_LT, TEAL_LT))
        self.interp_lbl.pack(side="left", padx=6)

        # Stat cards row (grid inside its own frame — no conflict)
        stats_frame = ctk.CTkFrame(right, fg_color="transparent")
        stats_frame.pack(fill="x", pady=(0,8))
        for i in range(4): stats_frame.grid_columnconfigure(i, weight=1)

        self.sc_se = StatCard(stats_frame, "Std Error", "—")
        self.sc_lo = StatCard(stats_frame, "CI Lower",  "—")
        self.sc_hi = StatCard(stats_frame, "CI Upper",  "—")
        self.sc_n  = StatCard(stats_frame, "N",         "—")
        self.sc_se.grid(row=0, column=0, sticky="ew", padx=(0,4))
        self.sc_lo.grid(row=0, column=1, sticky="ew", padx=2)
        self.sc_hi.grid(row=0, column=2, sticky="ew", padx=2)
        self.sc_n.grid (row=0, column=3, sticky="ew", padx=(4,0))

        # Interpretation guide card (pack before detail so detail expands)
        guide_card = SectionCard(right, title="Interpretation Guide (Landis & Koch, 1977)")
        guide_card.pack(fill="x", side="bottom", pady=(8,0))

        guide_data = [
            ("< 0.00",    "Poor",           "#EF4444"),
            ("0.00–0.20", "Slight",         "#F97316"),
            ("0.21–0.40", "Fair",           "#EAB308"),
            ("0.41–0.60", "Moderate",       "#3B82F6"),
            ("0.61–0.80", "Substantial",    "#8B5CF6"),
            ("0.81–1.00", "Almost Perfect", "#10B981"),
        ]
        guide_inner = ctk.CTkFrame(guide_card, fg_color="transparent")
        guide_inner.pack(fill="x", padx=12, pady=(0,12))
        for rng, label, clr in guide_data:
            col = ctk.CTkFrame(guide_inner, fg_color="transparent")
            col.pack(side="left", expand=True, fill="x")
            ctk.CTkLabel(col, text=rng,
                         font=ctk.CTkFont("Courier New", 9),
                         text_color=(SLATE_600, SLATE_400)).pack()
            ctk.CTkFrame(col, width=12, height=12,
                         corner_radius=6, fg_color=clr).pack(pady=1)
            ctk.CTkLabel(col, text=label,
                         font=ctk.CTkFont("Helvetica", 9, "bold"),
                         text_color=(SLATE_700, SLATE_200)).pack()

        # Detail textbox card (expands to fill remaining space)
        detail_card = SectionCard(right, title="Detailed Output")
        detail_card.pack(fill="both", expand=True, pady=(0,8))

        self.detail_box = ctk.CTkTextbox(detail_card,
                                          font=ctk.CTkFont("Courier New", 11),
                                          wrap="word",
                                          fg_color=("white", SLATE_900),
                                          corner_radius=8)
        self.detail_box.pack(fill="both", expand=True, padx=12, pady=(0,12))
        self._reset_detail_text()

    # ── ACTIONS ──────────────────────────────

    def _reset_detail_text(self):
        self.detail_box.configure(state="normal")
        self.detail_box.delete("1.0", "end")
        self.detail_box.insert("1.0",
            "╔══════════════════════════════════╗\n"
            "║   Cohen's Kappa Calculator       ║\n"
            "║   APA Format Edition             ║\n"
            "╚══════════════════════════════════╝\n\n"
            "Steps:\n"
            "  1. Enter rater name (sidebar)\n"
            "  2. Fill in Rater A / Rater B data\n"
            "     or import a CSV/XLSX file\n"
            "  3. Click  Compute Kappa\n"
            "  4. Export PDF or DOCX report\n\n"
            "Waiting for data...")
        self.detail_box.configure(state="disabled")

    def _set_status(self, msg): self.status_lbl.configure(text=msg)

    def _toggle_theme(self):
        if self._theme_mode == "dark":
            ctk.set_appearance_mode("light")
            self._theme_mode = "light"
            self._theme_btn.configure(text="🌙  Dark Mode")
        else:
            ctk.set_appearance_mode("dark")
            self._theme_mode = "dark"
            self._theme_btn.configure(text="☀  Light Mode")

    def _import_file(self):
        path = filedialog.askopenfilename(
            title="Import Data (2 columns: Rater A, Rater B)",
            filetypes=[("CSV/XLSX", "*.csv *.xlsx *.xls"), ("All files", "*.*")])
        if not path: return
        try:
            df = pd.read_csv(path, header=None) if path.lower().endswith(".csv") \
                 else pd.read_excel(path, header=None)
            if df.shape[1] < 2:
                raise ValueError("File must have at least 2 columns.")
            df = df.iloc[:, :2].copy()
            df.columns = ["Rater A", "Rater B"]
            self.grid.populate(df)
            self._set_status(f"Imported {len(df)} rows from\n{os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Import Error", str(e))

    def _reset(self):
        self.grid.clear()
        self.results = {}
        self.badge.update_badge("—", "Undefined")
        self.interp_lbl.configure(text="—")
        for sc in (self.sc_se, self.sc_lo, self.sc_hi, self.sc_n):
            sc.set_value("—")
        self._reset_detail_text()
        self.pdf_btn.configure(state="disabled")
        self.docx_btn.configure(state="disabled")
        self._set_status("Reset complete")

    def _compute(self):
        try:
            df = self.grid.get_data()
            ra, rb = df["Rater A"], df["Rater B"]
            cats = sorted(set(ra.dropna()) | set(rb.dropna()))
            kappa, Po, Pe, N, interp, details, se, lower, upper = \
                KappaStats.compute_kappa(ra.tolist(), rb.tolist(), cats)

            self.results = dict(kappa=kappa, Po=Po, Pe=Pe, N=N,
                                interpretation=interp, se=se,
                                ci_lower=lower, ci_upper=upper, details=details)

            # Update badge
            self.badge.update_badge(f"κ = {kappa:.4f}", interp)
            self.interp_lbl.configure(text=interp)

            # Update stat cards
            self.sc_se.set_value(f"{se:.4f}")
            self.sc_lo.set_value(f"{lower:.4f}")
            self.sc_hi.set_value(f"{upper:.4f}")
            self.sc_n.set_value(str(N))

            # Update detail box
            ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            rname = self.rater_entry.get().strip() or "Unnamed"
            detail = (
                f"╔══════════════════════════════════╗\n"
                f"║   COHEN'S KAPPA RESULTS          ║\n"
                f"╚══════════════════════════════════╝\n\n"
                f"Timestamp : {ts}\n"
                f"Rater     : {rname}\n\n"
                f"┌──────────────────────────────────┐\n"
                f"│  RELIABILITY COEFFICIENT         │\n"
                f"├──────────────────────────────────┤\n"
                f"  Cohen's κ       : {kappa:.4f}\n"
                f"  Interpretation  : {interp}\n"
                f"  Standard Error  : {se:.4f}\n"
                f"  95% CI Lower    : {lower:.4f}\n"
                f"  95% CI Upper    : {upper:.4f}\n"
                f"└──────────────────────────────────┘\n\n"
                f"┌──────────────────────────────────┐\n"
                f"│  AGREEMENT STATISTICS            │\n"
                f"├──────────────────────────────────┤\n"
                f"  Observed (Po)   : {Po:.4f}\n"
                f"  Expected (Pe)   : {Pe:.4f}\n"
                f"  N               : {N}\n"
                f"  Categories      : {', '.join(str(c) for c in details['categories'])}\n"
                f"└──────────────────────────────────┘\n\n"
                f"FORMULA:\n"
                f"  κ = (Po − Pe) / (1 − Pe)\n\n"
                f"INTERPRETATION SCALE:\n"
                f"  κ < 0.00  : Poor\n"
                f"  0.00–0.20 : Slight\n"
                f"  0.21–0.40 : Fair\n"
                f"  0.41–0.60 : Moderate\n"
                f"  0.61–0.80 : Substantial\n"
                f"  0.81–1.00 : Almost Perfect\n\n"
                f"✓ Ready to export PDF / DOCX report!"
            )
            self.detail_box.configure(state="normal")
            self.detail_box.delete("1.0", "end")
            self.detail_box.insert("1.0", detail)
            self.detail_box.configure(state="disabled")

            self.pdf_btn.configure(state="normal")
            self.docx_btn.configure(state="normal")
            self._set_status(f"Computed κ = {kappa:.4f}\n({interp})")

        except Exception as e:
            messagebox.showerror("Computation Error", str(e))
            self._set_status("Error — check data")

    def _export_pdf(self):
        if not self.results:
            messagebox.showerror("No Data", "Compute first!"); return
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=f"Kappa_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        if not path: return
        try:
            APAPDFExporter(path).generate(self.results, self.rater_entry.get().strip() or "Rater")
            self._set_status(f"PDF saved:\n{os.path.basename(path)}")
            messagebox.showinfo("Success", f"APA-formatted PDF saved!\n\n{os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("PDF Error", str(e))

    def _export_docx(self):
        if not self.results:
            messagebox.showerror("No Data", "Compute first!"); return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("DOCX", "*.docx")],
            initialfile=f"Kappa_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        if not path: return
        try:
            APADOCXExporter(path).generate(self.results, self.rater_entry.get().strip() or "Rater")
            self._set_status(f"DOCX saved:\n{os.path.basename(path)}")
            messagebox.showinfo("Success", f"APA-formatted DOCX saved!\n\n{os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("DOCX Error", str(e))


# ─────────────────────────────────────────────

def main():
    app = KappaApp()
    app.mainloop()

if __name__ == "__main__":
    main()