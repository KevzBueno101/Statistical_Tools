import customtkinter as ctk
from tkinter import messagebox, filedialog
from scipy.stats import f_oneway
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import os
from app_settings import SettingsManager, SettingsWindow

# ─── Theme ───────────────────────────────────────────────────────────────────
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")
plt.style.use('dark_background')

# Palette
BG_DEEP   = "#0d1117"
BG_CARD   = "#161b22"
BG_PANEL  = "#1c2230"
BG_INPUT  = "#1e2736"
ACCENT    = "#00c9a7"        # teal-green (matches Kappa screenshot)
ACCENT2   = "#4e9eff"        # blue highlight
DANGER    = "#ef4444"
WARN      = "#f59e0b"
SUCCESS   = "#22c55e"
PURPLE    = "#a855f7"
TEXT_PRI  = "#e6edf3"
TEXT_SEC  = "#8b949e"
BORDER    = "#30363d"

FONT_HEAD = ("Segoe UI", 22, "bold")
FONT_SUB  = ("Segoe UI", 11)
FONT_CARD = ("Segoe UI", 13, "bold")
FONT_BODY = ("Segoe UI", 11)
FONT_MONO = ("Consolas", 10)
FONT_BTN  = ("Segoe UI", 12, "bold")
FONT_TINY = ("Segoe UI", 9)


def fmt(val, decimals=2):
    """Round to 2 decimal places and return formatted string."""
    return f"{round(float(val), decimals):.{decimals}f}"


class Sidebar(ctk.CTkFrame):
    def __init__(self, master, **kw):
        super().__init__(master, width=220, fg_color=BG_CARD,
                         corner_radius=0, **kw)
        self.pack_propagate(False)
        self._build()

    def _build(self):
        # Logo area
        self.logo_frame = ctk.CTkFrame(self, fg_color=ACCENT, corner_radius=0, height=64)
        self.logo_frame.pack(fill="x")
        self.logo_frame.pack_propagate(False)
        ctk.CTkLabel(self.logo_frame, text="  F  ", font=("Segoe UI", 28, "bold"),
                     text_color="#0d1117", fg_color=ACCENT).pack(expand=True)

        self.title_lbl = ctk.CTkLabel(self, text="One-Way ANOVA", font=("Segoe UI", 14, "bold"),
                     text_color=TEXT_PRI, fg_color=BG_CARD)
        self.title_lbl.pack(pady=(18, 2))
        self.subtitle_lbl = ctk.CTkLabel(self, text="APA Format Analyzer", font=FONT_TINY,
                     text_color=TEXT_SEC, fg_color=BG_CARD)
        self.subtitle_lbl.pack(pady=(0, 24))

        divider(self)

        # Rater / Researcher name
        ctk.CTkLabel(self, text="RESEARCHER NAME", font=("Segoe UI", 9, "bold"),
                     text_color=TEXT_SEC, fg_color=BG_CARD).pack(anchor="w", padx=18, pady=(18, 4))
        self.researcher_entry = styled_entry(self, placeholder="e.g. Dr. John Smith")
        self.researcher_entry.pack(fill="x", padx=14, pady=(0, 14))

        divider(self)

        # Action buttons
        pad = {"fill": "x", "padx": 14, "pady": 5}

        self.import_btn = sidebar_btn(self, "📁  Import Excel / CSV",
                                      fg=ACCENT2, hover="#3b7ddd")
        self.import_btn.pack(**pad)

        self.add_btn = sidebar_btn(self, "＋  Add Group",
                                   fg="#2d6a4f", hover="#1b4332")
        self.add_btn.pack(**pad)

        self.clear_btn = sidebar_btn(self, "🗑  Clear All",
                                     fg=DANGER, hover="#b91c1c")
        self.clear_btn.pack(**pad)

        divider(self)

        self.run_btn = sidebar_btn(self, "▶   Compute ANOVA",
                                   fg=ACCENT, hover="#009e82",
                                   text_color="#0d1117",
                                   font=("Segoe UI", 13, "bold"), height=44)
        self.run_btn.pack(**pad)

        self.save_btn = sidebar_btn(self, "💾  Export DOCX", fg="#1d4ed8", hover="#1e3a8a")
        self.save_btn.pack(**pad)

        self.preview_btn = sidebar_btn(self, "✏️  Preview & Edit (APA)",
                                       fg=PURPLE, hover="#7e22ce")
        self.preview_btn.configure(state="disabled")
        self.preview_btn.pack(**pad)

        self.plots_btn = sidebar_btn(self, "📊  View Plots",
                                     fg="#0f766e", hover="#134e4a")
        self.plots_btn.configure(state="disabled")
        self.plots_btn.pack(**pad)

        self.reset_btn = sidebar_btn(self, "🔄  Reset", fg="#92400e", hover="#78350f")
        self.reset_btn.pack(**pad)

        # Settings button
        divider(self)
        self.settings_btn = sidebar_btn(self, "⚙   Settings",
                                        fg="#374151", hover="#4b5563",
                                        font=FONT_BODY, height=32)
        self.settings_btn.pack(fill="x", padx=14, pady=8)

        # Footer status
        self.status_label = ctk.CTkLabel(self, text="", font=FONT_TINY,
                                          text_color=ACCENT, fg_color=BG_CARD,
                                          wraplength=190)
        self.status_label.pack(side="bottom", padx=12, pady=10)


# ─── Helper widget factories ─────────────────────────────────────────────────

def divider(parent):
    ctk.CTkFrame(parent, height=1, fg_color=BORDER, corner_radius=0).pack(fill="x", padx=0, pady=4)


def styled_entry(parent, placeholder="", width=0):
    e = ctk.CTkEntry(parent, placeholder_text=placeholder,
                     fg_color=BG_INPUT, border_color=BORDER,
                     text_color=TEXT_PRI, placeholder_text_color=TEXT_SEC,
                     border_width=1, corner_radius=6, height=34,
                     font=FONT_BODY)
    if width:
        e.configure(width=width)
    return e


def sidebar_btn(parent, text, fg, hover, text_color=TEXT_PRI,
                font=FONT_BTN, height=36):
    return ctk.CTkButton(parent, text=text, fg_color=fg, hover_color=hover,
                         text_color=text_color, font=font, height=height,
                         corner_radius=8)


def card(parent, title="", **kw):
    frame = ctk.CTkFrame(parent, fg_color=BG_CARD, corner_radius=12,
                         border_width=1, border_color=BORDER, **kw)
    if title:
        ctk.CTkLabel(frame, text=title, font=FONT_CARD,
                     text_color=TEXT_PRI).pack(anchor="w", padx=16, pady=(12, 4))
    return frame


# ─── Main Application ─────────────────────────────────────────────────────────

class ANOVAAnalyzer(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("One-Way ANOVA Analyzer")
        self.geometry("1200x780")
        self.minsize(1100, 700)
        self.configure(fg_color=BG_DEEP)

        self.group_widgets = []
        self.anova_results = None

        self._build_ui()
        for _ in range(3):
            self.add_group()

    # ── Layout ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        # Sidebar
        self.sidebar = Sidebar(self)
        self.sidebar.pack(side="left", fill="y")

        # Wire sidebar buttons
        self.sidebar.import_btn.configure(command=self.import_excel)
        self.sidebar.add_btn.configure(command=self.add_group)
        self.sidebar.clear_btn.configure(command=self.clear_all)
        self.sidebar.run_btn.configure(command=self.run_anova)
        self.sidebar.save_btn.configure(command=self.save_to_docx)
        self.sidebar.preview_btn.configure(command=self.preview_edit_docx_report)
        self.sidebar.plots_btn.configure(command=self.preview_plots_graphs)
        self.sidebar.reset_btn.configure(command=self.reset_all)
        self.sidebar.settings_btn.configure(command=self.open_settings)

        # Main content
        content = ctk.CTkFrame(self, fg_color=BG_DEEP, corner_radius=0)
        content.pack(side="left", fill="both", expand=True)

        # ── Top header bar ────────────────────────────────────────────────────
        header = ctk.CTkFrame(content, fg_color=BG_CARD, corner_radius=0, height=64)
        header.pack(fill="x")
        header.pack_propagate(False)

        # LEFT side: Title + Subtitle entries
        meta_row = ctk.CTkFrame(header, fg_color=BG_CARD)
        meta_row.pack(side="left", padx=16)

        ctk.CTkLabel(meta_row, text="Title:", font=("Segoe UI", 12, "bold"),
                     text_color=TEXT_PRI).grid(row=0, column=0, padx=(0, 4))
        self.report_title_entry = styled_entry(meta_row,
                                               placeholder="ANOVA ANALYSIS RESULTS", width=200)
        self.report_title_entry.grid(row=0, column=1, padx=4)

        ctk.CTkLabel(meta_row, text="Subtitle:", font=("Segoe UI", 12, "bold"),
                     text_color=TEXT_PRI).grid(row=0, column=2, padx=(12, 4))
        self.report_subtitle_entry = styled_entry(meta_row,
                                                   placeholder="e.g. Variables A vs B", width=180)
        self.report_subtitle_entry.insert(0, "ANOVA")
        self.report_subtitle_entry.grid(row=0, column=3, padx=4)

        # RIGHT side: app name as subtitle label
        ctk.CTkLabel(header, text="One-Way ANOVA Analysis",
                     font=("Segoe UI", 13), text_color=TEXT_SEC).pack(side="right", padx=24)

        # ── Two-column resizable body ─────────────────────────────────────────
        import tkinter as tk
        from tkinter import ttk

        outer = ctk.CTkFrame(content, fg_color=BG_DEEP)
        outer.pack(fill="both", expand=True, padx=16, pady=16)

        style = ttk.Style()
        style.theme_use("default")
        style.configure("Sash", sashthickness=6, sashrelief="flat",
                        background="#30363d")

        pane = ttk.PanedWindow(outer, orient=tk.HORIZONTAL)
        pane.pack(fill="both", expand=True)

        # LEFT
        left_wrap = ctk.CTkFrame(pane, fg_color=BG_DEEP)
        left = card(left_wrap, title="📋  Input Data  (comma-separated values)")
        left.pack(fill="both", expand=True)
        self.groups_frame = ctk.CTkScrollableFrame(left, fg_color=BG_CARD,
                                                    scrollbar_button_color=BORDER)
        self.groups_frame.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        pane.add(left_wrap, weight=1)

        # RIGHT
        right_wrap = ctk.CTkFrame(pane, fg_color=BG_DEEP)
        right = card(right_wrap, title="📊  Analysis Results")
        right.pack(fill="both", expand=True)
        self.results_text = ctk.CTkTextbox(right, fg_color=BG_INPUT, text_color=TEXT_PRI,
                                            font=FONT_MONO, wrap="none",
                                            border_width=1, border_color=BORDER,
                                            corner_radius=8)
        self.results_text.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        pane.add(right_wrap, weight=1)

        def _set_sash(*_):
            total = pane.winfo_width()
            if total > 100:
                pane.sashpos(0, total // 2)
                pane.unbind("<Configure>")
        pane.bind("<Configure>", _set_sash)

        # ── Status bar ────────────────────────────────────────────────────────
        bar = ctk.CTkFrame(content, fg_color=BG_CARD, height=28, corner_radius=0)
        bar.pack(fill="x", side="bottom")
        bar.pack_propagate(False)
        self.file_label = ctk.CTkLabel(bar, text="No file saved yet",
                                        font=FONT_TINY, text_color=TEXT_SEC)
        self.file_label.pack(side="left", padx=12)
        self.kappa_label = ctk.CTkLabel(bar, text="",
                                         font=("Segoe UI", 9, "bold"), text_color=ACCENT)
        self.kappa_label.pack(side="right", padx=12)

    # ── Group widgets ─────────────────────────────────────────────────────────

    def add_group(self):
        group_num = len(self.group_widgets) + 1

        row = ctk.CTkFrame(self.groups_frame, fg_color=BG_PANEL,
                           corner_radius=8, border_width=1, border_color=BORDER)
        row.pack(fill="x", pady=4)

        badge = ctk.CTkEntry(row, width=48, height=32,
                             font=("Segoe UI", 11, "bold"),
                             fg_color=ACCENT, border_color=ACCENT,
                             text_color="#0d1117",
                             justify="center",
                             border_width=0, corner_radius=6)
        badge.insert(0, f"G{group_num}")
        badge.pack(side="left", padx=(10, 8), pady=8)

        entry = ctk.CTkEntry(row, placeholder_text="e.g., 12, 15, 14, 17",
                             fg_color=BG_INPUT, border_color=BORDER,
                             text_color=TEXT_PRI, placeholder_text_color=TEXT_SEC,
                             border_width=1, corner_radius=6, height=32,
                             font=FONT_BODY)
        entry.pack(side="left", fill="x", expand=True, padx=(0, 8), pady=8)

        remove_btn = ctk.CTkButton(row, text="✕", width=28, height=28,
                                   fg_color="#2d2d2d", hover_color=DANGER,
                                   text_color=TEXT_SEC, corner_radius=6,
                                   font=("Segoe UI", 11, "bold"))
        remove_btn.configure(command=lambda: self.remove_group(row, badge, entry, remove_btn))
        remove_btn.pack(side="right", padx=8, pady=8)

        self.group_widgets.append((row, badge, entry, remove_btn))

    def remove_group(self, row, badge, entry, remove_btn):
        if len(self.group_widgets) <= 2:
            messagebox.showwarning("Warning", "At least 2 groups are required for ANOVA!")
            return
        row.destroy()
        self.group_widgets = [w for w in self.group_widgets if w[0] != row]
        for i, (r, b, e, btn) in enumerate(self.group_widgets, 1):
            import re
            if re.fullmatch(r"G\d+", b.get()):
                b.delete(0, "end")
                b.insert(0, f"G{i}")

    def clear_all(self):
        for _, _, entry, _ in self.group_widgets:
            entry.delete(0, "end")

    def reset_all(self):
        self.clear_all()
        self.results_text.delete("1.0", "end")
        self.anova_results = None
        self.sidebar.preview_btn.configure(state="disabled")
        self.sidebar.plots_btn.configure(state="disabled")
        self.kappa_label.configure(text="")

    # ── Import ────────────────────────────────────────────────────────────────

    def open_settings(self):
        SettingsWindow(self, self)

    def apply_settings(self):
        sm = SettingsManager()
        fb, fc, fh, fm, fbt, ft = sm.fonts
        ff = sm.font_family

        # Theme
        ctk.set_appearance_mode(
            {"Dark": "dark", "Light": "light", "System": "system"}
            .get(sm.get("theme"), "dark")
        )

        # Sidebar width & accent
        self.sidebar.configure(width=sm.sidebar_width)
        self.sidebar.logo_frame.configure(fg_color=sm.accent)
        self.sidebar.run_btn.configure(
            fg_color=sm.accent, hover_color=sm.accent_hover,
            font=(ff, fbt, "bold"))

        # Sidebar labels
        self.sidebar.title_lbl.configure(font=(ff, fc, "bold"))
        self.sidebar.subtitle_lbl.configure(font=(ff, ft))
        self.sidebar.status_label.configure(font=(ff, ft))

        # All sidebar buttons font update
        for btn in [self.sidebar.import_btn, self.sidebar.add_btn,
                    self.sidebar.clear_btn, self.sidebar.save_btn,
                    self.sidebar.preview_btn, self.sidebar.plots_btn,
                    self.sidebar.reset_btn, self.sidebar.settings_btn]:
            try: btn.configure(font=(ff, fbt))
            except Exception: pass

        # Results text
        self.results_text.configure(font=(ff, fm), wrap=sm.wrap_mode)

        # Status bar
        self.kappa_label.configure(font=(ff, ft, "bold"))
        self.file_label.configure(font=(ff, ft))

        # Group data entry fonts
        for _, badge, entry, _ in self.group_widgets:
            entry.configure(font=(ff, fb))

    def import_excel(self):
        filepath = filedialog.askopenfilename(
            title="Select Excel / CSV File",
            filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if not filepath:
            return
        try:
            df = pd.read_excel(filepath) if filepath.endswith(".xlsx") else pd.read_csv(filepath)
            if len(df.columns) == 0:
                messagebox.showerror("Error", "File has no data columns!"); return

            self.clear_all()
            current = len(self.group_widgets)
            for _ in range(max(0, len(df.columns) - current)):
                self.add_group()

            ok = 0
            for col_idx, col in enumerate(df.columns):
                if col_idx >= len(self.group_widgets): break
                data = pd.to_numeric(df[col], errors="coerce").dropna()
                if len(data) == 0: continue
                _, _, entry, _ = self.group_widgets[col_idx]
                entry.delete(0, "end")
                entry.insert(0, ", ".join(str(v) for v in data.values))
                ok += 1

            if ok:
                messagebox.showinfo("Imported", f"Imported {ok} group(s) successfully.")
            else:
                messagebox.showwarning("Warning", "No valid numeric data found.")
        except Exception as e:
            messagebox.showerror("Error", f"Import failed:\n{e}")

    # ── Validation ────────────────────────────────────────────────────────────

    def validate_and_parse_inputs(self):
        if len(self.group_widgets) < 2:
            messagebox.showerror("Error", "At least 2 groups required!"); return None
        groups, names = [], []
        for i, (_, badge, entry, _) in enumerate(self.group_widgets, 1):
            text = entry.get().strip()
            if not text:
                messagebox.showerror("Error", f"Group {i} is empty!"); return None
            try:
                vals = [float(x.strip()) for x in text.split(",")]
                if len(vals) < 2:
                    messagebox.showerror("Error", f"Group {i} needs ≥ 2 values!"); return None
                group_name = badge.get().strip() or f"Group {i}"
                groups.append(vals); names.append(group_name)
            except ValueError:
                messagebox.showerror("Error", f"Group {i} has non-numeric values!"); return None
        return groups, names

    # ── ANOVA ─────────────────────────────────────────────────────────────────

    def run_anova(self):
        result = self.validate_and_parse_inputs()
        if result is None: return
        groups, group_names = result

        try:
            F_stat, p_val = f_oneway(*groups)
            all_data = np.concatenate(groups)
            grand_mean = np.mean(all_data)
            k, N = len(groups), len(all_data)

            SS_b = sum(len(g) * (np.mean(g) - grand_mean) ** 2 for g in groups)
            SS_t = np.sum((all_data - grand_mean) ** 2)
            SS_w = SS_t - SS_b
            df_b, df_w = k - 1, N - k
            MS_b, MS_w = SS_b / df_b, SS_w / df_w

            alpha = 0.05
            sig = p_val < alpha
            decision = "Reject H₀" if sig else "Fail to Reject H₀"
            conclusion = ("There is a statistically significant difference among group means."
                          if sig else
                          "There is no statistically significant difference among group means.")

            self.anova_results = {
                "groups": groups, "group_names": group_names,
                "F_statistic": F_stat, "p_value": p_val, "alpha": alpha,
                "decision": decision, "conclusion": conclusion, "is_significant": sig,
                "SS_between": SS_b, "SS_within": SS_w, "SS_total": SS_t,
                "df_between": df_b, "df_within": df_w,
                "MS_between": MS_b, "MS_within": MS_w, "all_data": all_data,
                "report_title": self.report_title_entry.get().strip() or "ANOVA ANALYSIS RESULTS",
                "report_subtitle": self.report_subtitle_entry.get().strip(),
                "researcher_name": self.sidebar.researcher_entry.get().strip(),
                "edited": False
            }

            if sig:
                try:
                    labels = []
                    for i, g in enumerate(groups):
                        labels.extend([group_names[i]] * len(g))
                    self.anova_results["tukey"] = pairwise_tukeyhsd(all_data, labels)
                except Exception:
                    pass

            self._display_results()
            self.sidebar.preview_btn.configure(state="normal")
            self.sidebar.plots_btn.configure(state="normal")
            sm2 = SettingsManager()
            self.kappa_label.configure(
                text=f"F({df_b}, {df_w}) = {sm2.fmt(F_stat)}   p = {sm2.fmt(p_val)}   {'✓ Significant' if sig else '✗ Not Significant'}"
            )

        except Exception as e:
            messagebox.showerror("Error", f"Analysis failed:\n{e}")

    # ── Display ───────────────────────────────────────────────────────────────

    def _display_results(self):
        self.results_text.delete("1.0", "end")
        r = self.anova_results
        if r is None: return
        # Use live decimal/wrap settings
        sm = SettingsManager()
        self.results_text.configure(wrap=sm.wrap_mode)
        def _f(v): return sm.fmt(v)

        line = "─" * 52
        out = f"{line}\n{r['report_title']}\n"
        if r["report_subtitle"]: out += f"{r['report_subtitle']}\n"
        if r["researcher_name"]: out += f"by: {r['researcher_name']}\n"
        out += f"{line}\n\n"

        out += "DESCRIPTIVE STATISTICS\n" + "─" * 30 + "\n"
        for name, g in zip(r["group_names"], r["groups"]):
            out += (f"{name}:  n={len(g)}  M={_f(np.mean(g))}  "
                    f"SD={_f(np.std(g, ddof=1))}\n")

        out += f"\n{'Source':<16}{'SS':>10}{'df':>6}{'MS':>12}{'F':>10}\n"
        out += "─" * 55 + "\n"
        out += (f"{'Between':<16}{_f(r['SS_between']):>10}{r['df_between']:>6}"
                f"{_f(r['MS_between']):>12}{_f(r['F_statistic']):>10}\n")
        out += (f"{'Within':<16}{_f(r['SS_within']):>10}{r['df_within']:>6}"
                f"{_f(r['MS_within']):>12}\n")
        out += (f"{'Total':<16}{_f(r['SS_total']):>10}"
                f"{r['df_between']+r['df_within']:>6}\n")

        out += f"\n{'═'*52}\nTEST RESULTS\n{'═'*52}\n"
        out += (f"F({r['df_between']}, {r['df_within']}) = {_f(r['F_statistic'])},  "
                f"p = {_f(r['p_value'])},  α = {r['alpha']}\n\n")
        out += f"Decision:   {r['decision']}\n\n"
        out += f"Conclusion:\n{r['conclusion']}\n"

        if r["is_significant"] and "tukey" in r:
            out += f"\n{'─'*52}\nPOST HOC — Tukey HSD\n{'─'*52}\n"
            out += str(r["tukey"]) + "\n"

        self.results_text.insert("1.0", out)

    # ── Preview & Edit ────────────────────────────────────────────────────────

    def preview_edit_docx_report(self):
        if not self.anova_results:
            messagebox.showwarning("Warning", "Run ANOVA first!"); return

        win = ctk.CTkToplevel(self)
        win.title("APA-Style Preview & Edit")
        win.geometry("1000x780")
        win.configure(fg_color=BG_DEEP)

        # Title bar
        top = ctk.CTkFrame(win, fg_color=BG_CARD, height=56, corner_radius=0)
        top.pack(fill="x")
        top.pack_propagate(False)
        ctk.CTkLabel(top, text="✏️  APA Report Preview & Edit",
                     font=FONT_HEAD, text_color=TEXT_PRI).pack(side="left", padx=20)

        scroll = ctk.CTkScrollableFrame(win, fg_color=BG_DEEP)
        scroll.pack(fill="both", expand=True, padx=16, pady=16)

        r = self.anova_results

        # ── Header section ────
        hdr = card(scroll, "Report Header")
        hdr.pack(fill="x", pady=(0, 12))

        def lbl_entry(parent, label, value, width=500):
            row = ctk.CTkFrame(parent, fg_color="transparent")
            row.pack(fill="x", padx=16, pady=4)
            ctk.CTkLabel(row, text=label, font=FONT_TINY, text_color=TEXT_SEC,
                         width=120).pack(side="left")
            e = styled_entry(row, width=width)
            e.insert(0, value)
            e.pack(side="left")
            return e

        title_e    = lbl_entry(hdr, "Report Title",    r["report_title"])
        subtitle_e = lbl_entry(hdr, "Subtitle",        r.get("report_subtitle", ""))
        author_e   = lbl_entry(hdr, "Researcher Name", r.get("researcher_name", ""))

        # ── Table 1: Descriptive ──
        t1 = card(scroll, "Table 1 — Descriptive Statistics")
        t1.pack(fill="x", pady=(0, 12))
        t1title_e = styled_entry(t1, width=500)
        t1title_e.insert(0, r.get("desc_table_title", "Descriptive Statistics for Groups"))
        t1title_e.pack(padx=16, pady=(0, 8))

        tbl1 = ctk.CTkFrame(t1, fg_color="transparent")
        tbl1.pack(padx=16, pady=(0, 12))
        for col, h in enumerate(["Group", "n", "M", "SD"]):
            ctk.CTkLabel(tbl1, text=h, font=FONT_CARD, text_color=ACCENT,
                         width=130).grid(row=0, column=col, padx=4, pady=4)

        desc_entries = []
        for idx, (name, g) in enumerate(zip(r["group_names"], r["groups"]), 1):
            ge = styled_entry(tbl1, width=130); ge.insert(0, name)
            ge.grid(row=idx, column=0, padx=4, pady=2)
            ctk.CTkLabel(tbl1, text=str(len(g)), width=130, text_color=TEXT_PRI).grid(row=idx, column=1, padx=4)
            ctk.CTkLabel(tbl1, text=fmt(np.mean(g)), width=130, text_color=TEXT_PRI).grid(row=idx, column=2, padx=4)
            ctk.CTkLabel(tbl1, text=fmt(np.std(g, ddof=1)), width=130, text_color=TEXT_PRI).grid(row=idx, column=3, padx=4)
            desc_entries.append(ge)

        # ── Table 2: ANOVA ──
        t2 = card(scroll, "Table 2 — ANOVA Summary")
        t2.pack(fill="x", pady=(0, 12))
        t2title_e = styled_entry(t2, width=500)
        t2title_e.insert(0, r.get("anova_table_title", "Analysis of Variance Summary Table"))
        t2title_e.pack(padx=16, pady=(0, 8))

        tbl2 = ctk.CTkFrame(t2, fg_color="transparent")
        tbl2.pack(padx=16, pady=(0, 12))
        for col, h in enumerate(["Source", "SS", "df", "MS", "F", "p"]):
            ctk.CTkLabel(tbl2, text=h, font=FONT_CARD, text_color=ACCENT,
                         width=110).grid(row=0, column=col, padx=4)

        between_e = styled_entry(tbl2, width=110); between_e.insert(0, r.get("anova_between_label", "Between Groups"))
        between_e.grid(row=1, column=0, padx=4, pady=2)
        for col, val in enumerate([fmt(r["SS_between"]), str(r["df_between"]),
                                    fmt(r["MS_between"]), fmt(r["F_statistic"]), fmt(r["p_value"])], 1):
            ctk.CTkLabel(tbl2, text=val, width=110, text_color=TEXT_PRI).grid(row=1, column=col, padx=4)

        within_e = styled_entry(tbl2, width=110); within_e.insert(0, r.get("anova_within_label", "Within Groups"))
        within_e.grid(row=2, column=0, padx=4, pady=2)
        for col, val in enumerate([fmt(r["SS_within"]), str(r["df_within"]), fmt(r["MS_within"])], 1):
            ctk.CTkLabel(tbl2, text=val, width=110, text_color=TEXT_PRI).grid(row=2, column=col, padx=4)

        total_e = styled_entry(tbl2, width=110); total_e.insert(0, r.get("anova_total_label", "Total"))
        total_e.grid(row=3, column=0, padx=4, pady=2)
        for col, val in enumerate([fmt(r["SS_total"]), str(r["df_between"]+r["df_within"])], 1):
            ctk.CTkLabel(tbl2, text=val, width=110, text_color=TEXT_PRI).grid(row=3, column=col, padx=4)

        # ── Post-hoc ──
        posthoc_text_widget = None
        posthoc_title_e = None
        if r["is_significant"] and "tukey" in r:
            t3 = card(scroll, "Table 3 — Post Hoc (Tukey HSD)")
            t3.pack(fill="x", pady=(0, 12))
            posthoc_title_e = styled_entry(t3, width=500)
            posthoc_title_e.insert(0, r.get("posthoc_title", "Post Hoc Comparisons (Tukey HSD)"))
            posthoc_title_e.pack(padx=16, pady=(0, 8))
            posthoc_text_widget = ctk.CTkTextbox(t3, height=150, font=FONT_MONO,
                                                  fg_color=BG_INPUT, text_color=TEXT_PRI)
            posthoc_text_widget.insert("1.0", str(r["tukey"]))
            posthoc_text_widget.pack(fill="x", padx=16, pady=(0, 12))

        # ── Conclusion ──
        tc = card(scroll, "Conclusion & Decision")
        tc.pack(fill="x", pady=(0, 12))
        conc_text = ctk.CTkTextbox(tc, height=180, font=FONT_BODY,
                                    fg_color=BG_INPUT, text_color=TEXT_PRI)
        default_conc = (f"Decision: {r['decision']}\n\n"
                        f"F({r['df_between']}, {r['df_within']}) = {fmt(r['F_statistic'])}, "
                        f"p = {fmt(r['p_value'])}\n\n"
                        f"Conclusion:\n{r['conclusion']}")
        conc_text.insert("1.0", r.get("conclusion_text", default_conc))
        conc_text.pack(fill="x", padx=16, pady=(0, 12))

        # ── Raw data ──
        tr = card(scroll, "Table 4 — Raw Data")
        tr.pack(fill="x", pady=(0, 12))
        ctk.CTkLabel(tr, text="⚠  Editing here won't recompute statistics.",
                     font=FONT_TINY, text_color=WARN).pack(anchor="w", padx=16)
        raw_title_e = styled_entry(tr, width=500)
        raw_title_e.insert(0, r.get("rawdata_table_title", "Raw Data by Group"))
        raw_title_e.pack(padx=16, pady=(0, 8))

        rtbl = ctk.CTkFrame(tr, fg_color="transparent"); rtbl.pack(padx=16, pady=(0, 12))
        for col, h in enumerate(["Group", "n", "Values"]):
            ctk.CTkLabel(rtbl, text=h, font=FONT_CARD, text_color=ACCENT,
                         width=140).grid(row=0, column=col, padx=4)
        raw_entries = []
        for idx, (name, g) in enumerate(zip(r["group_names"], r["groups"]), 1):
            rge = styled_entry(rtbl, width=140); rge.insert(0, name)
            rge.grid(row=idx, column=0, padx=4, pady=2)
            ctk.CTkLabel(rtbl, text=str(len(g)), width=140, text_color=TEXT_PRI).grid(row=idx, column=1, padx=4)
            rve = styled_entry(rtbl, width=380)
            rve.insert(0, ", ".join(fmt(v) for v in g))
            rve.grid(row=idx, column=2, padx=4, pady=2)
            raw_entries.append((rge, rve))

        ctk.CTkLabel(scroll, text="ℹ  Statistical values are display-only and not recomputed here.",
                     font=FONT_TINY, text_color=TEXT_SEC).pack(pady=8)

        def save_edits():
            self.anova_results.update({
                "report_title": title_e.get(),
                "report_subtitle": subtitle_e.get(),
                "researcher_name": author_e.get(),
                "group_names": [e.get() for e in desc_entries],
                "desc_table_title": t1title_e.get(),
                "anova_table_title": t2title_e.get(),
                "anova_between_label": between_e.get(),
                "anova_within_label": within_e.get(),
                "anova_total_label": total_e.get(),
                "conclusion_text": conc_text.get("1.0", "end-1c"),
                "rawdata_table_title": raw_title_e.get(),
                "raw_data_edits": [{"group_name": rge.get(), "values_text": rve.get()}
                                    for rge, rve in raw_entries],
                "edited": True
            })
            if posthoc_title_e:
                self.anova_results["posthoc_title"] = posthoc_title_e.get()
            if posthoc_text_widget:
                self.anova_results["posthoc_text"] = posthoc_text_widget.get("1.0", "end-1c")
            messagebox.showinfo("Saved", "Edits saved — will be included in DOCX export.")
            win.destroy()

        ctk.CTkButton(scroll, text="💾  Save Edits", command=save_edits,
                      height=42, font=FONT_BTN, fg_color=SUCCESS, hover_color="#16a34a",
                      text_color="#0d1117", corner_radius=8).pack(pady=16)

    # ── Plots ─────────────────────────────────────────────────────────────────

    def preview_plots_graphs(self):
        if not self.anova_results:
            messagebox.showwarning("Warning", "Run ANOVA first!"); return

        win = ctk.CTkToplevel(self)
        win.title("Statistical Plots")
        win.geometry("1260x680")
        win.configure(fg_color=BG_DEEP)

        top = ctk.CTkFrame(win, fg_color=BG_CARD, height=52, corner_radius=0)
        top.pack(fill="x"); top.pack_propagate(False)
        ctk.CTkLabel(top, text="📊  Statistical Plots Preview",
                     font=FONT_HEAD, text_color=TEXT_PRI).pack(side="left", padx=20)

        scroll = ctk.CTkScrollableFrame(win, fg_color=BG_DEEP)
        scroll.pack(fill="both", expand=True, padx=16, pady=16)
        scroll.grid_columnconfigure((0, 1, 2), weight=1)

        r = self.anova_results
        groups = r["groups"]; group_names = r["group_names"]
        COLORS = ["#00c9a7", "#4e9eff", "#f59e0b", "#a855f7", "#ef4444"]

        def make_col(parent, col, title, subtitle):
            f = card(parent, "")
            f.grid(row=0, column=col, padx=6, pady=4, sticky="nsew")
            ctk.CTkLabel(f, text=title, font=FONT_CARD, text_color=TEXT_PRI).pack(pady=(12, 0))
            ctk.CTkLabel(f, text=subtitle, font=FONT_TINY, text_color=TEXT_SEC).pack(pady=(0, 6))
            return f

        # ── Boxplot ──
        c1 = make_col(scroll, 0, "Distribution", "Median · IQR · Outliers")
        fig1 = Figure(figsize=(3.8, 4.6), facecolor=BG_CARD)
        ax1 = fig1.add_subplot(111, facecolor=BG_PANEL)
        bp = ax1.boxplot(groups, labels=group_names, patch_artist=True,
                         medianprops=dict(color="#f59e0b", linewidth=2.5),
                         whiskerprops=dict(color=TEXT_SEC),
                         capprops=dict(color=TEXT_SEC),
                         flierprops=dict(marker="o", markerfacecolor=DANGER, markersize=5, alpha=0.6))
        for patch, c in zip(bp["boxes"], COLORS):
            patch.set_facecolor(c); patch.set_alpha(0.55)
        ax1.set_xlabel("Groups", color=TEXT_SEC, fontsize=9)
        ax1.set_ylabel("Values", color=TEXT_SEC, fontsize=9)
        ax1.tick_params(colors=TEXT_SEC, labelsize=8)
        ax1.spines[["top","right"]].set_visible(False)
        ax1.spines[["left","bottom"]].set_color(BORDER)
        ax1.grid(axis="y", alpha=0.2, linestyle="--")
        fig1.tight_layout(pad=1.2)
        FigureCanvasTkAgg(fig1, c1).get_tk_widget().pack(fill="both", expand=True, padx=10, pady=10)

        # ── Bar + error ──
        c2 = make_col(scroll, 1, "Group Means", "Mean ± SD error bars")
        means = [np.mean(g) for g in groups]
        stds  = [np.std(g, ddof=1) for g in groups]
        fig2 = Figure(figsize=(3.8, 4.6), facecolor=BG_CARD)
        ax2 = fig2.add_subplot(111, facecolor=BG_PANEL)
        xp = np.arange(len(group_names))
        bars = ax2.bar(xp, means, yerr=stds, capsize=6, width=0.55,
                       color=COLORS[:len(groups)], alpha=0.80,
                       ecolor=TEXT_SEC, linewidth=0)
        for bar, m, s in zip(bars, means, stds):
            ax2.text(bar.get_x() + bar.get_width()/2, m + s + 0.3,
                     fmt(m), ha="center", va="bottom", color=TEXT_PRI, fontsize=8)
        ax2.set_xticks(xp); ax2.set_xticklabels(group_names)
        ax2.tick_params(colors=TEXT_SEC, labelsize=8)
        ax2.spines[["top","right"]].set_visible(False)
        ax2.spines[["left","bottom"]].set_color(BORDER)
        ax2.grid(axis="y", alpha=0.2, linestyle="--")
        fig2.tight_layout(pad=1.2)
        FigureCanvasTkAgg(fig2, c2).get_tk_widget().pack(fill="both", expand=True, padx=10, pady=10)

        # ── Strip + mean line ──
        c3 = make_col(scroll, 2, "Data Points", "Individual values + mean marker")
        fig3 = Figure(figsize=(3.8, 4.6), facecolor=BG_CARD)
        ax3 = fig3.add_subplot(111, facecolor=BG_PANEL)
        for i, (g, name) in enumerate(zip(groups, group_names)):
            jx = np.random.normal(i + 1, 0.06, size=len(g))
            ax3.scatter(jx, g, alpha=0.65, s=28, color=COLORS[i % len(COLORS)],
                        edgecolors="white", linewidth=0.4)
            mn = np.mean(g)
            ax3.plot([i+0.65, i+1.35], [mn, mn], color="#f59e0b", linewidth=3)
        ax3.set_xticks(range(1, len(group_names)+1)); ax3.set_xticklabels(group_names)
        ax3.tick_params(colors=TEXT_SEC, labelsize=8)
        ax3.spines[["top","right"]].set_visible(False)
        ax3.spines[["left","bottom"]].set_color(BORDER)
        ax3.grid(axis="y", alpha=0.2, linestyle="--")
        fig3.tight_layout(pad=1.2)
        FigureCanvasTkAgg(fig3, c3).get_tk_widget().pack(fill="both", expand=True, padx=10, pady=10)

        footer = ctk.CTkFrame(scroll, fg_color="transparent")
        footer.grid(row=1, column=0, columnspan=3, pady=10, sticky="ew")
        ctk.CTkLabel(footer, text="ℹ  Visual preview only — statistical values are not recomputed here.",
                     font=FONT_TINY, text_color=TEXT_SEC).pack(side="left", padx=8)
        ctk.CTkButton(footer, text="Close", command=win.destroy,
                      width=100, height=32, corner_radius=6,
                      fg_color=BG_PANEL, hover_color=BORDER).pack(side="right", padx=8)

    # ── Save DOCX ─────────────────────────────────────────────────────────────

    def save_to_docx(self):
        if not self.anova_results:
            messagebox.showwarning("Warning", "Run ANOVA first!"); return

        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            title="Save ANOVA Report"
        )
        if not filepath: return

        try:
            r = self.anova_results
            doc = Document()

            for section in doc.sections:
                section.top_margin    = Inches(0.60)
                section.bottom_margin = Inches(0.60)
                section.left_margin   = Inches(0.65)
                section.right_margin  = Inches(0.65)

            # ── APA border helper ──
            def apa_borders(table):
                tbl = table._tbl
                tblPr = tbl.tblPr
                tb = OxmlElement("w:tblBorders")
                for bn in ["top","left","bottom","right","insideH","insideV"]:
                    b = OxmlElement(f"w:{bn}")
                    b.set(qn("w:val"), "none")
                    tb.append(b)
                for bn, sz in [("top","12"), ("bottom","12")]:
                    b = OxmlElement(f"w:{bn}")
                    b.set(qn("w:val"), "single")
                    b.set(qn("w:sz"), sz)
                    tb.append(b)
                tblPr.append(tb)

            def add_header_sep(table):
                """Add thin bottom border to header row cells"""
                for cell in table.rows[0].cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement("w:tcBorders")
                    bot = OxmlElement("w:bottom")
                    bot.set(qn("w:val"), "single")
                    bot.set(qn("w:sz"), "6")
                    tcBorders.append(bot)
                    tcPr.append(tcBorders)

            def cell_fmt(cell, text, bold=False, italic=False,
                          size=9, align="left", color=None):
                para = cell.paragraphs[0]
                para.alignment = (WD_PARAGRAPH_ALIGNMENT.CENTER
                                   if align == "center" else WD_PARAGRAPH_ALIGNMENT.LEFT)
                run = para.add_run(text)
                run.font.size = Pt(size)
                run.bold = bold; run.italic = italic
                if color:
                    run.font.color.rgb = RGBColor(*color)

            # ── TITLE ──
            title_p = doc.add_heading(r["report_title"], 1)
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_p.runs:
                run.font.size = Pt(14); run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)

            if r.get("report_subtitle"):
                sp = doc.add_paragraph(r["report_subtitle"])
                sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in sp.runs: run.italic = True; run.font.size = Pt(11)

            if r.get("researcher_name"):
                np_ = doc.add_paragraph(r["researcher_name"])
                np_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in np_.runs: run.italic = True; run.font.size = Pt(10)

            # ── 2-column section ──
            s = doc.sections[0]
            sp = s._sectPr
            cols_el = sp.xpath("./w:cols")
            cols_el = cols_el[0] if cols_el else OxmlElement("w:cols")
            cols_el.set(qn("w:num"), "2")
            cols_el.set(qn("w:space"), "720")
            if not sp.xpath("./w:cols"):
                sp.append(cols_el)

            doc.add_paragraph()

            # ── Table 1: Descriptive ──
            p = doc.add_paragraph()
            p.add_run("Table 1\n").bold = True
            p.runs[0].font.size = Pt(10)
            tp = doc.add_paragraph(r.get("desc_table_title", "Descriptive Statistics for Groups"))
            tp.runs[0].italic = True; tp.runs[0].font.size = Pt(9)

            dt = doc.add_table(rows=len(r["groups"]) + 1, cols=4)
            apa_borders(dt)
            for cw, cell in zip([1.0, 0.5, 0.7, 0.7], dt.rows[0].cells):
                cell.width = Inches(cw)
            for i, h in enumerate(["Group","n","M","SD"]):
                cell_fmt(dt.rows[0].cells[i], h, bold=True, size=9, align="center")
            add_header_sep(dt)

            for i, (name, g) in enumerate(zip(r["group_names"], r["groups"]), 1):
                row = dt.rows[i]
                cell_fmt(row.cells[0], name, size=9)
                cell_fmt(row.cells[1], str(len(g)), size=9, align="center")
                cell_fmt(row.cells[2], fmt(np.mean(g)), size=9, align="center")
                cell_fmt(row.cells[3], fmt(np.std(g, ddof=1)), size=9, align="center")

            doc.add_paragraph()

            # ── Table 2: ANOVA ──
            tp2 = doc.add_paragraph(r.get("anova_table_title", "Analysis of Variance Summary Table"))
            tp2.runs[0].italic = True; tp2.runs[0].font.size = Pt(9)

            at = doc.add_table(rows=4, cols=6)
            apa_borders(at)
            for i, h in enumerate(["Source","SS","df","MS","F","p"]):
                cell_fmt(at.rows[0].cells[i], h, bold=True, size=9, align="center")
            add_header_sep(at)

            row1 = at.rows[1]
            cell_fmt(row1.cells[0], r.get("anova_between_label", "Between Groups"), size=9)
            for col, val in enumerate([fmt(r["SS_between"]), str(r["df_between"]),
                                        fmt(r["MS_between"]), fmt(r["F_statistic"]), fmt(r["p_value"])], 1):
                cell_fmt(row1.cells[col], val, size=9, align="center")

            row2 = at.rows[2]
            cell_fmt(row2.cells[0], r.get("anova_within_label", "Within Groups"), size=9)
            for col, val in enumerate([fmt(r["SS_within"]), str(r["df_within"]), fmt(r["MS_within"])], 1):
                cell_fmt(row2.cells[col], val, size=9, align="center")

            row3 = at.rows[3]
            cell_fmt(row3.cells[0], r.get("anova_total_label", "Total"), size=9)
            for col, val in enumerate([fmt(r["SS_total"]),
                                        str(r["df_between"] + r["df_within"])], 1):
                cell_fmt(row3.cells[col], val, size=9, align="center")

            doc.add_paragraph()

            # ── Results paragraph ──
            rh = doc.add_paragraph()
            rh.add_run("Test Results\n").bold = True
            rh.runs[0].font.size = Pt(10)

            if r.get("edited") and "conclusion_text" in r:
                cp = doc.add_paragraph(r["conclusion_text"])
                for run in cp.runs: run.font.size = Pt(9)
            else:
                cp = doc.add_paragraph()
                cp.add_run(f"F({r['df_between']}, {r['df_within']}) = ")
                fr = cp.add_run(fmt(r["F_statistic"])); fr.bold = True
                cp.add_run(", p = ")
                pr = cp.add_run(fmt(r["p_value"])); pr.bold = True
                cp.add_run("\n\nDecision: ")
                dr = cp.add_run(r["decision"]); dr.bold = True
                cp.add_run(f"\n\n{r['conclusion']}")
                for run in cp.runs: run.font.size = Pt(9)

            # ── Post-hoc ──
            if r["is_significant"] and "tukey" in r:
                doc.add_paragraph()
                php = doc.add_paragraph(r.get("posthoc_title", "Post Hoc Comparisons (Tukey HSD)"))
                php.runs[0].italic = True; php.runs[0].font.size = Pt(9)
                phtext = r.get("posthoc_text", str(r["tukey"]))
                pph = doc.add_paragraph(phtext)
                for run in pph.runs:
                    run.font.name = "Courier New"; run.font.size = Pt(7)

            # ── Raw data table ──
            doc.add_paragraph()
            rdt_title = doc.add_paragraph(r.get("rawdata_table_title", "Raw Data by Group"))
            rdt_title.runs[0].italic = True; rdt_title.runs[0].font.size = Pt(9)

            rdt = doc.add_table(rows=len(r["groups"]) + 1, cols=3)
            apa_borders(rdt)
            for i, h in enumerate(["Group","n","Values"]):
                cell_fmt(rdt.rows[0].cells[i], h, bold=True, size=8, align="center")
            add_header_sep(rdt)

            if "raw_data_edits" in r:
                for i, edit in enumerate(r["raw_data_edits"], 1):
                    row = rdt.rows[i]
                    cell_fmt(row.cells[0], edit["group_name"], size=7)
                    n_c = len(edit["values_text"].split(","))
                    cell_fmt(row.cells[1], str(n_c), size=7, align="center")
                    vt = edit["values_text"]
                    if len(vt) > 60: vt = vt[:57] + "…"
                    cell_fmt(row.cells[2], vt, size=7)
            else:
                for i, (name, g) in enumerate(zip(r["group_names"], r["groups"]), 1):
                    row = rdt.rows[i]
                    cell_fmt(row.cells[0], name, size=7)
                    cell_fmt(row.cells[1], str(len(g)), size=7, align="center")
                    vs = ", ".join(fmt(v) for v in g)
                    if len(vs) > 60: vs = vs[:57] + "…"
                    cell_fmt(row.cells[2], vs, size=7)

            doc.add_paragraph()

            # ── Footer ──
            fp = doc.add_paragraph()
            fr1 = fp.add_run(f"Saved: {os.path.abspath(filepath)}\n")
            fr1.font.size = Pt(7); fr1.font.color.rgb = RGBColor(128, 128, 128); fr1.italic = True
            fr2 = fp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}")
            fr2.font.size = Pt(7); fr2.font.color.rgb = RGBColor(128, 128, 128); fr2.italic = True

            doc.save(filepath)
            self.file_label.configure(text=f"Last saved: {filepath}")
            self.sidebar.status_label.configure(text=f"✓ Saved\n{os.path.basename(filepath)}")
            messagebox.showinfo("Saved", f"APA report saved successfully!\n\n{filepath}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document:\n{e}")


# ─── Entry Point ──────────────────────────────────────────────────────────────

def main():
    app = ANOVAAnalyzer()
    app.mainloop()


if __name__ == "__main__":
    main()