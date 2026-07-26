"""
cronbach_session_manager.py  —  Multi-Part Session Manager for Cronbach's Alpha
Adapted from the ANOVA session_manager.py.

Handles saving / loading / exporting of multiple reliability parts
(Part 1-A, 1-B, 2-A …) with DOCX and PDF batch export.
"""

import re
import numpy as np
import customtkinter as ctk
from tkinter import messagebox
from datetime import datetime
import os

# ── Palette (matches main app) ────────────────────────────────────────────────
from ui_theme import (
    BG_DEEP, BG_CARD, BG_PANEL, BG_INPUT,
    ACCENT, ACCENT2, DANGER, WARN, SUCCESS, PURPLE,
    TEXT_PRI, TEXT_SEC, BORDER,
)

FONT_HEAD = ("Segoe UI", 14, "bold")
FONT_CARD = ("Segoe UI", 12, "bold")
FONT_BODY = ("Segoe UI", 11)
FONT_BTN  = ("Segoe UI", 11, "bold")
FONT_MONO = ("Consolas", 9)
FONT_TINY = ("Segoe UI", 9)


# ── Label generator ───────────────────────────────────────────────────────────

def _next_part_label(existing_labels: list, current_label: str = None) -> str:
    """
    Generate next Part label by incrementing only the letter suffix.
    - If current_label is given (e.g. the value in the title entry field),
      parse its number and use that as the base, then find the next unused letter.
    - Falls back to scanning existing_labels for the highest used number.
    - Never resets the Part number back to 1 unless there are truly no labels.
    """
    base_num = None

    # 1. Try to parse the current title-entry label for the number
    if current_label:
        m = re.match(r"Part\s+(\d+)-([A-Z])", current_label.strip())
        if m:
            base_num = int(m.group(1))

    # 2. Fall back to the highest number seen in existing labels
    if base_num is None and existing_labels:
        for lbl in existing_labels:
            m = re.match(r"Part\s+(\d+)-([A-Z])", lbl)
            if m:
                base_num = max(base_num or 0, int(m.group(1)))

    # 3. Truly nothing to go on — start at 1
    if base_num is None:
        return "Part 1-A"

    # Collect all letters already used for this Part number
    used_letters = set()
    for lbl in existing_labels:
        m = re.match(r"Part\s+(\d+)-([A-Z])", lbl)
        if m and int(m.group(1)) == base_num:
            used_letters.add(m.group(2))

    # Find the next unused letter starting from A
    for ch in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
        if ch not in used_letters:
            return f"Part {base_num}-{ch}"

    # All 26 letters exhausted — bump Part number
    return f"Part {base_num + 1}-A"


def _fmt(v, d=4):
    return f"{float(v):.{d}f}" if v is not None else "N/A"


def _interp_color(interp: str) -> str:
    return {
        "Excellent":    ACCENT,
        "Good":         SUCCESS,
        "Acceptable":   ACCENT2,
        "Questionable": WARN,
        "Poor":         "#f97316",
        "Unacceptable": DANGER,
    }.get(interp, TEXT_SEC)


# ── Part Card ─────────────────────────────────────────────────────────────────

class PartCard(ctk.CTkFrame):
    """One row in the session panel."""

    def __init__(self, master, part_data: dict, index: int,
                 on_delete, on_view, **kw):
        super().__init__(master, fg_color=BG_PANEL, corner_radius=8,
                         border_width=1, border_color=BORDER, **kw)
        self.part_data = part_data
        self._build(index, on_delete, on_view)

    def _build(self, index, on_delete, on_view):
        top = ctk.CTkFrame(self, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=(8, 2))

        # Coloured part-label badge
        badge_color = [ACCENT, ACCENT2, WARN, PURPLE, SUCCESS][index % 5]
        ctk.CTkLabel(top, text=self.part_data["label"],
                     font=("Segoe UI", 10, "bold"),
                     fg_color=badge_color, text_color="#0d1117",
                     corner_radius=5, padx=8, pady=2).pack(side="left")

        # Interpretation badge
        interp = self.part_data.get("interpretation", "—")
        ic = _interp_color(interp)
        ctk.CTkLabel(top, text=interp,
                     font=FONT_TINY,
                     fg_color=ic,
                     text_color="#fff" if interp != "Excellent" else "#0d1117",
                     corner_radius=4, padx=6, pady=2).pack(side="left", padx=6)

        # Delete button
        ctk.CTkButton(top, text="✕", width=24, height=24,
                      fg_color="transparent", hover_color=DANGER,
                      text_color=TEXT_SEC, corner_radius=4,
                      font=("Segoe UI", 10, "bold"),
                      command=lambda: on_delete(self.part_data["label"])
                      ).pack(side="right")

        # View button
        ctk.CTkButton(top, text="👁", width=28, height=24,
                      fg_color=BG_INPUT, hover_color=ACCENT2,
                      text_color=TEXT_PRI, corner_radius=4,
                      font=("Segoe UI", 10),
                      command=lambda: on_view(self.part_data)
                      ).pack(side="right", padx=4)

        # Key stats row
        r = self.part_data
        stats = (f"α = {_fmt(r.get('alpha'))}   "
                 f"k = {r.get('n_items', '?')} items   "
                 f"N = {r.get('n_respondents', '?')} respondents   "
                 f"r̄ = {_fmt(r.get('avg_interitem_corr'))}")
        ctk.CTkLabel(self, text=stats, font=FONT_TINY,
                     text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 4))

        # Subtitle / title
        sub = r.get("report_subtitle") or r.get("report_title", "")
        if sub:
            ctk.CTkLabel(self, text=sub, font=("Segoe UI", 9, "italic"),
                         text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 6))


# ── Session Manager Panel ─────────────────────────────────────────────────────

class SessionManagerPanel(ctk.CTkToplevel):
    """
    Floating panel listing all saved Cronbach's Alpha parts.
    """

    def __init__(self, master, saved_parts: list,
                 on_delete_part, on_export_all_docx, on_export_all_pdf):
        super().__init__(master)
        self.title("📁  Session Manager — Saved Parts")
        self.geometry("480x660")
        self.minsize(400, 480)
        self.configure(fg_color=BG_DEEP)
        self.resizable(True, True)

        self.saved_parts    = saved_parts
        self.on_delete_part = on_delete_part
        self.on_export_docx = on_export_all_docx
        self.on_export_pdf  = on_export_all_pdf

        self.transient(master)
        self.lift()
        self.focus_force()

        self._build()

    def _build(self):
        # Header
        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=52)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        ctk.CTkLabel(hdr, text="📁  Session Manager",
                     font=FONT_HEAD, text_color=TEXT_PRI).pack(side="left", padx=16)

        # FIX: create count_label before packing it
        self.count_label = ctk.CTkLabel(hdr,
                     text=f"{len(self.saved_parts)} part(s) saved",
                     font=FONT_TINY, text_color=TEXT_SEC)
        self.count_label.pack(side="right", padx=16)

        # Scrollable list
        self.scroll = ctk.CTkScrollableFrame(self, fg_color=BG_DEEP,
                                              scrollbar_button_color=BORDER)
        self.scroll.pack(fill="both", expand=True, padx=12, pady=12)
        self._populate_cards()

        # Export / close buttons
        btn_frame = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=68)
        btn_frame.pack(fill="x")
        btn_frame.pack_propagate(False)

        ctk.CTkButton(btn_frame, text="📄  Export All → DOCX",
                      fg_color="#1d4ed8", hover_color="#1e3a8a",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_docx).pack(side="left", padx=(14, 6), pady=16)

        ctk.CTkButton(btn_frame, text="🖨  Export All → PDF",
                      fg_color=DANGER, hover_color="#b91c1c",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_pdf).pack(side="left", padx=6, pady=16)

        ctk.CTkButton(btn_frame, text="✕ Close",
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.destroy).pack(side="right", padx=14, pady=16)

    def _populate_cards(self):
        for w in self.scroll.winfo_children():
            w.destroy()
        if not self.saved_parts:
            ctk.CTkLabel(self.scroll,
                         text="No parts saved yet.\nCompute α and click  '💾 Save Part'.",
                         font=FONT_BODY, text_color=TEXT_SEC,
                         justify="center").pack(expand=True, pady=40)
            return
        for i, part in enumerate(self.saved_parts):
            c = PartCard(self.scroll, part, i,
                         on_delete=self._handle_delete,
                         on_view=self._handle_view)
            c.pack(fill="x", pady=4)

    def _handle_delete(self, label: str):
        if messagebox.askyesno("Delete Part",
                                f"Remove '{label}' from this session?"):
            self.on_delete_part(label)
            self.refresh()

    def _handle_view(self, part_data: dict):
        ViewPartWindow(self, part_data)

    def refresh(self):
        self._populate_cards()
        self.count_label.configure(text=f"{len(self.saved_parts)} part(s) saved")


# ── View Part Window ──────────────────────────────────────────────────────────

class ViewPartWindow(ctk.CTkToplevel):
    def __init__(self, master, part_data: dict):
        super().__init__(master)
        self.title(f"View — {part_data['label']}")
        self.geometry("560x520")
        self.configure(fg_color=BG_DEEP)

        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr,
                     text=f"  {part_data['label']}  —  {part_data.get('report_title', '')}",
                     font=FONT_CARD, text_color=TEXT_PRI).pack(side="left", padx=12)

        txt = ctk.CTkTextbox(self, fg_color=BG_INPUT, text_color=TEXT_PRI,
                              font=FONT_MONO, corner_radius=0)
        txt.pack(fill="both", expand=True, padx=12, pady=12)
        txt.insert("1.0", _build_part_summary(part_data))
        txt.configure(state="disabled")

        ctk.CTkButton(self, text="Close", command=self.destroy,
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=32).pack(pady=(0, 12))


# ── Text summary helper ───────────────────────────────────────────────────────

def _build_part_summary(r: dict) -> str:
    line = "─" * 52
    out  = f"{r.get('label', '')}\n{line}\n"
    out += f"{r.get('report_title', '')}\n"
    if r.get("report_subtitle"):
        out += f"{r['report_subtitle']}\n"
    if r.get("researcher_name"):
        out += f"by: {r['researcher_name']}\n"
    out += f"{line}\n\n"

    out += "RELIABILITY COEFFICIENT\n" + "─" * 28 + "\n"
    out += f"  Cronbach's Alpha (α): {_fmt(r.get('alpha'))}\n"
    out += f"  Interpretation:       {r.get('interpretation', '—')}\n\n"
    out += f"  Std. Error:           {_fmt(r.get('std_error'))}\n"
    out += f"  95% CI Lower:         {_fmt(r.get('ci_lower'))}\n"
    out += f"  95% CI Upper:         {_fmt(r.get('ci_upper'))}\n\n"

    out += "DESCRIPTIVE STATISTICS\n" + "─" * 28 + "\n"
    out += f"  Number of Items:       {r.get('n_items', '—')}\n"
    out += f"  Number of Respondents: {r.get('n_respondents', '—')}\n"
    out += f"  Avg Inter-item r:      {_fmt(r.get('avg_interitem_corr'))}\n\n"

    items = r.get("item_names", [])
    if items:
        out += f"ITEMS ({len(items)}):\n" + "─" * 28 + "\n"
        out += "\n".join(f"  {i+1}. {nm}" for i, nm in enumerate(items))
        out += "\n\n"

    if r.get("description"):
        out += f"DESCRIPTION:\n{r['description']}\n\n"

    return out


# ── DOCX Export (all parts) ───────────────────────────────────────────────────

def export_all_to_docx(parts: list, filepath: str):
    """
    Write all saved Cronbach's Alpha parts into a single APA-style DOCX.
    Two-column layout, same margins as the ANOVA exporter.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_section_margins(sect):
        sect.top_margin    = Inches(0.60)
        sect.bottom_margin = Inches(0.60)
        sect.left_margin   = Inches(0.65)
        sect.right_margin  = Inches(0.65)

    def _apply_2col(sectPr):
        old = sectPr.find(qn("w:cols"))
        if old is not None:
            sectPr.remove(old)
        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"),   "2")
        cols_el.set(qn("w:space"), "720")
        sectPr.append(cols_el)

    def _insert_section_break(doc, break_type="nextPage"):
        p_el = OxmlElement("w:p")
        pPr  = OxmlElement("w:pPr")
        sPr  = OxmlElement("w:sectPr")

        pg_mar = OxmlElement("w:pgMar")
        for attr, val in [("w:top",    "864"), ("w:bottom", "864"),
                           ("w:left",   "936"), ("w:right",  "936"),
                           ("w:header", "720"), ("w:footer", "720"),
                           ("w:gutter", "0")]:
            pg_mar.set(qn(attr), val)
        sPr.append(pg_mar)

        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"),   "2")
        cols_el.set(qn("w:space"), "720")
        sPr.append(cols_el)

        pg_type = OxmlElement("w:type")
        pg_type.set(qn("w:val"), break_type)
        sPr.append(pg_type)

        pPr.append(sPr)
        p_el.append(pPr)
        doc.element.body.append(p_el)

    def apa_borders(table):
        tbl  = table._tbl
        tblPr = tbl.tblPr
        tb   = OxmlElement("w:tblBorders")
        for bn in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tb.append(b)
        for bn, sz in [("top", "12"), ("bottom", "12")]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), sz)
            tb.append(b)
        tblPr.append(tb)

    def add_header_sep(table):
        for cell in table.rows[0].cells:
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            bot   = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            tcBorders.append(bot)
            tcPr.append(tcBorders)

    def cf(cell, text, bold=False, italic=False, size=9, align="left", color=None):
        para = cell.paragraphs[0]
        para.alignment = (WD_PARAGRAPH_ALIGNMENT.CENTER
                          if align == "center" else WD_PARAGRAPH_ALIGNMENT.LEFT)
        run = para.add_run(str(text))
        run.font.size = Pt(size)
        run.bold      = bold
        run.italic    = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    doc = Document()
    _set_section_margins(doc.sections[0])
    _apply_2col(doc.sections[0]._sectPr)

    for part_idx, r in enumerate(parts):
        if part_idx > 0:
            _insert_section_break(doc, "nextPage")

        lp = doc.add_heading(r.get("label", f"Part {part_idx + 1}"), level=1)
        lp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in lp.runs:
            run.font.size      = Pt(13)
            run.bold           = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        tp = doc.add_paragraph(r.get("report_title", "Unidimensional Reliability"))
        tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in tp.runs:
            run.bold           = True
            run.font.size      = Pt(14)

        if r.get("report_subtitle"):
            sp = doc.add_paragraph(r["report_subtitle"])
            sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sp.runs:
                run.italic    = True
                run.font.size = Pt(11)

        if r.get("researcher_name"):
            np_ = doc.add_paragraph(r["researcher_name"])
            np_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in np_.runs:
                run.italic    = True
                run.font.size = Pt(10)

        doc.add_paragraph()

        if r.get("description"):
            dp = doc.add_paragraph(r["description"])
            for run in dp.runs:
                run.font.size = Pt(9)
            doc.add_paragraph()

        ti = doc.add_paragraph("Frequentist Scale Reliability Statistics")
        if ti.runs:
            ti.runs[0].italic    = True
            ti.runs[0].font.size = Pt(9)

        tbl = doc.add_table(rows=3, cols=5)
        apa_borders(tbl)

        headers_row0 = ["", "", "", "95% CI", ""]
        for i, h in enumerate(headers_row0):
            cf(tbl.rows[0].cells[i], h, bold=True, size=8, align="center")
        tbl.rows[0].cells[3].merge(tbl.rows[0].cells[4])

        for i, h in enumerate(["Coefficient", "Estimate", "Std. Error", "Lower", "Upper"]):
            cf(tbl.rows[1].cells[i], h, bold=True, size=8, align="center")
        add_header_sep(tbl)

        cf(tbl.rows[2].cells[0], "Coefficient α",          size=9)
        cf(tbl.rows[2].cells[1], _fmt(r.get("alpha")),     size=9, align="center")
        cf(tbl.rows[2].cells[2], _fmt(r.get("std_error")), size=9, align="center")
        cf(tbl.rows[2].cells[3], _fmt(r.get("ci_lower")),  size=9, align="center")
        cf(tbl.rows[2].cells[4], _fmt(r.get("ci_upper")),  size=9, align="center")

        doc.add_paragraph()

        st_title = doc.add_paragraph("Summary Statistics")
        if st_title.runs:
            st_title.runs[0].italic    = True
            st_title.runs[0].font.size = Pt(9)

        st = doc.add_table(rows=5, cols=2)
        apa_borders(st)
        rows_data = [
            ("Statistic",               "Value",                           True),
            ("Number of Items",         str(r.get("n_items", "—")),        False),
            ("Number of Respondents",   str(r.get("n_respondents", "—")),  False),
            ("Average Inter-item Corr.",_fmt(r.get("avg_interitem_corr")), False),
            ("Reliability Interpretation", r.get("interpretation", "—"),   False),
        ]
        for row_idx, (label, value, is_hdr) in enumerate(rows_data):
            cf(st.rows[row_idx].cells[0], label, bold=is_hdr, size=9)
            cf(st.rows[row_idx].cells[1], value, bold=is_hdr, size=9, align="center")
        add_header_sep(st)

        doc.add_paragraph()

        items = r.get("item_names", [])
        if items:
            ip = doc.add_paragraph()
            ip.add_run(f"Items ({len(items)}): ").bold = True
            ip.runs[0].font.size = Pt(9)
            ip.add_run(", ".join(items)).font.size = Pt(9)

        if r.get("perfect_correlations"):
            pairs = [f"{p[0]} and {p[1]}" for p in r["perfect_correlations"]]
            note  = doc.add_paragraph(
                f"Note. Variables {', '.join(pairs)} correlated perfectly.")
            if note.runs:
                note.runs[0].italic    = True
                note.runs[0].font.size = Pt(8)

        doc.add_paragraph()

        fp_para = doc.add_paragraph()
        _f = fp_para.add_run(
            f"Part saved: {r.get('saved_at', '')}   |   "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}"
        )
        _f.font.size      = Pt(7)
        _f.font.color.rgb = RGBColor(128, 128, 128)
        _f.italic         = True

    body_sectPr = doc.element.body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        pg_mar = body_sectPr.find(qn("w:pgMar"))
        if pg_mar is None:
            pg_mar = OxmlElement("w:pgMar")
            body_sectPr.append(pg_mar)
        for attr, val in [("w:top",    "864"), ("w:bottom", "864"),
                           ("w:left",   "936"), ("w:right",  "936")]:
            pg_mar.set(qn(attr), val)
        _apply_2col(body_sectPr)

    doc.save(filepath)


# ── PDF Export (all parts) ────────────────────────────────────────────────────

def export_all_to_pdf(parts: list, filepath: str):
    """
    Write all saved Cronbach's Alpha parts into a single APA-style PDF.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, PageBreak, HRFlowable)
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc_pdf = SimpleDocTemplate(
        filepath, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch
    )

    styles = getSampleStyleSheet()

    Title      = ParagraphStyle("Title2",     parent=styles["Normal"],
                                fontSize=12,  alignment=TA_CENTER,
                                fontName="Helvetica-Bold", spaceAfter=2)
    Sub        = ParagraphStyle("Sub",        parent=styles["Normal"],
                                fontSize=10,  alignment=TA_CENTER,
                                fontName="Helvetica-Oblique", spaceAfter=2)
    SectionHdr = ParagraphStyle("SectionHdr", parent=styles["Normal"],
                                fontSize=10,  fontName="Helvetica-Bold",
                                spaceAfter=3, spaceBefore=8)
    Body       = ParagraphStyle("Body",       parent=styles["Normal"],
                                fontSize=9,   spaceAfter=3)
    Small      = ParagraphStyle("Small",      parent=styles["Normal"],
                                fontSize=7,   textColor=colors.grey)
    TableTitle = ParagraphStyle("TblTitle",   parent=styles["Normal"],
                                fontSize=9,   fontName="Helvetica-Oblique",
                                spaceAfter=3)

    story = []

    for part_idx, r in enumerate(parts):
        if part_idx > 0:
            story.append(PageBreak())

        story.append(Paragraph(
            r.get("report_title", "Unidimensional Reliability"), Title))
        if r.get("report_subtitle"):
            story.append(Paragraph(r["report_subtitle"], Sub))
        if r.get("researcher_name"):
            story.append(Paragraph(r["researcher_name"], Sub))
        story.append(Spacer(1, 8))

        if r.get("description"):
            story.append(Paragraph(r["description"], Body))
            story.append(Spacer(1, 6))

        story.append(Paragraph(
            "Frequentist Scale Reliability Statistics", TableTitle))

        t1_data = [
            ["", "", "", "95% CI", ""],
            ["Coefficient", "Estimate", "Std. Error", "Lower", "Upper"],
            ["Coefficient α",
             _fmt(r.get("alpha")),
             _fmt(r.get("std_error")),
             _fmt(r.get("ci_lower")),
             _fmt(r.get("ci_upper"))],
        ]

        # Soft highlight color keyed to interpretation
        _alpha_bg = {
            "Excellent":    "#d1fae5",
            "Good":         "#dcfce7",
            "Acceptable":   "#dbeafe",
            "Questionable": "#fef9c3",
            "Poor":         "#ffedd5",
            "Unacceptable": "#fee2e2",
        }.get(r.get("interpretation", ""), "#f3f4f6")

        t1 = Table(t1_data,
                   colWidths=[1.6*inch, 0.9*inch, 0.9*inch, 0.8*inch, 0.8*inch])
        t1.setStyle(TableStyle([
            ("SPAN",          (3, 0),  (4, 0)),
            ("FONTNAME",      (0, 0),  (-1, 1), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0),  (-1, -1), 8),
            ("ALIGN",         (0, 0),  (0, -1), "LEFT"),
            ("ALIGN",         (1, 0),  (-1, -1), "CENTER"),
            ("LINEABOVE",     (0, 0),  (-1, 0),  1.2, colors.black),
            ("LINEBELOW",     (0, 1),  (-1, 1),  0.6, colors.black),
            ("LINEBELOW",     (0, -1), (-1, -1), 1.2, colors.black),
            ("TOPPADDING",    (0, 0),  (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 4),
            ("BACKGROUND",    (1, 2),  (1, 2),   colors.HexColor(_alpha_bg)),
            ("FONTNAME",      (1, 2),  (1, 2),   "Helvetica-Bold"),
        ]))
        story.append(t1)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Summary Statistics", SectionHdr))
        t2_data = [
            ["Statistic",                  "Value"],
            ["Number of Items",            str(r.get("n_items", "—"))],
            ["Number of Respondents",      str(r.get("n_respondents", "—"))],
            ["Average Inter-item Corr.",   _fmt(r.get("avg_interitem_corr"))],
            ["Reliability Interpretation", r.get("interpretation", "—")],
        ]
        t2 = Table(t2_data, colWidths=[2.5*inch, 2.0*inch])
        t2.setStyle(TableStyle([
            ("FONTNAME",      (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0),  (-1, -1), 8),
            ("ALIGN",         (1, 0),  (-1, -1), "CENTER"),
            ("LINEABOVE",     (0, 0),  (-1, 0),  1.2, colors.black),
            ("LINEBELOW",     (0, 0),  (-1, 0),  0.6, colors.black),
            ("LINEBELOW",     (0, -1), (-1, -1), 1.2, colors.black),
            ("TOPPADDING",    (0, 0),  (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 4),
        ]))
        story.append(t2)
        story.append(Spacer(1, 10))

        items = r.get("item_names", [])
        if items:
            story.append(Paragraph("Items", SectionHdr))
            story.append(Paragraph(
                ", ".join(f"{i+1}. {nm}" for i, nm in enumerate(items)), Body))
            story.append(Spacer(1, 6))

        if r.get("perfect_correlations"):
            pairs      = [f"{p[0]} and {p[1]}" for p in r["perfect_correlations"]]
            note_style = ParagraphStyle("note", parent=styles["Normal"],
                                        fontSize=8, fontName="Helvetica-Oblique")
            story.append(Paragraph(
                f"Note. Variables {', '.join(pairs)} correlated perfectly.",
                note_style))
            story.append(Spacer(1, 6))

        # ── Interpretation Guide ──────────────────────────────────────────────
        _interp_row = {
            "Excellent": 1, "Good": 2, "Acceptable": 3,
            "Questionable": 4, "Poor": 5, "Unacceptable": 6,
        }
        highlight_row = _interp_row.get(r.get("interpretation", ""), 0)

        story.append(Paragraph("Interpretation Guide", SectionHdr))
        guide_data = [
            ["Range",            "Interpretation"],
            ["α ≥ 0.90",         "Excellent internal consistency"],
            ["0.80 ≤ α < 0.90",  "Good internal consistency"],
            ["0.70 ≤ α < 0.80",  "Acceptable internal consistency"],
            ["0.60 ≤ α < 0.70",  "Questionable internal consistency"],
            ["0.50 ≤ α < 0.60",  "Poor internal consistency"],
            ["α < 0.50",         "Unacceptable internal consistency"],
        ]
        tg_style = [
            ("FONTNAME",      (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0),  (-1, -1), 8),
            ("ALIGN",         (0, 0),  (-1, -1), "LEFT"),
            ("LINEABOVE",     (0, 0),  (-1, 0),  1.2, colors.black),
            ("LINEBELOW",     (0, 0),  (-1, 0),  0.6, colors.black),
            ("LINEBELOW",     (0, -1), (-1, -1), 1.2, colors.black),
            ("TOPPADDING",    (0, 0),  (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 3),
        ]
        if highlight_row:
            tg_style.append(
                ("BACKGROUND", (0, highlight_row), (-1, highlight_row),
                 colors.HexColor(_alpha_bg))
            )
        tg = Table(guide_data, colWidths=[1.8*inch, 3.0*inch])
        tg.setStyle(TableStyle(tg_style))
        story.append(tg)
        story.append(Spacer(1, 10))

    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}  |  "
        f"{len(parts)} part(s) total", Small))

    doc_pdf.build(story)