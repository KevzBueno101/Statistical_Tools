"""
Chi-Square Test Module
JASP-Style Statistical Analysis

Implements:
- Chi-Square Test of Independence
- McNemar's Test (for 2x2 paired data)
- Chi-Square Goodness-of-Fit Test
- APA-style reporting and tables
- Professional reporting (PDF/DOCX)

Author: Statistical Analysis Suite
Version: 2.0.0
"""

import customtkinter as ctk
from tkinter import messagebox, filedialog
import numpy as np
import pandas as pd
from scipy import stats
from datetime import datetime
import os

# PDF and Word export libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ChiSquareTestApp(ctk.CTk):
    """Chi-Square Test Application - JASP Style"""
    
    def __init__(self):
        super().__init__()
        
        # Configure window
        self.title("Chi-Square Test (JASP-Style)")
        self.geometry("700x700")
        self.resizable(True, True)
        self.minsize(1000, 700)
        
        # Initialize variables
        self.num_rows = 2
        self.num_cols = 2
        self.alpha = 0.05
        self.results = None
        
        # Center and create UI
        self.center_window()
        self.create_ui()
        
    def center_window(self):
        """Center window on screen"""
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (w // 2)
        y = (self.winfo_screenheight() // 2) - (h // 2)
        self.geometry(f'{w}x{h}+{x}+{y}')
    
    def create_ui(self):
        """Create main UI"""
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Header
        header = ctk.CTkFrame(main, fg_color="#1f538d", corner_radius=10)
        header.pack(fill="x", pady=(0, 20))
        ctk.CTkLabel(header, text="Chi-Square Test", font=ctk.CTkFont(size=28, weight="bold"), text_color="white").pack(pady=(15, 5))
        ctk.CTkLabel(header, text="JASP-Style Categorical Data Analysis", font=ctk.CTkFont(size=13), text_color="#e0e0e0").pack(pady=(0, 15))
        
        # Content
        content = ctk.CTkFrame(main, fg_color="transparent")
        content.pack(fill="both", expand=True)
        
        left = ctk.CTkFrame(content, fg_color="transparent")
        left.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        right = ctk.CTkFrame(content, fg_color="transparent")
        right.pack(side="right", fill="both", expand=True, padx=(10, 0))
        
        self.create_input_panel(left)
        self.create_results_panel(right)
        
    def create_input_panel(self, parent):
        """Create input panel"""
        ctrl = ctk.CTkFrame(parent)
        ctrl.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(ctrl, text="Test Configuration", font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=15, pady=(10, 5))
        
        # Title and Subtitle
        f0 = ctk.CTkFrame(ctrl, fg_color="transparent")
        f0.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(f0, text="Title:", width=100, anchor="w").pack(side="left", padx=(0, 10))
        self.title_entry = ctk.CTkEntry(f0, width=300, placeholder_text="Enter report title (optional)")
        self.title_entry.pack(side="left")
        
        f0b = ctk.CTkFrame(ctrl, fg_color="transparent")
        f0b.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(f0b, text="Subtitle:", width=100, anchor="w").pack(side="left", padx=(0, 10))
        self.subtitle_entry = ctk.CTkEntry(f0b, width=300, placeholder_text="Enter report subtitle (optional)")
        self.subtitle_entry.pack(side="left")
        
        # Test type
        f1 = ctk.CTkFrame(ctrl, fg_color="transparent")
        f1.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(f1, text="Test Type:", width=100, anchor="w").pack(side="left", padx=(0, 10))
        self.test_type_var = ctk.StringVar(value="independence")
        ctk.CTkOptionMenu(f1, variable=self.test_type_var, 
                         values=["independence", "mcnemar", "goodness-of-fit"], 
                         width=200,
                         command=self.on_test_type_change).pack(side="left")
        
        # Alpha
        f2 = ctk.CTkFrame(ctrl, fg_color="transparent")
        f2.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(f2, text="Alpha (α):", width=100, anchor="w").pack(side="left", padx=(0, 10))
        self.alpha_entry = ctk.CTkEntry(f2, width=100)
        self.alpha_entry.insert(0, "0.05")
        self.alpha_entry.pack(side="left")
        
        # Dimensions
        f3 = ctk.CTkFrame(ctrl, fg_color="transparent")
        f3.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(f3, text="Table Size:", width=100, anchor="w").pack(side="left", padx=(0, 10))
        ctk.CTkLabel(f3, text="Rows:").pack(side="left", padx=(0, 5))
        self.rows_entry = ctk.CTkEntry(f3, width=60)
        self.rows_entry.insert(0, "2")
        self.rows_entry.pack(side="left", padx=(0, 15))
        ctk.CTkLabel(f3, text="Cols:").pack(side="left", padx=(0, 5))
        self.cols_entry = ctk.CTkEntry(f3, width=60)
        self.cols_entry.insert(0, "2")
        self.cols_entry.pack(side="left", padx=(0, 15))
        ctk.CTkButton(f3, text="Update Table", command=self.update_table_size, width=120, height=28).pack(side="left")
        
        # File ops
        f4 = ctk.CTkFrame(ctrl, fg_color="transparent")
        f4.pack(fill="x", padx=15, pady=(10, 10))
        ctk.CTkButton(f4, text="📁 Load CSV/Excel", command=self.load_data_file, width=150, height=32).pack(side="left", padx=(0, 10))
        ctk.CTkButton(f4, text="🗑️ Clear Table", command=self.clear_table, width=150, height=32, fg_color="#d32f2f", hover_color="#b71c1c").pack(side="left")
        
        # Table
        inp = ctk.CTkFrame(parent)
        inp.pack(fill="both", expand=True)
        ctk.CTkLabel(inp, text="Contingency Table (Observed Frequencies)", font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=15, pady=(10, 10))
        self.table_scroll = ctk.CTkScrollableFrame(inp, fg_color="#2b2b2b")
        self.table_scroll.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        self.create_data_table()
    
    def on_test_type_change(self, choice):
        """Handle test type change"""
        if choice == "mcnemar":
            # Force 2x2 table for McNemar's test
            self.rows_entry.delete(0, 'end')
            self.rows_entry.insert(0, "2")
            self.cols_entry.delete(0, 'end')
            self.cols_entry.insert(0, "2")
            self.num_rows = 2
            self.num_cols = 2
            self.create_data_table()
            messagebox.showinfo("Info", "McNemar's test requires a 2x2 table for paired categorical data.")
        
    def create_results_panel(self, parent):
        """Create results panel"""
        res = ctk.CTkFrame(parent)
        res.pack(fill="both", expand=True)
        
        ctk.CTkLabel(res, text="Analysis Results", font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=15, pady=(10, 10))
        
        self.results_text = ctk.CTkTextbox(res, font=ctk.CTkFont(size=12), wrap="word", fg_color="#2b2b2b")
        self.results_text.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        act = ctk.CTkFrame(res, fg_color="transparent")
        act.pack(fill="x", padx=15, pady=(0, 15))
        
        ctk.CTkButton(act, text="▶ Run Analysis", command=self.run_analysis, width=180, height=40, font=ctk.CTkFont(size=14, weight="bold"), fg_color="#2e7d32", hover_color="#1b5e20").pack(pady=5)
        ctk.CTkLabel(act, text="Export Results:", font=ctk.CTkFont(size=12)).pack(pady=(15, 5))
        
        exp = ctk.CTkFrame(act, fg_color="transparent")
        exp.pack()
        ctk.CTkButton(exp, text="📄 Export PDF", command=self.export_pdf, width=140, height=35, fg_color="#1976d2", hover_color="#1565c0").pack(side="left", padx=5)
        ctk.CTkButton(exp, text="📝 Export Word", command=self.export_docx, width=140, height=35, fg_color="#7b1fa2", hover_color="#6a1b9a").pack(side="left", padx=5)
        
    def create_data_table(self):
        """Create data entry table"""
        for w in self.table_scroll.winfo_children():
            w.destroy()
        
        t = ctk.CTkFrame(self.table_scroll, fg_color="transparent")
        t.pack(fill="both", expand=True)
        
        self.cell_entries = []
        ctk.CTkLabel(t, text="", width=100).grid(row=0, column=0, padx=5, pady=5)
        
        self.col_label_entries = []
        for j in range(self.num_cols):
            e = ctk.CTkEntry(t, width=100)
            e.insert(0, f"Col_{j+1}")
            e.grid(row=0, column=j+1, padx=5, pady=5)
            self.col_label_entries.append(e)
        
        self.row_label_entries = []
        for i in range(self.num_rows):
            row_entries = []
            le = ctk.CTkEntry(t, width=100)
            le.insert(0, f"Row_{i+1}")
            le.grid(row=i+1, column=0, padx=5, pady=5)
            self.row_label_entries.append(le)
            
            for j in range(self.num_cols):
                ce = ctk.CTkEntry(t, width=100)
                ce.insert(0, "0")
                ce.grid(row=i+1, column=j+1, padx=5, pady=5)
                row_entries.append(ce)
            
            self.cell_entries.append(row_entries)
    
    def update_table_size(self):
        """Update table dimensions"""
        try:
            r, c = int(self.rows_entry.get()), int(self.cols_entry.get())
            
            # Check for McNemar's test
            if self.test_type_var.get() == "mcnemar":
                if r != 2 or c != 2:
                    messagebox.showerror("Error", "McNemar's test requires exactly a 2x2 table.")
                    return
            
            if r < 2 or c < 2:
                messagebox.showerror("Error", "Minimum 2x2 table required.")
                return
            if r > 10 or c > 10:
                messagebox.showerror("Error", "Maximum 10x10 table allowed.")
                return
            self.num_rows, self.num_cols = r, c
            self.create_data_table()
        except ValueError:
            messagebox.showerror("Error", "Enter valid integers.")
    
    def get_table_data(self):
        """Extract table data"""
        try:
            rows = [e.get().strip() for e in self.row_label_entries]
            cols = [e.get().strip() for e in self.col_label_entries]
            data = [[float(e.get().strip() or 0) for e in row] for row in self.cell_entries]
            return np.array(data), rows, cols
        except ValueError as e:
            messagebox.showerror("Error", f"Invalid data: {e}")
            return None, None, None
    
    def clear_table(self):
        """Clear all data"""
        for row in self.cell_entries:
            for e in row:
                e.delete(0, 'end')
                e.insert(0, "0")
    
    def load_data_file(self):
        """Load CSV/Excel"""
        fp = filedialog.askopenfilename(filetypes=[("CSV", "*.csv"), ("Excel", "*.xlsx *.xls"), ("All", "*.*")])
        if not fp:
            return
        try:
            df = pd.read_csv(fp, index_col=0) if fp.endswith('.csv') else pd.read_excel(fp, index_col=0)
            self.num_rows, self.num_cols = len(df), len(df.columns)
            
            # Check for McNemar's test
            if self.test_type_var.get() == "mcnemar" and (self.num_rows != 2 or self.num_cols != 2):
                messagebox.showerror("Error", "McNemar's test requires a 2x2 table.")
                return
            
            self.rows_entry.delete(0, 'end')
            self.rows_entry.insert(0, str(self.num_rows))
            self.cols_entry.delete(0, 'end')
            self.cols_entry.insert(0, str(self.num_cols))
            self.create_data_table()
            
            for i, rl in enumerate(df.index):
                self.row_label_entries[i].delete(0, 'end')
                self.row_label_entries[i].insert(0, str(rl))
            for j, cl in enumerate(df.columns):
                self.col_label_entries[j].delete(0, 'end')
                self.col_label_entries[j].insert(0, str(cl))
            for i in range(self.num_rows):
                for j in range(self.num_cols):
                    self.cell_entries[i][j].delete(0, 'end')
                    self.cell_entries[i][j].insert(0, str(df.iloc[i, j]))
            messagebox.showinfo("Success", f"Loaded: {fp}")
        except Exception as e:
            messagebox.showerror("Error", f"Load failed: {e}")
    
    def format_p_value(self, p):
        """Format p-value according to APA style (for narrative text)"""
        if p < 0.001:
            return "< .001"
        else:
            return f"= {p:.3f}"
    
    def format_p_value_table(self, p):
        """Format p-value for tables (without equals sign)"""
        if p < 0.001:
            return "< .001"
        else:
            # Format as .XXX (removing leading zero)
            formatted = f"{p:.3f}"
            if formatted.startswith("0."):
                return "." + formatted[2:]
            return formatted
    
    def run_analysis(self):
        """Run Chi-Square analysis"""
        try:
            self.alpha = float(self.alpha_entry.get())
            if not 0 < self.alpha < 1:
                messagebox.showerror("Error", "Alpha must be 0 < α < 1")
                return
        except ValueError:
            messagebox.showerror("Error", "Invalid alpha value")
            return
        
        obs, rows, cols = self.get_table_data()
        if obs is None:
            return
        
        if np.any(obs < 0):
            messagebox.showerror("Error", "Negative frequencies not allowed")
            return
        if np.sum(obs) == 0:
            messagebox.showerror("Error", "Table cannot be empty")
            return
        
        self.row_labels, self.col_labels = rows, cols
        
        test_type = self.test_type_var.get()
        
        if test_type == "independence":
            self.results = self.chi_square_independence(obs)
        elif test_type == "mcnemar":
            if obs.shape != (2, 2):
                messagebox.showerror("Error", "McNemar's test requires exactly a 2x2 table")
                return
            self.results = self.mcnemar_test(obs)
        else:
            self.results = self.chi_square_goodness_of_fit(obs)
        
        self.display_results()
    
    def chi_square_independence(self, obs):
        """Chi-Square test of independence"""
        chi2, p, dof, exp = stats.chi2_contingency(obs)
        n = np.sum(obs)
        min_dim = min(obs.shape[0] - 1, obs.shape[1] - 1)
        v = np.sqrt(chi2 / (n * min_dim))
        
        if min_dim == 1:
            eff = "negligible" if v < 0.10 else "small" if v < 0.30 else "medium" if v < 0.50 else "large"
        else:
            eff = "negligible" if v < 0.07 else "small" if v < 0.21 else "medium" if v < 0.35 else "large"
        
        warn = "Warning: Some expected frequencies < 5. Results may be unreliable." if np.any(exp < 5) else None
        
        return {
            'test_type': 'independence',
            'test_name': 'Chi-Square Test of Independence',
            'chi2': chi2,
            'p_value': p,
            'dof': dof,
            'observed': obs,
            'expected': exp,
            'cramers_v': v,
            'effect_size': eff,
            'alpha': self.alpha,
            'warning': warn
        }
    
    def mcnemar_test(self, obs):
        """McNemar's test for paired nominal data"""
        # obs should be 2x2: [[a, b], [c, d]]
        b = obs[0, 1]
        c = obs[1, 0]
        
        # McNemar's test statistic
        result = stats.mcnemar([[obs[0,0], obs[0,1]], [obs[1,0], obs[1,1]]], exact=False)
        chi2 = result.statistic
        p = result.pvalue
        
        # Effect size (odds ratio for discordant pairs)
        if c > 0:
            odds_ratio = b / c
        else:
            odds_ratio = float('inf') if b > 0 else 1.0
        
        warn = None
        if b + c < 25:
            warn = "Warning: Small number of discordant pairs (< 25). Consider using exact McNemar's test."
        
        return {
            'test_type': 'mcnemar',
            'test_name': "McNemar's Test",
            'chi2': chi2,
            'p_value': p,
            'dof': 1,
            'observed': obs,
            'b': b,
            'c': c,
            'odds_ratio': odds_ratio,
            'alpha': self.alpha,
            'warning': warn
        }
    
    def chi_square_goodness_of_fit(self, obs):
        """Chi-Square goodness-of-fit test"""
        flat = obs.flatten()
        exp_flat = np.full_like(flat, np.mean(flat))
        chi2, p = stats.chisquare(flat, exp_flat)
        dof = len(flat) - 1
        w = np.sqrt(chi2 / np.sum(flat))
        eff = "small" if w < 0.10 else "medium" if w < 0.30 else "large"
        exp = exp_flat.reshape(obs.shape)
        warn = "Warning: Some expected frequencies < 5. Results may be unreliable." if np.any(exp < 5) else None
        
        return {
            'test_type': 'goodness-of-fit',
            'test_name': 'Chi-Square Goodness-of-Fit Test',
            'chi2': chi2,
            'p_value': p,
            'dof': dof,
            'observed': obs,
            'expected': exp,
            'effect_size_w': w,
            'effect_size': eff,
            'alpha': self.alpha,
            'warning': warn
        }
    
    def display_results(self):
        """Display results"""
        self.results_text.delete("1.0", "end")
        r = self.results
        
        # Add custom title and subtitle if provided
        custom_title = self.title_entry.get().strip()
        custom_subtitle = self.subtitle_entry.get().strip()
        
        if custom_title:
            self.results_text.insert("end", "═" * 50 + "\n")
            self.results_text.insert("end", f"{custom_title.upper()}\n")
            self.results_text.insert("end", "═" * 50 + "\n\n")
        
        if custom_subtitle:
            self.results_text.insert("end", f"{custom_subtitle}\n\n")
        
        self.results_text.insert("end", "═" * 50 + "\n")
        self.results_text.insert("end", f"{r['test_name'].upper()}\n")
        self.results_text.insert("end", "═" * 50 + "\n\n")
        
        self.results_text.insert("end", f"Alpha: {r['alpha']}\n")
        self.results_text.insert("end", f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        self.results_text.insert("end", "─" * 50 + "\nCONTINGENCY TABLES:\n" + "─" * 50 + "\n")
        self.results_text.insert("end", self.format_table_with_totals(r['observed']) + "\n\n")
        
        self.results_text.insert("end", "─" * 50 + "\nTEST STATISTICS:\n" + "─" * 50 + "\n")
        stats_df = pd.DataFrame({
            'Statistic': ['χ²', 'df', 'p'],
            'Value': [f"{r['chi2']:.4f}", f"{r['dof']}", self.format_p_value_table(r['p_value'])]
        })
        self.results_text.insert("end", stats_df.to_string(index=False) + "\n\n")
        
        if r['test_type'] != 'mcnemar':
            self.results_text.insert("end", "─" * 50 + "\nEXPECTED FREQUENCIES:\n" + "─" * 50 + "\n")
            self.results_text.insert("end", self.format_table_with_totals(r['expected']) + "\n\n")
        
        self.results_text.insert("end", "─" * 50 + "\nEFFECT SIZE STATISTICS:\n" + "─" * 50 + "\n")
        
        if r['test_type'] == 'independence':
            self.results_text.insert("end", f"Cramer's V = {r['cramers_v']:.4f}\n")
            self.results_text.insert("end", f"Effect Size: {r['effect_size'].capitalize()}\n\n")
        elif r['test_type'] == 'mcnemar':
            self.results_text.insert("end", f"Discordant pairs (b): {r['b']:.0f}\n")
            self.results_text.insert("end", f"Discordant pairs (c): {r['c']:.0f}\n")
            if r['odds_ratio'] != float('inf'):
                self.results_text.insert("end", f"Odds Ratio: {r['odds_ratio']:.4f}\n\n")
            else:
                self.results_text.insert("end", f"Odds Ratio: Undefined (c = 0)\n\n")
        else:
            self.results_text.insert("end", f"Effect Size (w) = {r['effect_size_w']:.4f}\n")
            self.results_text.insert("end", f"Effect Size: {r['effect_size'].capitalize()}\n\n")
        
        if r.get('warning'):
            self.results_text.insert("end", f"⚠️ {r['warning']}\n\n")
        
        self.results_text.insert("end", "─" * 50 + "\nDECISION:\n" + "─" * 50 + "\n")
        if r['p_value'] < r['alpha']:
            self.results_text.insert("end", f"✓ REJECT H₀ (p {self.format_p_value(r['p_value'])} < α = {r['alpha']})\n\n")
        else:
            self.results_text.insert("end", f"✗ FAIL TO REJECT H₀ (p {self.format_p_value(r['p_value'])} ≥ α = {r['alpha']})\n\n")
        
        self.results_text.insert("end", "─" * 50 + "\nINTERPRETATION:\n" + "─" * 50 + "\n")
        self.results_text.insert("end", self.generate_interpretation() + "\n\n" + "═" * 50 + "\n")
    
    def format_table(self, data):
        """Format table"""
        return pd.DataFrame(data, index=self.row_labels, columns=self.col_labels).to_string()
    
    def format_table_with_totals(self, data):
        """Format table with row and column totals"""
        # Create dataframe
        df = pd.DataFrame(data, index=self.row_labels, columns=self.col_labels)
        
        # Add row totals
        df['Total'] = df.sum(axis=1)
        
        # Add column totals
        col_totals = df.sum(axis=0)
        df.loc['Total'] = col_totals
        
        return df.to_string()
    
    def generate_interpretation(self):
        """APA interpretation"""
        r = self.results
        p_formatted = self.format_p_value(r['p_value'])
        
        if r['test_type'] == 'independence':
            if r['p_value'] < r['alpha']:
                return f"A chi-square test of independence revealed a statistically significant association between the row and column variables, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}, Cramer's V = {r['cramers_v']:.2f}. The effect size was {r['effect_size']}, indicating a {r['effect_size']} relationship between the variables."
            else:
                return f"A chi-square test of independence showed no statistically significant association between the variables, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}, Cramer's V = {r['cramers_v']:.2f}. There was insufficient evidence to conclude that the variables are related."
        elif r['test_type'] == 'mcnemar':
            if r['p_value'] < r['alpha']:
                return f"McNemar's test indicated a statistically significant difference in the marginal frequencies, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}. The discordant pairs showed a significant asymmetry, suggesting that the change in response was not due to chance."
            else:
                return f"McNemar's test showed no statistically significant difference in the marginal frequencies, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}. There was insufficient evidence to conclude that the paired responses differed significantly."
        else:
            if r['p_value'] < r['alpha']:
                return f"A chi-square goodness-of-fit test indicated that observed frequencies differed significantly from expected frequencies, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}, w = {r['effect_size_w']:.2f}. The effect size was {r['effect_size']}."
            else:
                return f"A chi-square goodness-of-fit test showed that observed frequencies did not differ significantly from expected frequencies, χ²({r['dof']}) = {r['chi2']:.2f}, p {p_formatted}, w = {r['effect_size_w']:.2f}. The data were consistent with the expected distribution."
    
    def export_pdf(self):
        """Export to PDF"""
        if not self.results:
            messagebox.showerror("Error", "Run analysis first")
            return
        if not PDF_AVAILABLE:
            messagebox.showerror("Error", "Install reportlab: pip install reportlab")
            return
        
        # Ask user where to save the file
        default_filename = f"chi_square_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.pdf"
        fp = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            initialfile=default_filename,
            title="Save PDF Report"
        )
        
        if not fp:  # User cancelled
            return
        
        try:
            doc = SimpleDocTemplate(fp, pagesize=letter)
            story, styles = [], getSampleStyleSheet()
            
            # Add custom title and subtitle if provided
            custom_title = self.title_entry.get().strip()
            custom_subtitle = self.subtitle_entry.get().strip()
            
            if custom_title:
                story.append(Paragraph(custom_title, styles['Title']))
                story.append(Spacer(1, 0.1*inch))
            
            if custom_subtitle:
                subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#555555'))
                story.append(Paragraph(custom_subtitle, subtitle_style))
                story.append(Spacer(1, 0.2*inch))
            
            # Test name
            story.append(Paragraph(f"{self.results['test_name']}", styles['Title']))
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph(f"Alpha: {self.results['alpha']}", styles['Normal']))
            story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Observed - APA Style with totals (no colored tables)
            story.append(Paragraph("Contingency Tables", styles['Heading2']))
            
            # Build table with totals
            obs_data = [[''] + self.col_labels + ['Total']]
            for i in range(self.num_rows):
                row_total = sum(self.results['observed'][i,:])
                obs_data.append([self.row_labels[i]] + [f"{self.results['observed'][i,j]:.0f}" for j in range(self.num_cols)] + [f"{row_total:.0f}"])
            
            # Add column totals
            col_totals = ['Total'] + [f"{sum(self.results['observed'][:,j]):.0f}" for j in range(self.num_cols)]
            grand_total = f"{np.sum(self.results['observed']):.0f}"
            col_totals.append(grand_total)
            obs_data.append(col_totals)
            
            t1 = Table(obs_data)
            t1.setStyle(TableStyle([
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
                ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                ('FONTNAME', (-1,0), (-1,-1), 'Helvetica-Bold'),
                ('LINEBELOW', (0,0), (-1,0), 1, colors.black),
                ('LINEABOVE', (0,0), (-1,0), 1, colors.black),
                ('LINEABOVE', (0,-1), (-1,-1), 1, colors.black),
                ('LINEBELOW', (0,-1), (-1,-1), 1, colors.black),
            ]))
            story.append(t1)
            story.append(Spacer(1, 0.2*inch))
            
            # Add statistics table (APA style)
            story.append(Paragraph("Test Statistics", styles['Heading2']))
            stats_data = [
                ['Statistic', 'Value'],
                ['χ²', f'{self.results["chi2"]:.4f}'],
                ['df', f'{self.results["dof"]}'],
                ['p', f'{self.format_p_value_table(self.results["p_value"])}']
            ]
            stats_table = Table(stats_data)
            stats_table.setStyle(TableStyle([
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('LINEBELOW', (0,0), (-1,0), 1, colors.black),
                ('LINEABOVE', (0,0), (-1,0), 1, colors.black),
                ('LINEBELOW', (0,-1), (-1,-1), 1, colors.black),
            ]))
            story.append(stats_table)
            story.append(Spacer(1, 0.2*inch))

            # Expected (if not McNemar's)
            if self.results['test_type'] != 'mcnemar':
                story.append(Paragraph("Expected Frequencies", styles['Heading2']))
                
                # Build table with totals
                exp_data = [[''] + self.col_labels + ['Total']]
                for i in range(self.num_rows):
                    row_total = sum(self.results['expected'][i,:])
                    exp_data.append([self.row_labels[i]] + [f"{self.results['expected'][i,j]:.2f}" for j in range(self.num_cols)] + [f"{row_total:.2f}"])
                
                # Add column totals
                col_totals = ['Total'] + [f"{sum(self.results['expected'][:,j]):.2f}" for j in range(self.num_cols)]
                grand_total = f"{np.sum(self.results['expected']):.2f}"
                col_totals.append(grand_total)
                exp_data.append(col_totals)
                
                t2 = Table(exp_data)
                t2.setStyle(TableStyle([
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
                    ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                    ('FONTNAME', (-1,0), (-1,-1), 'Helvetica-Bold'),
                    ('LINEBELOW', (0,0), (-1,0), 1, colors.black),
                    ('LINEABOVE', (0,0), (-1,0), 1, colors.black),
                    ('LINEABOVE', (0,-1), (-1,-1), 1, colors.black),
                    ('LINEBELOW', (0,-1), (-1,-1), 1, colors.black),
                ]))
                story.append(t2)
                story.append(Spacer(1, 0.2*inch))
            
            # Stats (Effect sizes only, since chi2, df, p are in table)
            story.append(Paragraph("Effect Size Statistics", styles['Heading2']))
            if self.results['test_type'] == 'independence':
                story.append(Paragraph(f"Cramer's V = {self.results['cramers_v']:.4f}", styles['Normal']))
            elif self.results['test_type'] == 'mcnemar':
                story.append(Paragraph(f"Discordant pairs (b) = {self.results['b']:.0f}", styles['Normal']))
                story.append(Paragraph(f"Discordant pairs (c) = {self.results['c']:.0f}", styles['Normal']))
            else:
                story.append(Paragraph(f"Effect Size (w) = {self.results['effect_size_w']:.4f}", styles['Normal']))
            
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph("Interpretation", styles['Heading2']))
            story.append(Paragraph(self.generate_interpretation(), styles['Normal']))
            
            if self.results.get('warning'):
                story.append(Spacer(1, 0.2*inch))
            
            story.append(Spacer(1, 0.3*inch))
            footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, fontName='Helvetica-Oblique')
            story.append(Paragraph(f"Saved: {fp}", footer_style))
            story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
            
            doc.build(story)
            messagebox.showinfo("Success", f"PDF exported:\n{fp}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")
    
    def export_docx(self):
        """Export to Word"""
        if not self.results:
            messagebox.showerror("Error", "Run analysis first")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "Install python-docx: pip install python-docx")
            return
        
        # Ask user where to save the file
        default_filename = f"chi_square_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        fp = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=default_filename,
            title="Save Word Report"
        )
        
        if not fp:  # User cancelled
            return
        
        try:
            doc = Document()
            
            # Add custom title and subtitle if provided
            custom_title = self.title_entry.get().strip()
            custom_subtitle = self.subtitle_entry.get().strip()
            
            if custom_title:
                title_para = doc.add_heading(custom_title, 0)
            
            if custom_subtitle:
                subtitle_para = doc.add_paragraph(custom_subtitle)
                subtitle_para.runs[0].font.size = Pt(14)
                subtitle_para.runs[0].font.color.rgb = RGBColor(85, 85, 85)
                subtitle_para.runs[0].bold = True
            
            # Test name
            doc.add_heading(f"{self.results['test_name']}", 0)
            doc.add_paragraph(f"Alpha: {self.results['alpha']}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Observed - APA Style table with totals
            doc.add_heading('Observed Frequencies', 2)
            t1 = doc.add_table(rows=self.num_rows+2, cols=self.num_cols+2)
            t1.style = 'Light Grid'
            
            # Header row
            for j, c in enumerate(self.col_labels):
                t1.rows[0].cells[j+1].text = c
            t1.rows[0].cells[self.num_cols+1].text = 'Total'
            
            # Data rows with row totals
            for i, r in enumerate(self.row_labels):
                t1.rows[i+1].cells[0].text = r
                row_total = 0
                for j in range(self.num_cols):
                    value = self.results['observed'][i,j]
                    t1.rows[i+1].cells[j+1].text = f"{value:.0f}"
                    row_total += value
                t1.rows[i+1].cells[self.num_cols+1].text = f"{row_total:.0f}"
            
            # Total row
            t1.rows[self.num_rows+1].cells[0].text = 'Total'
            grand_total = 0
            for j in range(self.num_cols):
                col_total = sum(self.results['observed'][:,j])
                t1.rows[self.num_rows+1].cells[j+1].text = f"{col_total:.0f}"
                grand_total += col_total
            t1.rows[self.num_rows+1].cells[self.num_cols+1].text = f"{grand_total:.0f}"
            
            # Add statistics table (APA style)
            doc.add_heading('Test Statistics', 2)
            stats_table = doc.add_table(rows=4, cols=2)
            stats_table.style = 'Light Grid'
            
            # Header
            stats_table.rows[0].cells[0].text = 'Statistic'
            stats_table.rows[0].cells[1].text = 'Value'
            
            # Data
            stats_table.rows[1].cells[0].text = 'χ²'
            stats_table.rows[1].cells[1].text = f'{self.results["chi2"]:.4f}'
            stats_table.rows[2].cells[0].text = 'df'
            stats_table.rows[2].cells[1].text = f'{self.results["dof"]}'
            stats_table.rows[3].cells[0].text = 'p'
            stats_table.rows[3].cells[1].text = f'{self.format_p_value_table(self.results["p_value"])}'
            
            # Expected (if not McNemar's)
            if self.results['test_type'] != 'mcnemar':
                doc.add_heading('Expected Frequencies', 2)
                t2 = doc.add_table(rows=self.num_rows+2, cols=self.num_cols+2)
                t2.style = 'Light Grid'
                
                # Header row
                for j, c in enumerate(self.col_labels):
                    t2.rows[0].cells[j+1].text = c
                t2.rows[0].cells[self.num_cols+1].text = 'Total'
                
                # Data rows with row totals
                for i, r in enumerate(self.row_labels):
                    t2.rows[i+1].cells[0].text = r
                    row_total = 0
                    for j in range(self.num_cols):
                        value = self.results['expected'][i,j]
                        t2.rows[i+1].cells[j+1].text = f"{value:.2f}"
                        row_total += value
                    t2.rows[i+1].cells[self.num_cols+1].text = f"{row_total:.2f}"
                
                # Total row
                t2.rows[self.num_rows+1].cells[0].text = 'Total'
                grand_total = 0
                for j in range(self.num_cols):
                    col_total = sum(self.results['expected'][:,j])
                    t2.rows[self.num_rows+1].cells[j+1].text = f"{col_total:.2f}"
                    grand_total += col_total
                t2.rows[self.num_rows+1].cells[self.num_cols+1].text = f"{grand_total:.2f}"
            
            doc.add_heading('Effect Size Statistics', 2)
            if self.results['test_type'] == 'independence':
                doc.add_paragraph(f"Cramer's V = {self.results['cramers_v']:.4f}")
                doc.add_paragraph(f"Effect Size: {self.results['effect_size'].capitalize()}")
            elif self.results['test_type'] == 'mcnemar':
                doc.add_paragraph(f"Discordant pairs (b) = {self.results['b']:.0f}")
                doc.add_paragraph(f"Discordant pairs (c) = {self.results['c']:.0f}")
                if self.results['odds_ratio'] != float('inf'):
                    doc.add_paragraph(f"Odds Ratio = {self.results['odds_ratio']:.4f}")
            else:
                doc.add_paragraph(f"Effect Size (w) = {self.results['effect_size_w']:.4f}")
                doc.add_paragraph(f"Effect Size: {self.results['effect_size'].capitalize()}")

            doc.add_heading('Decision', 2)
            if self.results['p_value'] < self.results['alpha']:
                doc.add_paragraph(
                    f"Reject the null hypothesis (p {self.format_p_value(self.results['p_value'])} < α = {self.results['alpha']})."
                )
            else:
                doc.add_paragraph(
                    f"Fail to reject the null hypothesis (p {self.format_p_value(self.results['p_value'])} ≥ α = {self.results['alpha']})."
                )

            doc.add_heading('Interpretation', 2)
            doc.add_paragraph(self.generate_interpretation())

            if self.results.get('warning'):
                warn_p.runs[0].font.color.rgb = RGBColor(180, 0, 0)

            # Add saved file info with italic formatting
            saved_p = doc.add_paragraph(f"\nSaved file: {fp}")
            saved_p.runs[0].italic = True
            saved_p.runs[0].font.size = Pt(8)
            
            date_p = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_p.runs[0].italic = True
            date_p.runs[0].font.size = Pt(8)

            doc.save(fp)
            messagebox.showinfo("Success", f"Word document exported:\n{fp}")

        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")


# ===============================
# Application Entry Point
# ===============================
if __name__ == "__main__":
    ctk.set_appearance_mode("System")   # Light / Dark / System
    ctk.set_default_color_theme("blue") # Theme

    app = ChiSquareTestApp()
    app.mainloop()