"""
ttest_session_manager.py  —  Multi-Part Session Manager for T-Test Analysis
Handles saving / loading / exporting of multiple t-test parts
(Part 1-A, 1-B, 2-A …) with DOCX and PDF batch export.
"""

import re
import customtkinter as ctk
from tkinter import messagebox
from datetime import datetime

# ── Palette ───────────────────────────────────────────────────────────────────
BG_DEEP   = "#0d1117"
BG_CARD   = "#161b22"
BG_PANEL  = "#1c2230"
BG_INPUT  = "#1e2736"
ACCENT    = "#00c9a7"
ACCENT2   = "#4e9eff"
DANGER    = "#ef4444"
WARN      = "#f59e0b"
SUCCESS   = "#22c55e"
PURPLE    = "#a855f7"
TEXT_PRI  = "#e6edf3"
TEXT_SEC  = "#8b949e"
BORDER    = "#30363d"

FONT_HEAD = ("Segoe UI", 14, "bold")
FONT_CARD = ("Segoe UI", 12, "bold")
FONT_BODY = ("Segoe UI", 11)
FONT_BTN  = ("Segoe UI", 11, "bold")
FONT_MONO = ("Consolas", 9)
FONT_TINY = ("Segoe UI", 9)


# ── Label generator ───────────────────────────────────────────────────────────

def _next_part_label(existing_labels: list) -> str:
    """Generate next Part label: Part 1-A → Part 1-B → … → Part 2-A"""
    if not existing_labels:
        return "Part 1-A"
    max_num, max_letter = 1, "A"
    for lbl in existing_labels:
        m = re.match(r"Part (\d+)-([A-Z])", str(lbl))
        if m:
            n, l = int(m.group(1)), m.group(2)
            if n > max_num or (n == max_num and l > max_letter):
                max_num, max_letter = n, l
    next_letter = chr(ord(max_letter) + 1)
    if next_letter > "Z":
        return f"Part {max_num + 1}-A"
    return f"Part {max_num}-{next_letter}"


def _fmt(v, d=4):
    """Format a numeric value to d decimal places, or return 'N/A'."""
    try:
        return f"{float(v):.{d}f}"
    except (TypeError, ValueError):
        return "N/A"


def _fmt2(v):
    return _fmt(v, 2)


def _fmt_p(p):
    try:
        p = float(p)
        return "< .001" if p < 0.001 else f"= {p:.3f}"
    except (TypeError, ValueError):
        return "N/A"


# ── Part Card ─────────────────────────────────────────────────────────────────

class PartCard(ctk.CTkFrame):
    """Single saved-part row in the session panel."""

    def __init__(self, master, part_data: dict, index: int,
                 on_delete, on_view, on_load, **kw):
        super().__init__(master, fg_color=BG_PANEL, corner_radius=8,
                         border_width=1, border_color=BORDER, **kw)
        self.part_data = part_data
        self._build(index, on_delete, on_view, on_load)

    def _build(self, index, on_delete, on_view, on_load):
        top = ctk.CTkFrame(self, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=(8, 2))

        # Coloured part-label badge
        badge_color = [ACCENT, ACCENT2, WARN, PURPLE, SUCCESS][index % 5]
        ctk.CTkLabel(top, text=self.part_data.get("label", f"Part {index+1}"),
                     font=("Segoe UI", 10, "bold"),
                     fg_color=badge_color, text_color="#0d1117",
                     corner_radius=5, padx=8, pady=2).pack(side="left")

        # Test type badge
        test_type = self.part_data.get("test_type", "t-test")
        type_labels = {
            "one-sample":  "One-Sample",
            "independent": "Independent",
            "paired":      "Paired",
        }
        type_text = type_labels.get(test_type, test_type.title())
        ctk.CTkLabel(top, text=type_text,
                     font=FONT_TINY, fg_color=BG_INPUT, text_color=ACCENT2,
                     corner_radius=4, padx=6, pady=2).pack(side="left", padx=4)

        # Significant / Not Significant badge
        is_sig    = self.part_data.get("is_significant", False)
        sig_text  = "✓ Significant" if is_sig else "✗ Not Significant"
        sig_color = SUCCESS if is_sig else DANGER
        ctk.CTkLabel(top, text=sig_text, font=FONT_TINY,
                     fg_color=sig_color, text_color="#fff",
                     corner_radius=4, padx=6, pady=2).pack(side="left", padx=4)

        # Action buttons (right side)
        ctk.CTkButton(top, text="✕", width=24, height=24,
                      fg_color="transparent", hover_color=DANGER,
                      text_color=TEXT_SEC, corner_radius=4,
                      font=("Segoe UI", 10, "bold"),
                      command=lambda: on_delete(self.part_data["label"])
                      ).pack(side="right")

        ctk.CTkButton(top, text="↩ Load", width=52, height=24,
                      fg_color=BG_INPUT, hover_color=ACCENT,
                      text_color=TEXT_PRI, corner_radius=4,
                      font=("Segoe UI", 10),
                      command=lambda: on_load(self.part_data)
                      ).pack(side="right", padx=4)

        ctk.CTkButton(top, text="👁 View", width=52, height=24,
                      fg_color=BG_INPUT, hover_color=ACCENT2,
                      text_color=TEXT_PRI, corner_radius=4,
                      font=("Segoe UI", 10),
                      command=lambda: on_view(self.part_data)
                      ).pack(side="right", padx=4)

        # Key stats row
        r  = self.part_data
        t  = _fmt(r.get("t_statistic"), 2)
        df = _fmt(r.get("df"), 2)
        p  = _fmt(r.get("p_value"), 3)
        d  = _fmt(r.get("cohens_d"), 2)
        stats_str = (f"t({df}) = {t}   p {_fmt_p(r.get('p_value'))}   "
                     f"d = {d}   {r.get('decision', '')}")
        ctk.CTkLabel(self, text=stats_str, font=FONT_TINY,
                     text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 2))

        # Title / subtitle line
        sub = r.get("subtitle") or r.get("title", "")
        if sub:
            ctk.CTkLabel(self, text=sub, font=("Segoe UI", 9, "italic"),
                         text_color=TEXT_SEC).pack(anchor="w", padx=10, pady=(0, 6))


# ── Session Manager Panel ─────────────────────────────────────────────────────

class SessionManagerPanel(ctk.CTkToplevel):
    """Floating panel listing all saved t-test parts."""

    def __init__(self, master, saved_parts: list,
                 on_delete_part, on_load_part,
                 on_export_all_docx, on_export_all_pdf):
        super().__init__(master)
        self.title("📁  Session Manager — Saved t-Test Parts")
        self.geometry("560x680")
        self.minsize(440, 480)
        self.configure(fg_color=BG_DEEP)
        self.resizable(True, True)

        self.saved_parts      = saved_parts
        self.on_delete_part   = on_delete_part
        self.on_load_part     = on_load_part
        self.on_export_docx   = on_export_all_docx
        self.on_export_pdf    = on_export_all_pdf

        self.transient(master)
        self.lift()
        self.focus_force()

        self._build()

    def _build(self):
        # Header
        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=52)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="📁  Session Manager — t-Test Parts",
                     font=FONT_HEAD, text_color=TEXT_PRI).pack(side="left", padx=16)
        self.count_label = ctk.CTkLabel(
            hdr, text=f"{len(self.saved_parts)} part(s) saved",
            font=FONT_TINY, text_color=TEXT_SEC)
        self.count_label.pack(side="right", padx=16)

        # Scrollable list
        self.scroll = ctk.CTkScrollableFrame(self, fg_color=BG_DEEP,
                                             scrollbar_button_color=BORDER)
        self.scroll.pack(fill="both", expand=True, padx=12, pady=12)
        self._populate_cards()

        # Bottom button bar
        btn_frame = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=68)
        btn_frame.pack(fill="x")
        btn_frame.pack_propagate(False)

        ctk.CTkButton(btn_frame, text="💾  Export All → DOCX",
                      fg_color="#1d4ed8", hover_color="#1e3a8a",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_docx
                      ).pack(side="left", padx=(14, 6), pady=16)

        ctk.CTkButton(btn_frame, text="📄  Export All → PDF",
                      fg_color=DANGER, hover_color="#b91c1c",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.on_export_pdf
                      ).pack(side="left", padx=6, pady=16)

        ctk.CTkButton(btn_frame, text="🗑  Clear All",
                      fg_color="#7f1d1d", hover_color="#450a0a",
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self._clear_all
                      ).pack(side="left", padx=6, pady=16)

        ctk.CTkButton(btn_frame, text="✕ Close",
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=36, corner_radius=8,
                      command=self.destroy
                      ).pack(side="right", padx=14, pady=16)

    def _populate_cards(self):
        for w in self.scroll.winfo_children():
            w.destroy()
        if not self.saved_parts:
            ctk.CTkLabel(self.scroll,
                         text="No parts saved yet.\nRun a t-test and click  '🗂 Save Part'.",
                         font=FONT_BODY, text_color=TEXT_SEC,
                         justify="center").pack(expand=True, pady=40)
            return
        for i, part in enumerate(self.saved_parts):
            c = PartCard(self.scroll, part, i,
                         on_delete=self._handle_delete,
                         on_view=self._handle_view,
                         on_load=self._handle_load)
            c.pack(fill="x", pady=4)

    def _handle_delete(self, label: str):
        if messagebox.askyesno("Delete Part",
                               f"Remove '{label}' from this session?"):
            self.on_delete_part(label)
            self.refresh()

    def _handle_view(self, part_data: dict):
        ViewPartWindow(self, part_data)

    def _handle_load(self, part_data: dict):
        self.on_load_part(part_data)
        self.destroy()

    def _clear_all(self):
        if not self.saved_parts:
            messagebox.showinfo("Nothing to Clear", "There are no saved parts to clear.")
            return
        n = len(self.saved_parts)
        if messagebox.askyesno("Clear All Parts",
                               f"Remove all {n} saved part(s)?\n\nThis cannot be undone."):
            for part in list(self.saved_parts):
                self.on_delete_part(part["label"])
            self.refresh()

    def refresh(self):
        self._populate_cards()
        self.count_label.configure(text=f"{len(self.saved_parts)} part(s) saved")


# ── View Part Window ──────────────────────────────────────────────────────────

class ViewPartWindow(ctk.CTkToplevel):
    def __init__(self, master, part_data: dict):
        super().__init__(master)
        self.title(f"View — {part_data.get('label', 'Part')}")
        self.geometry("580x540")
        self.configure(fg_color=BG_DEEP)

        hdr = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0, height=48)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr,
                     text=f"  {part_data.get('label', '')}  —  {part_data.get('title', '')}",
                     font=FONT_CARD, text_color=TEXT_PRI).pack(side="left", padx=12)

        txt = ctk.CTkTextbox(self, fg_color=BG_INPUT, text_color=TEXT_PRI,
                             font=FONT_MONO, corner_radius=0)
        txt.pack(fill="both", expand=True, padx=12, pady=12)
        txt.insert("1.0", _build_part_summary(part_data))
        txt.configure(state="disabled")

        ctk.CTkButton(self, text="Close", command=self.destroy,
                      fg_color=BG_PANEL, hover_color=BORDER,
                      font=FONT_BTN, height=32).pack(pady=(0, 12))


# ── Text summary builder ──────────────────────────────────────────────────────

def _build_part_summary(r: dict) -> str:
    line = "─" * 52
    out  = f"{r.get('label', '')}\n{line}\n"
    out += f"Test: {r.get('test_name', r.get('test_type', ''))}\n"
    if r.get("title"):      out += f"Title: {r['title']}\n"
    if r.get("subtitle"):   out += f"Subtitle: {r['subtitle']}\n"
    if r.get("researcher"): out += f"by: {r['researcher']}\n"
    out += f"α = {r.get('alpha', 0.05)}\n"
    out += f"{line}\n\n"

    test_type = r.get("test_type", "")

    out += "DESCRIPTIVE STATISTICS\n" + "─" * 30 + "\n"
    if test_type == "one-sample":
        out += (f"  {r.get('group1_name', 'Sample')}:  "
                f"M = {_fmt2(r.get('mean1'))},  "
                f"SD = {_fmt2(r.get('std1'))},  "
                f"N = {r.get('n1', '')}\n")
        out += f"  Test Value (μ₀) = {r.get('test_value', 0)}\n"
    elif test_type == "independent":
        out += (f"  {r.get('group1_name', 'Group 1')}:  "
                f"M = {_fmt2(r.get('mean1'))},  "
                f"SD = {_fmt2(r.get('std1'))},  "
                f"N = {r.get('n1', '')}\n")
        out += (f"  {r.get('group2_name', 'Group 2')}:  "
                f"M = {_fmt2(r.get('mean2'))},  "
                f"SD = {_fmt2(r.get('std2'))},  "
                f"N = {r.get('n2', '')}\n")
    elif test_type == "paired":
        out += (f"  {r.get('group1_name', 'Pre')}:  "
                f"M = {_fmt2(r.get('mean1'))},  "
                f"N = {r.get('n1', '')}\n")
        out += (f"  {r.get('group2_name', 'Post')}:  "
                f"M = {_fmt2(r.get('mean2'))},  "
                f"N = {r.get('n2', r.get('n1', ''))}\n")
        out += (f"  Mean Difference = {_fmt2(r.get('mean_diff'))},  "
                f"SD(diff) = {_fmt2(r.get('std_diff'))}\n")

    out += f"\n{'═'*52}\nTEST RESULTS\n{'═'*52}\n"
    out += f"t({_fmt(r.get('df'), 2)}) = {_fmt(r.get('t_statistic'), 4)}\n"
    out += f"p {_fmt_p(r.get('p_value'))}\n"
    out += f"Cohen's d = {_fmt(r.get('cohens_d'), 4)}\n\n"
    out += f"Decision:   {r.get('decision', '')}\n\n"
    out += f"Interpretation:\n{r.get('interpretation', '')}\n\n"
    out += f"Saved: {r.get('saved_at', '—')}\n"
    return out


# ── DOCX Export (all parts) ───────────────────────────────────────────────────

def export_all_to_docx(parts: list, filepath: str):
    """Write all saved t-test parts into a single APA-style DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_margins(sect):
        sect.top_margin    = Inches(0.75)
        sect.bottom_margin = Inches(0.75)
        sect.left_margin   = Inches(1.0)
        sect.right_margin  = Inches(1.0)

    def _page_break(doc):
        doc.add_page_break()

    def apa_borders(table):
        tbl   = table._tbl
        tblPr = tbl.tblPr
        tb    = OxmlElement("w:tblBorders")
        for bn in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "none")
            tb.append(b)
        for bn, sz in [("top", "12"), ("bottom", "12")]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), sz)
            tb.append(b)
        tblPr.append(tb)

    def hdr_sep(table):
        for cell in table.rows[0].cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            tcB.append(bot)
            tcPr.append(tcB)

    def cf(cell, text, bold=False, italic=False, size=10, align="left"):
        para = cell.paragraphs[0]
        para.alignment = (WD_PARAGRAPH_ALIGNMENT.CENTER
                          if align == "center" else WD_PARAGRAPH_ALIGNMENT.LEFT)
        run = para.add_run(str(text))
        run.font.size = Pt(size)
        run.bold      = bold
        run.italic    = italic

    doc = Document()
    _set_margins(doc.sections[0])

    for part_idx, r in enumerate(parts):
        if part_idx > 0:
            _page_break(doc)

        test_type = r.get("test_type", "")

        # ── Part label ────────────────────────────────────────────────────────
        lp = doc.add_heading(r.get("label", f"Part {part_idx + 1}"), level=1)
        lp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in lp.runs:
            run.font.size      = Pt(13)
            run.bold           = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # ── Title / subtitle / author ─────────────────────────────────────────
        if r.get("title"):
            tp = doc.add_paragraph(r["title"])
            tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in tp.runs:
                run.bold          = True
                run.font.size     = Pt(14)

        if r.get("subtitle"):
            sp = doc.add_paragraph(r["subtitle"])
            sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sp.runs:
                run.italic    = True
                run.font.size = Pt(11)

        if r.get("researcher"):
            ap = doc.add_paragraph(f"by: {r['researcher']}")
            ap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in ap.runs:
                run.italic    = True
                run.font.size = Pt(10)

        # ── Test heading ──────────────────────────────────────────────────────
        doc.add_heading(r.get("test_name", "t-Test"), 2)
        doc.add_paragraph(f"Alpha: α = {r.get('alpha', 0.05)}")
        doc.add_paragraph()

        # ── Table 1: Descriptive Statistics ───────────────────────────────────
        doc.add_heading("Descriptive Statistics", 3)

        if test_type == "one-sample":
            dt = doc.add_table(rows=2, cols=4)
            apa_borders(dt); hdr_sep(dt)
            for i, h in enumerate(["Group", "M", "SD", "N"]):
                cf(dt.rows[0].cells[i], h, bold=True, align="center")
            cf(dt.rows[1].cells[0], r.get("group1_name", "Sample"))
            cf(dt.rows[1].cells[1], _fmt2(r.get("mean1")),       align="center")
            cf(dt.rows[1].cells[2], _fmt2(r.get("std1")),        align="center")
            cf(dt.rows[1].cells[3], str(r.get("n1", "")),        align="center")
            doc.add_paragraph(f"Test Value (μ₀) = {r.get('test_value', 0)}")

        elif test_type == "independent":
            dt = doc.add_table(rows=3, cols=4)
            apa_borders(dt); hdr_sep(dt)
            for i, h in enumerate(["Group", "M", "SD", "N"]):
                cf(dt.rows[0].cells[i], h, bold=True, align="center")
            cf(dt.rows[1].cells[0], r.get("group1_name", "Group 1"))
            cf(dt.rows[1].cells[1], _fmt2(r.get("mean1")), align="center")
            cf(dt.rows[1].cells[2], _fmt2(r.get("std1")),  align="center")
            cf(dt.rows[1].cells[3], str(r.get("n1", "")),  align="center")
            cf(dt.rows[2].cells[0], r.get("group2_name", "Group 2"))
            cf(dt.rows[2].cells[1], _fmt2(r.get("mean2")), align="center")
            cf(dt.rows[2].cells[2], _fmt2(r.get("std2")),  align="center")
            cf(dt.rows[2].cells[3], str(r.get("n2", "")),  align="center")

        else:  # paired
            dt = doc.add_table(rows=4, cols=3)
            apa_borders(dt); hdr_sep(dt)
            for i, h in enumerate(["Measurement", "M", "N"]):
                cf(dt.rows[0].cells[i], h, bold=True, align="center")
            cf(dt.rows[1].cells[0], r.get("group1_name", "Pre"))
            cf(dt.rows[1].cells[1], _fmt2(r.get("mean1")), align="center")
            cf(dt.rows[1].cells[2], str(r.get("n1", "")),  align="center")
            cf(dt.rows[2].cells[0], r.get("group2_name", "Post"))
            cf(dt.rows[2].cells[1], _fmt2(r.get("mean2")), align="center")
            cf(dt.rows[2].cells[2], str(r.get("n2", r.get("n1", ""))), align="center")
            cf(dt.rows[3].cells[0], "Mean Difference")
            cf(dt.rows[3].cells[1], _fmt2(r.get("mean_diff")), align="center")
            cf(dt.rows[3].cells[2],
               f"SD(diff) = {_fmt2(r.get('std_diff'))}", align="center")

        doc.add_paragraph()

        # ── Table 2: Test Statistics ───────────────────────────────────────────
        doc.add_heading("Test Statistics", 3)
        st = doc.add_table(rows=5, cols=2)
        apa_borders(st); hdr_sep(st)
        for i, h in enumerate(["Statistic", "Value"]):
            cf(st.rows[0].cells[i], h, bold=True, align="center")
        for row_i, (stat, val) in enumerate([
            ("t",         _fmt(r.get("t_statistic"), 4)),
            ("df",        _fmt(r.get("df"), 2)),
            ("p",         _fmt_p(r.get("p_value"))),
            ("Cohen's d", _fmt(r.get("cohens_d"), 4)),
        ], 1):
            cf(st.rows[row_i].cells[0], stat)
            cf(st.rows[row_i].cells[1], val, align="center")

        doc.add_paragraph()

        # ── Decision & Interpretation ─────────────────────────────────────────
        doc.add_heading("Decision", 3)
        doc.add_paragraph(r.get("decision", ""))

        doc.add_heading("Interpretation", 3)
        doc.add_paragraph(r.get("interpretation", ""))

        doc.add_paragraph()

        # ── Raw data (abbreviated) ────────────────────────────────────────────
        if r.get("test_type") in ("independent", "paired"):
            doc.add_heading("Raw Data (abbreviated)", 3)
            raw_tbl = doc.add_table(rows=3, cols=3)
            apa_borders(raw_tbl); hdr_sep(raw_tbl)
            for i, h in enumerate(["Group", "N", "Values (first 10)"]):
                cf(raw_tbl.rows[0].cells[i], h, bold=True, align="center", size=9)
            for row_i, (grp, data_key, n_key) in enumerate([
                (r.get("group1_name", "Group 1"), "data1", "n1"),
                (r.get("group2_name", "Group 2"), "data2", "n2"),
            ], 1):
                data = r.get(data_key, [])
                preview = ", ".join(_fmt2(v) for v in data[:10])
                if len(data) > 10:
                    preview += " …"
                cf(raw_tbl.rows[row_i].cells[0], grp,             size=9)
                cf(raw_tbl.rows[row_i].cells[1], str(r.get(n_key, len(data))),
                   align="center", size=9)
                cf(raw_tbl.rows[row_i].cells[2], preview,         size=9)
            doc.add_paragraph()

        # ── Footer ────────────────────────────────────────────────────────────
        fp_p = doc.add_paragraph(
            f"Part saved: {r.get('saved_at', '')}   |   "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}")
        for run in fp_p.runs:
            run.font.size      = Pt(7)
            run.italic         = True
            run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(filepath)


# ── PDF Export (all parts) ────────────────────────────────────────────────────

def export_all_to_pdf(parts: list, filepath: str):
    """Write all saved t-test parts into a single APA-style PDF."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, Table, TableStyle,
        HRFlowable, NextPageTemplate, PageBreak,
    )
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    PAGE_W, PAGE_H = letter
    ML = MR = 0.75 * inch
    MT = MB = 0.60 * inch
    COL_W = PAGE_W - ML - MR
    COL_H = PAGE_H - MT - MB

    ACCENT_CL = colors.HexColor("#00796b")
    GREY_TEXT  = colors.HexColor("#555555")
    BLACK      = colors.black

    styles = getSampleStyleSheet()

    def PS(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    PartLbl  = PS("PartLbl",  fontSize=13, fontName="Helvetica-Bold",
                              spaceAfter=2, textColor=colors.HexColor("#003366"))
    TitleSt  = PS("TitleSt",  fontSize=12, fontName="Helvetica-Bold",
                              alignment=TA_CENTER, spaceAfter=2)
    SubSt    = PS("SubSt",    fontSize=10, fontName="Helvetica-Oblique",
                              alignment=TA_CENTER, spaceAfter=1)
    SecHdr   = PS("SecHdr",   fontSize=10, fontName="Helvetica-Bold",
                              spaceAfter=2, spaceBefore=6,
                              textColor=colors.HexColor("#1a1a2e"))
    BodySt   = PS("BodySt",   fontSize=9,  spaceAfter=2,  leading=13)
    ResultSt = PS("ResultSt", fontSize=11, spaceAfter=3,  leading=15)
    SmallSt  = PS("SmallSt",  fontSize=7.5, textColor=GREY_TEXT,
                              spaceAfter=0, leading=9)

    def tbl_style_fn(extra=None):
        base = [
            ("FONTNAME",      (0,  0), (-1,  0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,  0), (-1, -1),  9),
            ("LEADING",       (0,  0), (-1, -1),  12),
            ("ALIGN",         (1,  0), (-1, -1),  "CENTER"),
            ("ALIGN",         (0,  0), (0,  -1),  "LEFT"),
            ("LINEABOVE",     (0,  0), (-1,  0),  1.2, BLACK),
            ("LINEBELOW",     (0,  0), (-1,  0),  0.6, BLACK),
            ("LINEBELOW",     (0, -1), (-1, -1),  1.2, BLACK),
            ("TOPPADDING",    (0,  0), (-1, -1),  3),
            ("BOTTOMPADDING", (0,  0), (-1, -1),  3),
            ("LEFTPADDING",   (0,  0), (-1, -1),  5),
            ("RIGHTPADDING",  (0,  0), (-1, -1),  5),
        ]
        if extra:
            base.extend(extra)
        return TableStyle(base)

    # One PageTemplate per part
    doc = BaseDocTemplate(filepath, pagesize=letter,
                          leftMargin=ML, rightMargin=MR,
                          topMargin=MT,  bottomMargin=MB)
    templates = []
    for i in range(len(parts)):
        frame = Frame(ML, MB, COL_W, COL_H,
                      leftPadding=0, rightPadding=0,
                      topPadding=0,  bottomPadding=0,
                      id=f"p{i}_main")
        templates.append(PageTemplate(id=f"part{i}", frames=[frame]))
    doc.addPageTemplates(templates)

    story = []

    for part_idx, r in enumerate(parts):
        test_type = r.get("test_type", "")

        story.append(NextPageTemplate(f"part{part_idx}"))
        if part_idx > 0:
            story.append(PageBreak())

        # Part label + rule
        story.append(Paragraph(r.get("label", f"Part {part_idx + 1}"), PartLbl))
        story.append(HRFlowable(width=COL_W, thickness=1.5,
                                color=ACCENT_CL, spaceAfter=4))

        # Title / subtitle / researcher
        story.append(Paragraph(r.get("title", "t-Test Analysis Results"), TitleSt))
        if r.get("subtitle"):
            story.append(Paragraph(r["subtitle"], SubSt))
        if r.get("researcher"):
            story.append(Paragraph(r["researcher"], SubSt))
        story.append(Spacer(1, 4))

        # Test name & alpha
        story.append(Paragraph(r.get("test_name", "t-Test"), SecHdr))
        story.append(Paragraph(f"α = {r.get('alpha', 0.05)}", BodySt))
        story.append(Spacer(1, 4))

        # ── Descriptive Statistics table ───────────────────────────────────────
        story.append(Paragraph("Descriptive Statistics", SecHdr))

        if test_type == "one-sample":
            t1_data = [["Group", "M", "SD", "N"],
                       [r.get("group1_name", "Sample"),
                        _fmt2(r.get("mean1")),
                        _fmt2(r.get("std1")),
                        str(r.get("n1", ""))]]
            story.append(Paragraph(f"Test Value (μ₀) = {r.get('test_value', 0)}", BodySt))
            cw = [COL_W*0.40, COL_W*0.20, COL_W*0.20, COL_W*0.20]

        elif test_type == "independent":
            t1_data = [["Group", "M", "SD", "N"],
                       [r.get("group1_name", "Group 1"),
                        _fmt2(r.get("mean1")), _fmt2(r.get("std1")), str(r.get("n1", ""))],
                       [r.get("group2_name", "Group 2"),
                        _fmt2(r.get("mean2")), _fmt2(r.get("std2")), str(r.get("n2", ""))]]
            cw = [COL_W*0.40, COL_W*0.20, COL_W*0.20, COL_W*0.20]

        else:  # paired
            t1_data = [["Measurement", "M", "N"],
                       [r.get("group1_name", "Pre"),
                        _fmt2(r.get("mean1")), str(r.get("n1", ""))],
                       [r.get("group2_name", "Post"),
                        _fmt2(r.get("mean2")), str(r.get("n2", r.get("n1", "")))],
                       ["Mean Difference",
                        _fmt2(r.get("mean_diff")),
                        f"SD(diff)={_fmt2(r.get('std_diff'))}"]]
            cw = [COL_W*0.45, COL_W*0.25, COL_W*0.30]

        t1 = Table(t1_data, colWidths=cw)
        t1.setStyle(tbl_style_fn())
        story.append(t1)
        story.append(Spacer(1, 6))

        # ── Test Statistics table ──────────────────────────────────────────────
        story.append(Paragraph("Test Statistics", SecHdr))
        t2_data = [
            ["Statistic", "Value"],
            ["t",         _fmt(r.get("t_statistic"), 4)],
            ["df",        _fmt(r.get("df"), 2)],
            ["p",         _fmt_p(r.get("p_value"))],
            ["Cohen's d", _fmt(r.get("cohens_d"), 4)],
        ]
        cw2 = [COL_W * 0.50, COL_W * 0.50]
        t2 = Table(t2_data, colWidths=cw2)
        t2.setStyle(tbl_style_fn())
        story.append(t2)
        story.append(Spacer(1, 6))

        # ── Decision & result line ─────────────────────────────────────────────
        story.append(Paragraph("Decision", SecHdr))
        story.append(Paragraph(
            f'<b>t</b>({_fmt(r.get("df"), 2)}) = '
            f'<b>{_fmt(r.get("t_statistic"), 4)}</b>,&nbsp;&nbsp;'
            f'p {_fmt_p(r.get("p_value"))},&nbsp;&nbsp;'
            f"d = {_fmt(r.get('cohens_d'), 2)}",
            ResultSt))
        story.append(Paragraph(
            f'<b>Decision:&nbsp;</b><b>{r.get("decision", "")}</b>',
            ResultSt))
        story.append(Spacer(1, 4))

        # ── Interpretation ────────────────────────────────────────────────────
        story.append(Paragraph("Interpretation", SecHdr))
        story.append(Paragraph(r.get("interpretation", ""), BodySt))
        story.append(Spacer(1, 6))

        # ── Footer ─────────────────────────────────────────────────────────────
        story.append(HRFlowable(width=COL_W, thickness=0.5,
                                color=colors.grey, spaceAfter=2))
        story.append(Paragraph(
            f"Part saved: {r.get('saved_at', '—')}  &nbsp;|&nbsp;  "
            f"Generated: {datetime.now().strftime('%Y-%m-%d %I:%M %p')}  "
            f"&nbsp;|&nbsp;  Part {part_idx + 1} of {len(parts)}",
            SmallSt))

    doc.build(story)