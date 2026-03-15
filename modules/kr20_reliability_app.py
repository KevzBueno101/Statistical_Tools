"""
Kuder-Richardson Formula 20 (KR-20) Reliability Test Application
=================================================================
A desktop statistical application for computing KR-20 reliability
from binary (dichotomous) test data and generating professional
PDF and DOCX reports for academic submission.

Formula:
    KR-20 = (k / (k - 1)) * (1 - (Σ(p*q) / σ²_total))

    k          = number of items
    p          = proportion correct for each item
    q          = 1 - p
    Σ(p*q)     = sum of item variances (dichotomous)
    σ²_total   = variance of the total test scores

Requirements:
    pip install customtkinter pandas numpy reportlab python-docx openpyxl
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import pandas as pd
import numpy as np
from datetime import datetime
import os

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# python-docx import for DOCX generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# Binary Frequency Expander Module
# ============================================================================

class BinaryFrequencyExpander:
    """
    Module for expanding binary frequency data into respondent-level data.
    Converts aggregated correct/incorrect counts into individual 0/1 records.

    Example:
        Item1: Correct=40, Incorrect=10
        → Expands to: [1, 1, 1, ...(40 times)..., 0, 0, ...(10 times)]
    """

    @staticmethod
    def expand_item(correct: int, incorrect: int) -> list:
        """
        Expand a single binary item from frequency counts to respondent values.

        Parameters
        ----------
        correct   : int  – number of respondents who answered correctly (1)
        incorrect : int  – number of respondents who answered incorrectly (0)

        Returns
        -------
        list of int  – individual respondent responses [1, 1, ..., 0, 0, ...]
        """
        if correct < 0 or incorrect < 0:
            raise ValueError("Frequencies must be non-negative integers.")
        return [1] * correct + [0] * incorrect

    @staticmethod
    def expand_multiple_items(items_freq_data: dict) -> pd.DataFrame:
        """
        Expand multiple binary items into a respondent-level DataFrame.

        Parameters
        ----------
        items_freq_data : dict
            {item_name: {"correct": int, "incorrect": int}, ...}

        Returns
        -------
        pd.DataFrame  – rows = respondents, columns = items (values 0 or 1)
        """
        expanded_data = {}
        respondent_counts = {}

        for item_name, freq in items_freq_data.items():
            expanded = BinaryFrequencyExpander.expand_item(
                freq["correct"], freq["incorrect"]
            )
            expanded_data[item_name] = expanded
            respondent_counts[item_name] = len(expanded)

        # All items must share the same total respondent count
        unique_counts = set(respondent_counts.values())
        if len(unique_counts) > 1:
            raise ValueError(
                f"Inconsistent respondent counts: {respondent_counts}. "
                "All items must have the same total (correct + incorrect)."
            )

        return pd.DataFrame(expanded_data)

    @staticmethod
    def validate_item_input(correct: int, incorrect: int):
        """
        Validate a single item's frequency inputs.

        Returns
        -------
        (bool, str|None)  – (is_valid, error_message)
        """
        if correct < 0 or incorrect < 0:
            return False, "Frequencies must be non-negative integers."
        total = correct + incorrect
        if total < 2:
            return False, f"Total respondents ({total}) must be at least 2."
        return True, None


# ============================================================================
# KR-20 Statistical Calculator
# ============================================================================

class KR20Calculator:
    """
    Statistical computation module for the Kuder-Richardson Formula 20 (KR-20).

    KR-20 is the appropriate reliability measure for tests composed
    entirely of dichotomously scored items (0 = incorrect, 1 = correct).
    """

    @staticmethod
    def validate_binary_data(data: pd.DataFrame):
        """
        Ensure all values in the DataFrame are binary (0 or 1) or NaN.

        Parameters
        ----------
        data : pd.DataFrame

        Returns
        -------
        (bool, str|None)  – (is_valid, error_message)
        """
        numeric = data.select_dtypes(include=[np.number])
        if numeric.empty:
            return False, "No numeric columns found in the dataset."

        # Drop NaN for value checking
        flat_values = numeric.values.flatten()
        non_nan = flat_values[~np.isnan(flat_values.astype(float))]
        unique_vals = set(non_nan)

        invalid = unique_vals - {0, 1, 0.0, 1.0}
        if invalid:
            return (
                False,
                f"Non-binary values detected: {sorted(invalid)}. "
                "KR-20 requires only 0 and 1 values."
            )
        return True, None

    @staticmethod
    def compute_kr20(data: pd.DataFrame) -> dict:
        """
        Compute the KR-20 reliability coefficient.

        Formula:
            KR-20 = (k / (k-1)) * (1 - Σ(p*q) / σ²_total)

        Parameters
        ----------
        data : pd.DataFrame
            Respondent-level binary data (rows=respondents, cols=items)

        Returns
        -------
        dict with keys:
            kr20, n_items, n_respondents, sum_pq, total_variance,
            item_stats, interpretation
        """
        # Work with numeric columns only
        numeric_data = data.select_dtypes(include=[np.number])
        if numeric_data.empty:
            raise ValueError("No numeric columns found in the dataset.")

        # Validate binary data
        is_valid, error_msg = KR20Calculator.validate_binary_data(numeric_data)
        if not is_valid:
            raise ValueError(error_msg)

        # Drop rows that are entirely NaN
        clean_data = numeric_data.dropna(how="all")
        if len(clean_data) < 2:
            raise ValueError("Need at least 2 valid respondents after removing missing rows.")

        k = clean_data.shape[1]   # number of items
        n = clean_data.shape[0]   # number of respondents

        if k < 2:
            raise ValueError("At least 2 items are required to compute KR-20.")

        # Per-item p and q (using pairwise complete cases per item)
        item_stats = {}
        sum_pq = 0.0

        for col in clean_data.columns:
            col_data = clean_data[col].dropna()
            p = col_data.mean()       # proportion correct
            q = 1.0 - p               # proportion incorrect
            pq = p * q
            sum_pq += pq
            item_stats[col] = {
                "p": round(p, 4),
                "q": round(q, 4),
                "pq": round(pq, 4),
                "n_valid": len(col_data)
            }

        # Total test score per respondent
        total_scores = clean_data.sum(axis=1)

        # Population variance of total scores  (ddof=0 matches classical test theory)
        total_variance = float(total_scores.var(ddof=1))

        if total_variance == 0:
            raise ValueError(
                "Total score variance is zero. "
                "All respondents have identical total scores – KR-20 is undefined."
            )

        # KR-20 formula
        kr20 = (k / (k - 1)) * (1.0 - sum_pq / total_variance)

        # Clamp to [-1, 1] (negative values can occur with poor items)
        kr20 = max(-1.0, min(1.0, kr20))

        return {
            "kr20": kr20,
            "n_items": k,
            "n_respondents": n,
            "sum_pq": sum_pq,
            "total_variance": total_variance,
            "item_stats": item_stats,
            "interpretation": KR20Calculator.interpret_kr20(kr20),
        }

    @staticmethod
    def interpret_kr20(kr20: float) -> str:
        """Return a verbal interpretation of the KR-20 value."""
        if kr20 >= 0.90:
            return "Excellent"
        elif kr20 >= 0.80:
            return "Good"
        elif kr20 >= 0.70:
            return "Acceptable"
        elif kr20 >= 0.60:
            return "Questionable"
        elif kr20 >= 0.50:
            return "Poor"
        else:
            return "Unacceptable"


# ============================================================================
# PDF Report Generator
# ============================================================================

class PDFReportGenerator:
    """
    Generates a professional APA-style PDF report for KR-20 analysis
    using ReportLab.
    """

    @staticmethod
    def generate_report(results: dict, description: str, filename: str,
                        title: str = None, subtitle: str = None, byline: str = None):
        """
        Build and save the PDF report.

        Parameters
        ----------
        results     : dict   – output from KR20Calculator.compute_kr20()
        description : str    – free-text study description
        filename    : str    – output file path
        title       : str    – report title
        subtitle    : str    – report subtitle
        byline      : str    – author name
        """
        doc = SimpleDocTemplate(
            filename, pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=36
        )

        elements = []
        styles = getSampleStyleSheet()

        # ── Custom Paragraph Styles ──────────────────────────────────────────
        title_style = ParagraphStyle(
            "KRTitle", parent=styles["Heading1"],
            fontSize=16, textColor=colors.black,
            spaceAfter=10, alignment=TA_CENTER,
            fontName="Helvetica-Bold"
        )
        byline_style = ParagraphStyle(
            "KRByline", parent=styles["Normal"],
            fontSize=12, textColor=colors.black,
            alignment=TA_CENTER, spaceAfter=18,
            fontName="Helvetica"
        )
        subtitle_style = ParagraphStyle(
            "KRSubtitle", parent=styles["Normal"],
            fontSize=12, textColor=colors.Color(0.27, 0.27, 0.27),
            alignment=TA_CENTER, spaceAfter=6,
            fontName="Helvetica-Oblique"
        )
        heading_style = ParagraphStyle(
            "KRHeading", parent=styles["Heading2"],
            fontSize=12, textColor=colors.black,
            spaceAfter=6, spaceBefore=12,
            fontName="Helvetica-Bold", alignment=TA_LEFT
        )
        normal_style = ParagraphStyle(
            "KRNormal", parent=styles["Normal"],
            fontSize=10, spaceAfter=6,
            alignment=TA_LEFT, fontName="Helvetica"
        )
        italic_style = ParagraphStyle(
            "KRItalic", parent=styles["Normal"],
            fontSize=9, spaceAfter=6,
            alignment=TA_LEFT, fontName="Helvetica-Oblique"
        )
        table_title_style = ParagraphStyle(
            "KRTableTitle", parent=styles["Normal"],
            fontSize=11, spaceAfter=6, spaceBefore=12,
            fontName="Helvetica-Oblique", alignment=TA_LEFT
        )
        footer_style = ParagraphStyle(
            "KRFooter", parent=styles["Normal"],
            fontSize=8, textColor=colors.grey,
            alignment=TA_LEFT, fontName="Helvetica-Oblique"
        )

        # ── Title, Subtitle & Author ─────────────────────────────────────────
        report_title = title if title else "KR-20 Reliability Analysis"
        elements.append(Paragraph(report_title, title_style))
        elements.append(Spacer(1, 0.05 * inch))

        if subtitle and subtitle.strip():
            elements.append(Paragraph(subtitle, subtitle_style))
            elements.append(Spacer(1, 0.05 * inch))

        if byline and byline.strip():
            elements.append(Paragraph(byline, byline_style))

        elements.append(Spacer(1, 0.15 * inch))

        # ── Description ──────────────────────────────────────────────────────
        if description and description.strip():
            elements.append(Paragraph(description, normal_style))
            elements.append(Spacer(1, 0.2 * inch))

        # ── ROW 1: KR-20 Coefficient (left) + Summary Statistics (right) ─────
        # Wrapped in a 2-column outer Table so both blocks sit side by side.
        elements.append(
            Paragraph("<i>Frequentist Scale Reliability Statistics</i>", table_title_style)
        )

        # Build inner coefficient table data
        coef_data = [
            ["Coefficient", "Estimate"],
            ["KR-20", f'{results["kr20"]:.3f}'],
        ]
        coef_inner = Table(coef_data, colWidths=[1.3 * inch, 1.1 * inch])
        coef_inner.setStyle(TableStyle([
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",      (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.white),
            ("LINEABOVE",     (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",     (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",     (0, -1), (-1, -1), 0.5, colors.black),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ]))

        # Build inner summary table data
        summary_data = [
            ["Statistic", "Value"],
            ["No. of Items",          str(results["n_items"])],
            ["No. of Respondents",    str(results["n_respondents"])],
            ["Total Score Variance",  f'{results["total_variance"]:.4f}'],
            ["Sum of p·q",       f'{results["sum_pq"]:.4f}'],
            ["Interpretation",        results["interpretation"]],
        ]
        sum_inner = Table(summary_data, colWidths=[1.6 * inch, 1.3 * inch])
        sum_inner.setStyle(TableStyle([
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",      (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.white),
            ("LINEABOVE",     (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",     (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",     (0, -1), (-1, -1), 0.5, colors.black),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ]))

        # Outer 2-column table holding both inner tables
        outer_row1 = Table(
            [[coef_inner, sum_inner]],
            colWidths=[2.6 * inch, 3.1 * inch],
            hAlign="LEFT"
        )
        outer_row1.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        elements.append(outer_row1)
        elements.append(Spacer(1, 0.25 * inch))

        # ── ROW 2: Item Difficulty split into two side-by-side half-tables ───
        elements.append(Paragraph("Item Difficulty Statistics (p values)", heading_style))

        item_col_headers = ["Item", "p", "q", "p·q"]
        all_item_rows = []
        for item_name, stats in results["item_stats"].items():
            all_item_rows.append([
                str(item_name),
                f'{stats["p"]:.4f}',
                f'{stats["q"]:.4f}',
                f'{stats["pq"]:.4f}',
            ])

        # Split items into two halves
        mid = (len(all_item_rows) + 1) // 2
        left_items  = [item_col_headers] + all_item_rows[:mid]
        right_items = [item_col_headers] + (all_item_rows[mid:] if all_item_rows[mid:] else [["—", "—", "—", "—"]])

        item_style = TableStyle([
            ("FONTNAME",         (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",         (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",         (0, 0), (-1, -1), 8),
            ("ALIGN",            (0, 0), (0, -1),  "LEFT"),
            ("ALIGN",            (1, 0), (-1, -1), "CENTER"),
            ("BACKGROUND",       (0, 0), (-1, 0),  colors.white),
            ("LINEABOVE",        (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",        (0, 0), (-1, 0),  0.5, colors.black),
            ("LINEBELOW",        (0, -1), (-1, -1), 0.5, colors.black),
            ("ROWBACKGROUNDS",   (0, 1), (-1, -1),
             [colors.white, colors.Color(0.95, 0.95, 0.95)]),
            ("TOPPADDING",       (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",    (0, 0), (-1, -1), 3),
            ("LEFTPADDING",      (0, 0), (-1, -1), 3),
        ])

        item_cw = [0.55 * inch, 0.6 * inch, 0.6 * inch, 0.55 * inch]
        left_tbl  = Table(left_items,  colWidths=item_cw)
        right_tbl = Table(right_items, colWidths=item_cw)
        left_tbl.setStyle(item_style)
        right_tbl.setStyle(item_style)

        outer_row2 = Table(
            [[left_tbl, right_tbl]],
            colWidths=[2.55 * inch, 2.55 * inch],
            hAlign="LEFT"
        )
        outer_row2.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        elements.append(outer_row2)
        elements.append(Spacer(1, 0.25 * inch))

        # ── Interpretation Guide (two columns) ────────────────────────────────
        elements.append(Paragraph("Interpretation Guide", heading_style))

        scale_items = [
            ("KR-20 ≥ 0.90",        "Excellent"),
            ("0.80 ≤ KR-20 < 0.90", "Good"),
            ("0.70 ≤ KR-20 < 0.80", "Acceptable"),
            ("0.60 ≤ KR-20 < 0.70", "Questionable"),
            ("0.50 ≤ KR-20 < 0.60", "Poor"),
            ("KR-20 < 0.50",             "Unacceptable"),
        ]

        mid_g = (len(scale_items) + 1) // 2
        left_guide  = scale_items[:mid_g]
        right_guide = scale_items[mid_g:]

        guide_label_style = ParagraphStyle(
            "KRGuideLabel", parent=styles["Normal"],
            fontSize=9, fontName="Helvetica-Bold"
        )
        guide_val_style = ParagraphStyle(
            "KRGuideVal", parent=styles["Normal"],
            fontSize=9, fontName="Helvetica"
        )

        def _guide_cell(threshold, meaning):
            return Paragraph(
                f"<b>{threshold}:</b> {meaning}", 
                ParagraphStyle("gi", parent=styles["Normal"], fontSize=9, fontName="Helvetica")
            )

        # Pad shorter column with empty string
        while len(right_guide) < len(left_guide):
            right_guide.append(("", ""))

        guide_rows = [
            [_guide_cell(l[0], l[1]), _guide_cell(r[0], r[1])]
            for l, r in zip(left_guide, right_guide)
        ]

        guide_tbl = Table(guide_rows, colWidths=[2.8 * inch, 2.8 * inch], hAlign="LEFT")
        guide_tbl.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING",   (0, 0), (-1, -1), 2),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ]))
        elements.append(guide_tbl)

        # ── Footer ────────────────────────────────────────────────────────────
        elements.append(Spacer(1, 0.5 * inch))
        timestamp = datetime.now().strftime("%B %d, %Y at %H:%M")
        footer_text = (
            f"File: {os.path.abspath(filename)}<br/>"
            f"Generated: {timestamp}"
        )
        elements.append(Paragraph(footer_text, footer_style))

        doc.build(elements)


# ============================================================================
# DOCX Report Generator
# ============================================================================

class DOCXReportGenerator:
    """
    Generates a professional DOCX report for KR-20 analysis using python-docx.
    """

    @staticmethod
    def _set_cell_bg(cell, hex_color: str):
        """Helper to set table cell background colour."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    @staticmethod
    def generate_report(results: dict, description: str, filename: str,
                        title: str = None, subtitle: str = None, byline: str = None):
        """
        Build and save the DOCX report.

        Parameters
        ----------
        results     : dict
        description : str
        filename    : str
        title       : str
        subtitle    : str
        byline      : str
        """
        doc = Document()

        # Page margins
        section = doc.sections[0]
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

        # ── Title ─────────────────────────────────────────────────────────────
        report_title = title if title else "KR-20 Reliability Analysis"
        h = doc.add_heading(report_title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # ── Subtitle ──────────────────────────────────────────────────────────
        if subtitle and subtitle.strip():
            sp = doc.add_paragraph(subtitle)
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sp.runs:
                run.font.size = Pt(13)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── Author ────────────────────────────────────────────────────────────
        if byline and byline.strip():
            p = doc.add_paragraph(byline)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(12)

        doc.add_paragraph()  # spacer

        # ── Description ───────────────────────────────────────────────────────
        if description and description.strip():
            doc.add_heading("Study Description", level=2)
            p = doc.add_paragraph(description)
            p.paragraph_format.space_after = Pt(12)

        # ── ROW 1: KR-20 Coefficient (left) + Summary Statistics (right) ────────
        # Uses a 2-column outer table so both blocks sit side by side on one page.
        doc.add_heading("Reliability Statistics", level=2)

        outer1 = doc.add_table(rows=1, cols=2)
        outer1.style = "Table Grid"
        outer1.alignment = WD_TABLE_ALIGNMENT.LEFT

        # ── Left cell: KR-20 coefficient ────────────────────────────────────
        left_cell1 = outer1.cell(0, 0)
        left_cell1.width = Inches(2.8)

        # Section label inside left cell
        lp = left_cell1.paragraphs[0]
        lr = lp.add_run("Coefficient")
        lr.bold = True
        lr.font.size = Pt(9)
        DOCXReportGenerator._set_cell_bg(left_cell1, "EAF0FB")

        inner_coef = left_cell1.add_table(rows=2, cols=2)
        inner_coef.style = "Table Grid"
        for i, hdr in enumerate(["Coefficient", "Estimate"]):
            c = inner_coef.cell(0, i)
            c.text = hdr
            DOCXReportGenerator._set_cell_bg(c, "D9E1F2")
            for run in c.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for i, val in enumerate(["KR-20", f'{results["kr20"]:.3f}']):
            c = inner_coef.cell(1, i)
            c.text = val
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)

        # ── Right cell: Summary Statistics ───────────────────────────────────
        right_cell1 = outer1.cell(0, 1)
        right_cell1.width = Inches(3.6)

        rp = right_cell1.paragraphs[0]
        rr = rp.add_run("Summary Statistics")
        rr.bold = True
        rr.font.size = Pt(9)
        DOCXReportGenerator._set_cell_bg(right_cell1, "EAF0FB")

        stats_rows = [
            ("No. of Items",            str(results["n_items"])),
            ("No. of Respondents",      str(results["n_respondents"])),
            ("Total Score Variance",    f'{results["total_variance"]:.4f}'),
            ("Sum of p·q",              f'{results["sum_pq"]:.4f}'),
            ("Interpretation",          results["interpretation"]),
        ]

        inner_sum = right_cell1.add_table(rows=len(stats_rows) + 1, cols=2)
        inner_sum.style = "Table Grid"
        for i, hdr in enumerate(["Statistic", "Value"]):
            c = inner_sum.cell(0, i)
            c.text = hdr
            DOCXReportGenerator._set_cell_bg(c, "D9E1F2")
            for run in c.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for r_idx, (stat, val) in enumerate(stats_rows, start=1):
            inner_sum.cell(r_idx, 0).text = stat
            inner_sum.cell(r_idx, 1).text = val
            for ci in range(2):
                for run in inner_sum.cell(r_idx, ci).paragraphs[0].runs:
                    run.font.size = Pt(9)

        doc.add_paragraph()

        # ── ROW 2: Item Difficulty (left half + right half side by side) ─────
        doc.add_heading("Item Difficulty Statistics (p values)", level=2)

        item_col_headers = ["Item", "p", "q", "p·q"]
        item_rows_data = [
            (name,
             f'{s["p"]:.4f}',
             f'{s["q"]:.4f}',
             f'{s["pq"]:.4f}')
            for name, s in results["item_stats"].items()
        ]

        # Split items into two halves for side-by-side columns
        mid = (len(item_rows_data) + 1) // 2
        left_items  = item_rows_data[:mid]
        right_items = item_rows_data[mid:]

        outer2 = doc.add_table(rows=1, cols=2)
        outer2.style = "Table Grid"
        outer2.alignment = WD_TABLE_ALIGNMENT.LEFT

        def _fill_item_half(container_cell, items_subset):
            """Fill one half-column with an item difficulty mini-table."""
            DOCXReportGenerator._set_cell_bg(container_cell, "FAFAFA")
            inner = container_cell.add_table(
                rows=len(items_subset) + 1, cols=4
            )
            inner.style = "Table Grid"
            for ci, hdr in enumerate(item_col_headers):
                c = inner.cell(0, ci)
                c.text = hdr
                DOCXReportGenerator._set_cell_bg(c, "D9E1F2")
                for run in c.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
            for r_idx, row_vals in enumerate(items_subset, start=1):
                for ci, val in enumerate(row_vals):
                    c = inner.cell(r_idx, ci)
                    c.text = val
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(8)

        _fill_item_half(outer2.cell(0, 0), left_items)
        _fill_item_half(outer2.cell(0, 1), right_items if right_items else [("—", "—", "—", "—")])

        doc.add_paragraph()

        # ── Interpretation Guide (two-column bullet layout) ───────────────────
        doc.add_heading("Interpretation Guide", level=2)

        scale_items = [
            ("KR-20 ≥ 0.90",        "Excellent"),
            ("0.80 ≤ KR-20 < 0.90", "Good"),
            ("0.70 ≤ KR-20 < 0.80", "Acceptable"),
            ("0.60 ≤ KR-20 < 0.70", "Questionable"),
            ("0.50 ≤ KR-20 < 0.60", "Poor"),
            ("KR-20 < 0.50",         "Unacceptable"),
        ]

        # Two-column guide table (no borders for clean look)
        mid_g = (len(scale_items) + 1) // 2
        left_guide  = scale_items[:mid_g]
        right_guide = scale_items[mid_g:]

        guide_tbl = doc.add_table(rows=max(len(left_guide), len(right_guide)), cols=2)
        guide_tbl.style = "Table Grid"
        guide_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        def _write_guide_cell(cell, threshold, meaning):
            cell.paragraphs[0].clear()
            run_t = cell.paragraphs[0].add_run(f"{threshold}: ")
            run_t.bold = True
            run_t.font.size = Pt(9)
            run_m = cell.paragraphs[0].add_run(meaning)
            run_m.font.size = Pt(9)

        for r_idx, (threshold, meaning) in enumerate(left_guide):
            _write_guide_cell(guide_tbl.cell(r_idx, 0), threshold, meaning)
        for r_idx, (threshold, meaning) in enumerate(right_guide):
            _write_guide_cell(guide_tbl.cell(r_idx, 1), threshold, meaning)

        # ── Footer ────────────────────────────────────────────────────────────
        doc.add_paragraph()
        timestamp = datetime.now().strftime("%B %d, %Y at %H:%M")
        footer_p = doc.add_paragraph(
            f"File: {os.path.abspath(filename)}\nGenerated: {timestamp}"
        )
        for run in footer_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.save(filename)


# ============================================================================
# Data Table Frame – Scrollable Dataset Preview
# ============================================================================

class DataTableFrame(ctk.CTkScrollableFrame):
    """
    Scrollable frame for previewing the loaded dataset.
    Displays all rows with row numbers and column headers.
    """

    def __init__(self, master, **kwargs):
        super().__init__(master, **kwargs)

    def display_data(self, df: pd.DataFrame, max_rows: int = None):
        """
        Render the DataFrame inside the scrollable frame.

        Parameters
        ----------
        df       : pd.DataFrame – data to display
        max_rows : int|None     – if set, only first max_rows rows are shown
        """
        # Clear previous widgets
        for widget in self.winfo_children():
            widget.destroy()

        if df is None or df.empty:
            ctk.CTkLabel(
                self, text="No data loaded",
                font=("Arial", 13), text_color="gray"
            ).pack(pady=40)
            return

        # Info bar
        ctk.CTkLabel(
            self,
            text=f"📊 {len(df)} rows × {len(df.columns)} columns  (showing all data)",
            font=("Arial", 13, "bold")
        ).pack(pady=(8, 10))

        preview_df = df if max_rows is None else df.head(max_rows)
        show_more  = (max_rows is not None and len(df) > max_rows)

        # Header row
        header_frame = ctk.CTkFrame(self, fg_color=("gray80", "gray30"))
        header_frame.pack(fill="x", padx=5, pady=5)

        ctk.CTkLabel(
            header_frame, text="#",
            font=("Arial", 11, "bold"), width=50
        ).pack(side="left", padx=2, pady=4)

        for col in preview_df.columns:
            col_name = str(col)[:14] if col is not None else "N/A"
            ctk.CTkLabel(
                header_frame, text=col_name,
                font=("Arial", 11, "bold"), width=90
            ).pack(side="left", padx=2, pady=4)

        # Data rows
        for display_idx, (orig_idx, row) in enumerate(preview_df.iterrows()):
            row_frame = ctk.CTkFrame(self, fg_color="transparent")
            row_frame.pack(fill="x", padx=5, pady=1)

            ctk.CTkLabel(
                row_frame, text=str(display_idx + 1),
                font=("Arial", 10), width=50, text_color="gray"
            ).pack(side="left", padx=2)

            for val in row:
                val_str = str(val)[:14] if pd.notna(val) else ""
                ctk.CTkLabel(
                    row_frame, text=val_str,
                    font=("Arial", 10), width=90
                ).pack(side="left", padx=2)

        if show_more:
            ctk.CTkLabel(
                self,
                text=f"… and {len(df) - max_rows} more rows",
                font=("Arial", 10, "italic"),
                text_color="gray"
            ).pack(pady=8)


# ============================================================================
# Main Application
# ============================================================================

class KR20ReliabilityApp(ctk.CTk):
    """
    Main application window for the KR-20 Reliability Test Tool.

    Layout:
        LEFT   – Data input, binary frequency generator, report config
        CENTER – Dataset preview (scrollable table)
        RIGHT  – Analysis results

    Data flow:
        1. User imports CSV/XLSX  OR  enters binary frequencies
        2. Dataset is previewed in the centre column
        3. User clicks Compute KR-20 → results shown on right
        4. User exports PDF and/or DOCX report
    """

    def __init__(self):
        super().__init__()

        self.title("KR-20 Reliability Test")
        self.geometry("1440x860")
        self.minsize(1100, 700)

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # ── State ────────────────────────────────────────────────────────────
        self.df: pd.DataFrame | None = None
        self.results: dict | None    = None
        self.current_mode            = "light"

        # Binary frequency generator state
        self.freq_expanded   = False
        self.freq_entries    = {}   # {item_name: {"correct": Entry, "incorrect": Entry}}
        self.freq_num_items  = 0

        self._build_ui()

    # ────────────────────────────────────────────────────────────────────────
    # UI Construction
    # ────────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        """Construct the entire UI."""
        self._build_header()
        self._build_main_columns()

    def _build_header(self):
        """Top bar: title + theme toggle."""
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=15, pady=(15, 10))

        ctk.CTkLabel(
            header, text="KR-20 Reliability Test",
            font=("Arial", 26, "bold")
        ).pack(side="left", padx=10)

        ctk.CTkLabel(
            header, text="Kuder-Richardson Formula 20 · Binary Test Data",
            font=("Arial", 12, "italic"), text_color="gray"
        ).pack(side="left", padx=6)

        self.theme_btn = ctk.CTkButton(
            header, text="🌙 Dark Mode",
            command=self._toggle_theme,
            width=140, height=38,
            font=("Arial", 12, "bold"),
            corner_radius=8
        )
        self.theme_btn.pack(side="right", padx=10)

    def _build_main_columns(self):
        """Three-column main layout."""
        container = ctk.CTkFrame(self, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        container.grid_columnconfigure(0, weight=1, minsize=380)
        container.grid_columnconfigure(1, weight=1, minsize=380)
        container.grid_columnconfigure(2, weight=1, minsize=380)
        container.grid_rowconfigure(0, weight=1)

        self._build_left_column(container)
        self._build_center_column(container)
        self._build_right_column(container)

    # ── LEFT COLUMN ──────────────────────────────────────────────────────────

    def _build_left_column(self, parent):
        left = ctk.CTkScrollableFrame(parent, fg_color=("gray90", "gray20"))
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 5))

        ctk.CTkLabel(
            left, text="📊 Data Input",
            font=("Arial", 18, "bold")
        ).pack(pady=(10, 15), padx=10)

        # ── Import Section ────────────────────────────────────────────────────
        import_sec = ctk.CTkFrame(left)
        import_sec.pack(fill="x", padx=10, pady=(0, 15))

        ctk.CTkLabel(
            import_sec, text="Import Existing Data",
            font=("Arial", 14, "bold")
        ).pack(pady=(10, 8), padx=10, anchor="w")

        ctk.CTkButton(
            import_sec, text="📁 Import CSV / Excel",
            command=self._import_data,
            height=40, font=("Arial", 13, "bold"), corner_radius=8
        ).pack(fill="x", padx=10, pady=(0, 10))

        # ── Binary Frequency Generator ────────────────────────────────────────
        self.freq_frame = ctk.CTkFrame(left)
        self.freq_frame.pack(fill="x", padx=10, pady=(0, 15))

        self.freq_toggle_btn = ctk.CTkButton(
            self.freq_frame,
            text="▶ Generate from Binary Frequencies",
            command=self._toggle_freq_section,
            font=("Arial", 13, "bold"),
            fg_color="#6a4c93", hover_color="#553c7a",
            anchor="w", height=40, corner_radius=8
        )
        self.freq_toggle_btn.pack(fill="x", padx=10, pady=10)

        self.freq_content = ctk.CTkFrame(self.freq_frame, fg_color="transparent")

        ctk.CTkLabel(
            self.freq_content,
            text=(
                "Enter the number of respondents who answered\n"
                "each item Correctly (1) or Incorrectly (0)."
            ),
            font=("Arial", 11, "italic"),
            text_color="gray", wraplength=340
        ).pack(pady=(4, 10), padx=10)

        # Number of items
        cfg = ctk.CTkFrame(self.freq_content, fg_color="transparent")
        cfg.pack(fill="x", padx=10, pady=5)

        ctk.CTkLabel(
            cfg, text="Number of Items:", font=("Arial", 12, "bold")
        ).pack(anchor="w", pady=(0, 3))

        self.num_items_entry = ctk.CTkEntry(
            cfg, height=32, placeholder_text="e.g., 20",
            font=("Arial", 12)
        )
        self.num_items_entry.pack(fill="x", pady=(0, 10))

        ctk.CTkButton(
            cfg, text="Create Input Fields",
            command=self._create_freq_fields,
            font=("Arial", 12, "bold"), height=36, corner_radius=6
        ).pack(fill="x")

        # Scrollable input grid
        self.freq_grid = ctk.CTkScrollableFrame(
            self.freq_content, height=230,
            fg_color=("white", "gray25")
        )
        self.freq_grid.pack(fill="both", expand=True, padx=10, pady=10)

        # Action buttons
        act = ctk.CTkFrame(self.freq_content, fg_color="transparent")
        act.pack(fill="x", padx=10, pady=(5, 10))

        ctk.CTkButton(
            act, text="✓ Generate Dataset",
            command=self._generate_binary_dataset,
            font=("Arial", 12, "bold"),
            fg_color="#2a9d8f", hover_color="#238276",
            height=38, corner_radius=6
        ).pack(fill="x", pady=(0, 5))

        btn_row = ctk.CTkFrame(act, fg_color="transparent")
        btn_row.pack(fill="x")

        ctk.CTkButton(
            btn_row, text="✗ Clear",
            command=self._clear_freq_fields,
            font=("Arial", 11, "bold"),
            fg_color="#e76f51", hover_color="#d45a3f",
            height=34, corner_radius=6
        ).pack(side="left", fill="x", expand=True, padx=(0, 3))

        ctk.CTkButton(
            btn_row, text="💾 Export Dataset",
            command=self._export_dataset,
            font=("Arial", 11, "bold"),
            height=34, corner_radius=6
        ).pack(side="left", fill="x", expand=True, padx=(3, 0))

        # ── Report Configuration ──────────────────────────────────────────────
        meta = ctk.CTkFrame(left)
        meta.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(
            meta, text="Report Configuration",
            font=("Arial", 14, "bold")
        ).pack(pady=(10, 10), padx=10, anchor="w")

        ctk.CTkLabel(meta, text="Title:", font=("Arial", 11, "bold")).pack(
            anchor="w", padx=10, pady=(0, 3))
        self.title_entry = ctk.CTkEntry(
            meta, font=("Arial", 12), height=32, placeholder_text="Report title"
        )
        self.title_entry.pack(fill="x", padx=10, pady=(0, 10))
        self.title_entry.insert(0, "KR-20 Reliability Analysis")

        ctk.CTkLabel(meta, text="Subtitle:", font=("Arial", 11, "bold")).pack(
            anchor="w", padx=10, pady=(0, 3))
        self.subtitle_entry = ctk.CTkEntry(
            meta, font=("Arial", 12), height=32, placeholder_text="e.g., A Study on..."
        )
        self.subtitle_entry.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(meta, text="Author:", font=("Arial", 11, "bold")).pack(
            anchor="w", padx=10, pady=(0, 3))
        self.byline_entry = ctk.CTkEntry(
            meta, font=("Arial", 12), height=32, placeholder_text="Author name"
        )
        self.byline_entry.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(meta, text="Description:", font=("Arial", 11, "bold")).pack(
            anchor="w", padx=10, pady=(0, 3))
        self.description_text = ctk.CTkTextbox(meta, height=90, font=("Arial", 12))
        self.description_text.pack(fill="x", padx=10, pady=(0, 10))

    # ── CENTER COLUMN ─────────────────────────────────────────────────────────

    def _build_center_column(self, parent):
        center = ctk.CTkFrame(parent, fg_color=("gray90", "gray20"))
        center.grid(row=0, column=1, sticky="nsew", padx=5)

        ctk.CTkLabel(
            center, text="📋 Dataset Preview",
            font=("Arial", 18, "bold")
        ).pack(pady=(10, 15), padx=10)

        # Preview table
        preview_container = ctk.CTkFrame(center, fg_color=("white", "gray25"))
        preview_container.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.data_frame = DataTableFrame(
            preview_container, fg_color=("white", "gray25")
        )
        self.data_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # Action buttons
        btn_container = ctk.CTkFrame(center, fg_color="transparent")
        btn_container.pack(fill="x", padx=10, pady=(5, 10))

        ctk.CTkButton(
            btn_container, text="📊 Compute KR-20",
            command=self._compute_kr20,
            height=45, font=("Arial", 14, "bold"),
            fg_color="#2a9d8f", hover_color="#238276",
            corner_radius=8
        ).pack(fill="x", pady=(0, 8))

        export_row = ctk.CTkFrame(btn_container, fg_color="transparent")
        export_row.pack(fill="x")

        ctk.CTkButton(
            export_row, text="📄 Export PDF",
            command=self._export_pdf,
            height=40, font=("Arial", 12, "bold"),
            fg_color="#e76f51", hover_color="#d45a3f",
            corner_radius=8
        ).pack(side="left", fill="x", expand=True, padx=(0, 4))

        ctk.CTkButton(
            export_row, text="📝 Export DOCX",
            command=self._export_docx,
            height=40, font=("Arial", 12, "bold"),
            fg_color="#457b9d", hover_color="#356080",
            corner_radius=8
        ).pack(side="left", fill="x", expand=True, padx=(4, 0))

    # ── RIGHT COLUMN ──────────────────────────────────────────────────────────

    def _build_right_column(self, parent):
        right = ctk.CTkFrame(parent, fg_color=("gray90", "gray20"))
        right.grid(row=0, column=2, sticky="nsew", padx=(5, 0))

        ctk.CTkLabel(
            right, text="📈 Analysis Results",
            font=("Arial", 18, "bold")
        ).pack(pady=(10, 15), padx=10)

        results_container = ctk.CTkFrame(right, fg_color=("white", "gray25"))
        results_container.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.results_text = ctk.CTkTextbox(
            results_container, font=("Courier", 11), wrap="word"
        )
        self.results_text.pack(fill="both", expand=True, padx=8, pady=8)
        self.results_text.insert("1.0",
            "═══════════════════════════════════════\n"
            "  KR-20 Reliability Calculator\n"
            "  Kuder-Richardson Formula 20\n"
            "═══════════════════════════════════════\n\n"
            "Ready to analyze your binary test data!\n\n"
            "Steps:\n"
            "  1. Import a CSV/Excel file  OR\n"
            "     generate data from binary\n"
            "     frequency counts\n"
            "  2. Click 'Compute KR-20'\n"
            "  3. Review results here\n"
            "  4. Export PDF or DOCX report\n\n"
            "Note: Dataset must contain only\n"
            "binary values (0 = incorrect,\n"
            "1 = correct).\n\n"
            "Waiting for data…"
        )
        self.results_text.configure(state="disabled")

    # ────────────────────────────────────────────────────────────────────────
    # Event Handlers – Theme
    # ────────────────────────────────────────────────────────────────────────

    def _toggle_theme(self):
        """Switch between light and dark appearance."""
        if self.current_mode == "light":
            ctk.set_appearance_mode("dark")
            self.current_mode = "dark"
            self.theme_btn.configure(text="☀️ Light Mode")
        else:
            ctk.set_appearance_mode("light")
            self.current_mode = "light"
            self.theme_btn.configure(text="🌙 Dark Mode")

    # ────────────────────────────────────────────────────────────────────────
    # Event Handlers – Data Import
    # ────────────────────────────────────────────────────────────────────────

    def _import_data(self):
        """Open a file dialog to import CSV or Excel data."""
        filename = filedialog.askopenfilename(
            title="Select binary test data file",
            filetypes=[
                ("CSV files",   "*.csv"),
                ("Excel files", "*.xlsx"),
                ("All files",   "*.*"),
            ]
        )
        if not filename:
            return

        try:
            if filename.lower().endswith(".csv"):
                self.df = pd.read_csv(filename)
            elif filename.lower().endswith(".xlsx"):
                self.df = pd.read_excel(filename)
            else:
                messagebox.showerror("Error", "Unsupported format. Use CSV or XLSX.")
                return

            # Validate binary
            is_valid, err = KR20Calculator.validate_binary_data(self.df)
            if not is_valid:
                if not messagebox.askyesno(
                    "⚠ Non-binary Data",
                    f"{err}\n\nProceed anyway? (Non-binary values will cause errors.)"
                ):
                    self.df = None
                    return

            self.data_frame.display_data(self.df)
            self._set_results_text(
                f"═══════════════════════════════════════\n"
                f"  Data Imported Successfully!\n"
                f"═══════════════════════════════════════\n\n"
                f"File:    {os.path.basename(filename)}\n"
                f"Rows:    {len(self.df)}\n"
                f"Columns: {len(self.df.columns)}\n\n"
                f"Columns:\n"
                f"  {', '.join(str(c) for c in self.df.columns)}\n\n"
                f"✓ Ready for KR-20 analysis!\n\n"
                f"Click 'Compute KR-20' to calculate\n"
                f"the reliability coefficient."
            )
            messagebox.showinfo(
                "Import Successful",
                f"Loaded {len(self.df)} rows × {len(self.df.columns)} columns."
            )

        except Exception as exc:
            messagebox.showerror("Import Error", f"Failed to load file:\n{exc}")

    # ────────────────────────────────────────────────────────────────────────
    # Event Handlers – Binary Frequency Generator
    # ────────────────────────────────────────────────────────────────────────

    def _toggle_freq_section(self):
        """Show or hide the binary frequency generator panel."""
        if self.freq_expanded:
            self.freq_content.pack_forget()
            self.freq_toggle_btn.configure(text="▶ Generate from Binary Frequencies")
            self.freq_expanded = False
        else:
            self.freq_content.pack(fill="both", expand=True, padx=5, pady=5)
            self.freq_toggle_btn.configure(text="▼ Generate from Binary Frequencies")
            self.freq_expanded = True

    def _create_freq_fields(self):
        """Dynamically create one input row per item in the grid."""
        raw = self.num_items_entry.get().strip()
        try:
            num_items = int(raw)
        except ValueError:
            messagebox.showerror("Invalid Input", "Please enter a valid integer for number of items.")
            return

        if num_items < 2:
            messagebox.showwarning("Too Few Items", "Please enter at least 2 items.")
            return
        if num_items > 100:
            messagebox.showwarning("Too Many Items", "Maximum supported is 100 items.")
            return

        self.freq_num_items = num_items

        # Clear existing widgets
        for w in self.freq_grid.winfo_children():
            w.destroy()
        self.freq_entries = {}

        # Header
        hdr = ctk.CTkFrame(self.freq_grid, fg_color=("gray85", "gray30"))
        hdr.pack(fill="x", pady=(2, 6))

        ctk.CTkLabel(hdr, text="Item",       font=("Arial", 11, "bold"), width=70).pack(side="left", padx=2)
        ctk.CTkLabel(hdr, text="Correct (1)", font=("Arial", 11, "bold"), width=90).pack(side="left", padx=2)
        ctk.CTkLabel(hdr, text="Wrong (0)",  font=("Arial", 11, "bold"), width=90).pack(side="left", padx=2)

        # One row per item
        for idx in range(num_items):
            item_name = f"I{idx + 1}"
            row = ctk.CTkFrame(self.freq_grid, fg_color="transparent")
            row.pack(fill="x", pady=2)

            ctk.CTkLabel(row, text=item_name, font=("Arial", 11, "bold"), width=70).pack(side="left", padx=2)

            entry_correct   = ctk.CTkEntry(row, width=90, placeholder_text="0", height=28)
            entry_correct.pack(side="left", padx=2)

            entry_incorrect = ctk.CTkEntry(row, width=90, placeholder_text="0", height=28)
            entry_incorrect.pack(side="left", padx=2)

            self.freq_entries[item_name] = {
                "correct":   entry_correct,
                "incorrect": entry_incorrect,
            }

        messagebox.showinfo(
            "Fields Created",
            f"Input fields created for {num_items} items.\n"
            "Enter the frequency counts and click 'Generate Dataset'."
        )

    def _clear_freq_fields(self):
        """Reset all frequency entry fields to empty."""
        for item_name, entries in self.freq_entries.items():
            entries["correct"].delete(0, "end")
            entries["incorrect"].delete(0, "end")

    def _parse_freq_int(self, entry: ctk.CTkEntry, label: str) -> int | None:
        """Parse a non-negative integer from an Entry widget."""
        text = entry.get().strip()
        if not text:
            return 0
        try:
            val = int(text)
            if val < 0:
                raise ValueError
            return val
        except ValueError:
            messagebox.showerror(
                "Invalid Input",
                f"'{text}' is not a valid count for {label}. "
                "Please enter a non-negative integer."
            )
            return None

    def _generate_binary_dataset(self):
        """Read frequency fields, validate, expand to DataFrame, and display."""
        if not self.freq_entries:
            messagebox.showwarning("No Fields", "Please create input fields first.")
            return

        items_freq_data = {}

        for item_name, entries in self.freq_entries.items():
            correct   = self._parse_freq_int(entries["correct"],   f"{item_name} Correct")
            incorrect = self._parse_freq_int(entries["incorrect"], f"{item_name} Incorrect")

            if correct is None or incorrect is None:
                return  # error already shown

            is_valid, err = BinaryFrequencyExpander.validate_item_input(correct, incorrect)
            if not is_valid:
                messagebox.showerror("Validation Error", f"{item_name}: {err}")
                return

            items_freq_data[item_name] = {"correct": correct, "incorrect": incorrect}

        # Check for consistent totals
        totals = {n: d["correct"] + d["incorrect"] for n, d in items_freq_data.items()}
        if len(set(totals.values())) > 1:
            summary = "\n".join(f"  {n}: {t}" for n, t in totals.items())
            if not messagebox.askyesno(
                "⚠ Inconsistent Totals",
                f"Items have different total respondent counts:\n\n{summary}\n\n"
                "Rows with missing values will be padded with NaN.\n\n"
                "Proceed anyway?"
            ):
                return

        try:
            # Build expanded data with NaN padding for unequal totals
            expanded = {}
            max_len  = max(d["correct"] + d["incorrect"] for d in items_freq_data.values())

            for item_name, freq in items_freq_data.items():
                col = BinaryFrequencyExpander.expand_item(freq["correct"], freq["incorrect"])
                if len(col) < max_len:
                    col += [np.nan] * (max_len - len(col))
                expanded[item_name] = col

            self.df = pd.DataFrame(expanded)
            self.data_frame.display_data(self.df)

            self._set_results_text(
                f"═══════════════════════════════════════\n"
                f"  Dataset Generated!\n"
                f"═══════════════════════════════════════\n\n"
                f"Method:      Binary Frequency Expansion\n"
                f"Respondents: {len(self.df)}\n"
                f"Items:       {len(self.df.columns)}\n\n"
                f"Items: {', '.join(self.df.columns)}\n\n"
                f"✓ Ready for KR-20 analysis!\n\n"
                f"Click 'Compute KR-20' to continue."
            )
            messagebox.showinfo(
                "Dataset Ready",
                f"Generated {len(self.df)} respondents × {len(self.df.columns)} items.\n"
                "Scroll the centre panel to view all rows."
            )

        except Exception as exc:
            messagebox.showerror("Generation Error", f"Failed to generate dataset:\n{exc}")

    def _export_dataset(self):
        """Save the current DataFrame to XLSX or CSV."""
        if self.df is None:
            messagebox.showwarning("No Data", "No dataset available to export.")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")],
            initialfile=f"Binary_Dataset_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        if not filename:
            return

        try:
            if filename.lower().endswith(".csv"):
                self.df.to_csv(filename, index=False)
            else:
                self.df.to_excel(filename, index=False)
            messagebox.showinfo("Exported", f"Dataset saved:\n{os.path.basename(filename)}")
        except Exception as exc:
            messagebox.showerror("Export Error", f"Failed to export:\n{exc}")

    # ────────────────────────────────────────────────────────────────────────
    # Event Handlers – KR-20 Computation
    # ────────────────────────────────────────────────────────────────────────

    def _compute_kr20(self):
        """Run the KR-20 computation and display results."""
        if self.df is None:
            messagebox.showwarning("No Data", "Please import or generate a dataset first.")
            return

        try:
            self.results = KR20Calculator.compute_kr20(self.df)
        except ValueError as exc:
            messagebox.showerror("Computation Error", str(exc))
            return
        except Exception as exc:
            messagebox.showerror("Unexpected Error", f"KR-20 computation failed:\n{exc}")
            return

        r = self.results
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Build item difficulty block
        item_lines = []
        for item_name, stats in r["item_stats"].items():
            item_lines.append(
                f"  {item_name:<10}  p={stats['p']:.4f}  "
                f"q={stats['q']:.4f}  pq={stats['pq']:.4f}"
            )
        item_block = "\n".join(item_lines)

        output = (
            f"╔════════════════════════════════════════════╗\n"
            f"║   KR-20 RELIABILITY RESULTS                ║\n"
            f"║   Kuder-Richardson Formula 20              ║\n"
            f"╚════════════════════════════════════════════╝\n\n"
            f"Timestamp: {ts}\n\n"
            f"┌────────────────────────────────────────────┐\n"
            f"│ RELIABILITY COEFFICIENT                    │\n"
            f"├────────────────────────────────────────────┤\n"
            f"│ KR-20:               {r['kr20']:.4f}\n"
            f"│\n"
            f"│ Interpretation:      {r['interpretation']}\n"
            f"└────────────────────────────────────────────┘\n\n"
            f"┌────────────────────────────────────────────┐\n"
            f"│ DESCRIPTIVE STATISTICS                     │\n"
            f"├────────────────────────────────────────────┤\n"
            f"│ Number of Items:     {r['n_items']}\n"
            f"│ Respondents:         {r['n_respondents']}\n"
            f"│ Total Score Var:     {r['total_variance']:.4f}\n"
            f"│ Sum of p·q:          {r['sum_pq']:.4f}\n"
            f"└────────────────────────────────────────────┘\n\n"
            f"FORMULA USED:\n"
            f"  KR-20 = (k/(k-1)) × (1 - Σ(p·q) / σ²)\n\n"
            f"  k        = {r['n_items']} items\n"
            f"  Σ(p·q)   = {r['sum_pq']:.4f}\n"
            f"  σ²_total = {r['total_variance']:.4f}\n\n"
            f"ITEM DIFFICULTY (p values):\n"
            f"{'─'*44}\n"
            f"{item_block}\n\n"
            f"INTERPRETATION SCALE:\n"
            f"  • KR-20 ≥ 0.90 : Excellent\n"
            f"  • KR-20 ≥ 0.80 : Good\n"
            f"  • KR-20 ≥ 0.70 : Acceptable\n"
            f"  • KR-20 ≥ 0.60 : Questionable\n"
            f"  • KR-20 ≥ 0.50 : Poor\n"
            f"  • KR-20 <  0.50 : Unacceptable\n\n"
            f"✓ Ready to export PDF / DOCX report!"
        )

        self._set_results_text(output)
        messagebox.showinfo(
            "Analysis Complete",
            f"KR-20 computed successfully!\n\n"
            f"KR-20 = {r['kr20']:.4f}\n"
            f"Interpretation: {r['interpretation']}\n\n"
            f"Use the Export buttons to save your report."
        )

    # ────────────────────────────────────────────────────────────────────────
    # Event Handlers – Export
    # ────────────────────────────────────────────────────────────────────────

    def _get_report_meta(self) -> tuple[str, str, str, str]:
        """Return (title, subtitle, byline, description) from the form widgets."""
        title       = self.title_entry.get().strip() or "KR-20 Reliability Analysis"
        subtitle    = self.subtitle_entry.get().strip()
        byline      = self.byline_entry.get().strip()
        description = self.description_text.get("1.0", "end-1c")
        return title, subtitle, byline, description

    def _export_pdf(self):
        """Generate and save a PDF report."""
        if self.results is None:
            messagebox.showwarning("No Results", "Please compute KR-20 first.")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"KR20_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        )
        if not filename:
            return

        title, subtitle, byline, description = self._get_report_meta()
        try:
            PDFReportGenerator.generate_report(
                self.results, description, filename, title, subtitle, byline
            )
            messagebox.showinfo(
                "PDF Exported",
                f"Report saved:\n{os.path.basename(filename)}"
            )
        except Exception as exc:
            messagebox.showerror("Export Error", f"PDF export failed:\n{exc}")

    def _export_docx(self):
        """Generate and save a DOCX report."""
        if self.results is None:
            messagebox.showwarning("No Results", "Please compute KR-20 first.")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile=f"KR20_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        if not filename:
            return

        title, subtitle, byline, description = self._get_report_meta()
        try:
            DOCXReportGenerator.generate_report(
                self.results, description, filename, title, subtitle, byline
            )
            messagebox.showinfo(
                "DOCX Exported",
                f"Report saved:\n{os.path.basename(filename)}"
            )
        except Exception as exc:
            messagebox.showerror("Export Error", f"DOCX export failed:\n{exc}")

    # ────────────────────────────────────────────────────────────────────────
    # Utility
    # ────────────────────────────────────────────────────────────────────────

    def _set_results_text(self, text: str):
        """Replace the results textbox content."""
        self.results_text.configure(state="normal")
        self.results_text.delete("1.0", "end")
        self.results_text.insert("1.0", text)
        self.results_text.configure(state="disabled")


# ============================================================================
# Entry Point
# ============================================================================

def main():
    """Launch the KR-20 Reliability Test application."""
    app = KR20ReliabilityApp()
    app.mainloop()


if __name__ == "__main__":
    main()