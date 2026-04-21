"""
T-Test Analysis — Modern Sleek UI
Implements:
- One-Sample t-test
- Independent Samples t-test
- Paired Samples t-test
- APA-style reporting
- DOCX / PDF export
- Excel / CSV import
- Summary Statistics input mode (M, SD, N)
- Multi-Part Session Manager
"""

import customtkinter as ctk
from tkinter import messagebox, filedialog
import numpy as np
import pandas as pd
from scipy import stats
from datetime import datetime
import re
import os

try:
    from app_settings import SettingsManager, SettingsWindow
    SETTINGS_AVAILABLE = True
except ImportError:
    SETTINGS_AVAILABLE = False

try:
    from ttest_session_manager import _next_part_label, SessionManagerPanel
    SESSION_AVAILABLE = True
except ImportError:
    SESSION_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


# ─── Palette ──────────────────────────────────────────────────────────────────
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

BG_DEEP   = "#0d1117"
BG_CARD   = "#161b22"
BG_PANEL  = "#1c2230"
BG_INPUT  = "#1e2736"
ACCENT    = "#00c9a7"
ACCENT2   = "#4e9eff"
DANGER    = "#ef4444"
WARN      = "#f59e0b"
SUCCESS   = "#22c55e"
PURPLE    = "#a855f7"
TEXT_PRI  = "#e6edf3"
TEXT_SEC  = "#8b949e"
BORDER    = "#30363d"
TAB_ACT   = "#00c9a7"
TAB_INACT = "#1c2230"

FONT_HEAD = ("Segoe UI", 26, "bold")
FONT_CARD = ("Segoe UI", 15, "bold")
FONT_BODY = ("Segoe UI", 13)
FONT_MONO = ("Consolas", 12)
FONT_BTN  = ("Segoe UI", 13, "bold")
FONT_TINY = ("Segoe UI", 11)
FONT_LBL  = ("Segoe UI", 12, "bold")


# ─── Helpers ──────────────────────────────────────────────────────────────────

def divider(parent):
    ctk.CTkFrame(parent, height=1, fg_color=BORDER, corner_radius=0).pack(fill="x")


def styled_entry(parent, placeholder="", width=0, height=34):
    kw = dict(placeholder_text=placeholder, fg_color=BG_INPUT,
              border_color=BORDER, text_color=TEXT_PRI,
              placeholder_text_color=TEXT_SEC, border_width=1,
              corner_radius=6, height=height, font=FONT_BODY)
    e = ctk.CTkEntry(parent, **kw)
    if width:
        e.configure(width=width)
    return e


def sidebar_btn(parent, text, fg, hover, text_color=TEXT_PRI,
                font=FONT_BTN, height=36, state="normal"):
    return ctk.CTkButton(parent, text=text, fg_color=fg, hover_color=hover,
                         text_color=text_color, font=font, height=height,
                         corner_radius=8, state=state)


def card(parent, title="", **kw):
    frame = ctk.CTkFrame(parent, fg_color=BG_CARD, corner_radius=12,
                         border_width=1, border_color=BORDER, **kw)
    if title:
        ctk.CTkLabel(frame, text=title, font=FONT_CARD,
                     text_color=TEXT_PRI).pack(anchor="w", padx=16, pady=(12, 4))
    return frame


def section_label(parent, text):
    ctk.CTkLabel(parent, text=text, font=("Segoe UI", 11, "bold"),
                 text_color=TEXT_SEC).pack(anchor="w", padx=18, pady=(14, 3))


# ─── Summary Stats Input Widget ───────────────────────────────────────────────

class SummaryStatsFrame(ctk.CTkFrame):
    def __init__(self, parent, show_sd=True, **kw):
        super().__init__(parent, fg_color="transparent", **kw)
        self.show_sd = show_sd
        self._build()

    def _build(self):
        row = ctk.CTkFrame(self, fg_color="transparent")
        row.pack(fill="x", padx=12, pady=(4, 8))

        ctk.CTkLabel(row, text="M", font=("Segoe UI", 11, "bold"),
                     text_color=TEXT_SEC, width=18).pack(side="left", padx=(0, 4))
        self.mean_entry = styled_entry(row, placeholder="Mean", width=90, height=30)
        self.mean_entry.pack(side="left", padx=(0, 10))

        if self.show_sd:
            ctk.CTkLabel(row, text="SD", font=("Segoe UI", 11, "bold"),
                         text_color=TEXT_SEC, width=22).pack(side="left", padx=(0, 4))
            self.sd_entry = styled_entry(row, placeholder="Std Dev", width=90, height=30)
            self.sd_entry.pack(side="left", padx=(0, 10))
        else:
            self.sd_entry = None

        ctk.CTkLabel(row, text="N", font=("Segoe UI", 11, "bold"),
                     text_color=TEXT_SEC, width=18).pack(side="left", padx=(0, 4))
        self.n_entry = styled_entry(row, placeholder="Sample size", width=80, height=30)
        self.n_entry.pack(side="left")

    def get_values(self):
        m = float(self.mean_entry.get())
        n = int(self.n_entry.get())
        sd = float(self.sd_entry.get()) if self.sd_entry else None
        return m, sd, n

    def clear(self):
        self.mean_entry.delete(0, "end")
        if self.sd_entry:
            self.sd_entry.delete(0, "end")
        self.n_entry.delete(0, "end")


# ─── Group Card Widget ────────────────────────────────────────────────────────

class GroupCard(ctk.CTkFrame):
    def __init__(self, parent, badge_text, badge_color, group_key,
                 show_sd_in_summary=True, **kw):
        super().__init__(parent, fg_color=BG_PANEL, corner_radius=8,
                         border_width=1, border_color=BORDER, **kw)
        self.badge_color = badge_color
        self.group_key = group_key
        self.show_sd_in_summary = show_sd_in_summary
        self._mode = "raw"
        self._build(badge_text)

    def _build(self, badge_text):
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=12, pady=(10, 4))

        self.badge = ctk.CTkEntry(header, width=80, height=28,
                                  font=("Segoe UI", 10, "bold"),
                                  fg_color=self.badge_color,
                                  border_color=self.badge_color,
                                  text_color="#0d1117", justify="center",
                                  border_width=0, corner_radius=6)
        self.badge.insert(0, badge_text)
        self.badge.pack(side="left", padx=(0, 8))

        self.name_entry = styled_entry(header, placeholder="Group name", height=28)
        self.name_entry.insert(0, self.group_key)
        self.name_entry.pack(side="left", fill="x", expand=True)

        tab_bar = ctk.CTkFrame(self, fg_color=BG_INPUT, corner_radius=6)
        tab_bar.pack(fill="x", padx=12, pady=(2, 6))

        self.raw_btn = ctk.CTkButton(
            tab_bar, text="📋  Raw Data", height=26,
            font=("Segoe UI", 11, "bold"),
            fg_color=TAB_ACT, hover_color="#009e82",
            text_color="#0d1117", corner_radius=5,
            command=self._set_raw)
        self.raw_btn.pack(side="left", padx=(4, 2), pady=4, fill="x", expand=True)

        self.sum_btn = ctk.CTkButton(
            tab_bar, text="∑  Summary Stats", height=26,
            font=("Segoe UI", 11, "bold"),
            fg_color=TAB_INACT, hover_color=BORDER,
            text_color=TEXT_SEC, corner_radius=5,
            command=self._set_summary)
        self.sum_btn.pack(side="left", padx=(2, 4), pady=4, fill="x", expand=True)

        self.raw_frame = ctk.CTkFrame(self, fg_color="transparent")

        raw_header = ctk.CTkFrame(self.raw_frame, fg_color="transparent")
        raw_header.pack(fill="x", padx=12, pady=(2, 2))
        ctk.CTkLabel(raw_header, text="Data (comma-separated):",
                     font=FONT_TINY, text_color=TEXT_SEC).pack(side="left")
        ctk.CTkButton(raw_header, text="🗑 Clear", width=68, height=22,
                      font=("Segoe UI", 10, "bold"),
                      fg_color=BG_INPUT, hover_color=DANGER,
                      text_color=TEXT_SEC, corner_radius=5, border_width=1,
                      border_color=BORDER,
                      command=self._clear_raw_data).pack(side="right")

        self.data_text = ctk.CTkTextbox(self.raw_frame, height=80,
                                        fg_color=BG_INPUT, text_color=TEXT_PRI,
                                        border_width=1, border_color=BORDER,
                                        font=FONT_BODY, corner_radius=6)
        self.data_text.pack(fill="x", padx=12, pady=(0, 10))
        self.raw_frame.pack(fill="x")

        self.sum_frame = ctk.CTkFrame(self, fg_color="transparent")
        ctk.CTkLabel(self.sum_frame, text="Enter your pre-computed statistics below:",
                     font=FONT_TINY, text_color=TEXT_SEC).pack(anchor="w", padx=12, pady=(2, 4))
        self.stats_input = SummaryStatsFrame(self.sum_frame,
                                             show_sd=self.show_sd_in_summary)
        self.stats_input.pack(fill="x")

        self._sum_note = ctk.CTkLabel(
            self.sum_frame,
            text="ℹ  Paired t-test requires same N for both groups",
            font=("Segoe UI", 10), text_color=WARN)
        self._sum_note.pack(anchor="w", padx=12, pady=(0, 6))
        self._sum_note.pack_forget()

    def _clear_raw_data(self):
        self.data_text.delete("1.0", "end")

    def set_show_paired_note(self, show):
        if show:
            self._sum_note.pack(anchor="w", padx=12, pady=(0, 6))
        else:
            self._sum_note.pack_forget()

    def _set_raw(self):
        self._mode = "raw"
        self.raw_btn.configure(fg_color=TAB_ACT, text_color="#0d1117")
        self.sum_btn.configure(fg_color=TAB_INACT, text_color=TEXT_SEC)
        self.sum_frame.pack_forget()
        self.raw_frame.pack(fill="x")

    def _set_summary(self):
        self._mode = "summary"
        self.sum_btn.configure(fg_color=ACCENT2, text_color="#0d1117")
        self.raw_btn.configure(fg_color=TAB_INACT, text_color=TEXT_SEC)
        self.raw_frame.pack_forget()
        self.sum_frame.pack(fill="x")

    @property
    def mode(self):
        return self._mode

    def get_name(self):
        return self.name_entry.get().strip() or self.group_key

    def get_raw_data(self):
        text = self.data_text.get("1.0", "end").strip()
        text = re.sub(r'[,\n\r\t]+', ' ', text)
        return [float(v) for v in re.findall(r'-?\d+\.?\d*', text)]

    def get_summary(self):
        return self.stats_input.get_values()

    def clear(self):
        self.data_text.delete("1.0", "end")
        self.stats_input.clear()
        self._set_raw()


# ─── Sidebar ──────────────────────────────────────────────────────────────────

class Sidebar(ctk.CTkFrame):
    def __init__(self, master, **kw):
        super().__init__(master, width=230, fg_color=BG_CARD,
                         corner_radius=0, **kw)
        self.pack_propagate(False)
        self._build()

    def _build(self):
        logo = ctk.CTkFrame(self, fg_color=ACCENT, corner_radius=0, height=64)
        logo.pack(fill="x")
        logo.pack_propagate(False)
        ctk.CTkLabel(logo, text="  t  ", font=("Segoe UI", 30, "bold"),
                     text_color="#0d1117", fg_color=ACCENT).pack(expand=True)

        ctk.CTkLabel(self, text="t-Test Analysis", font=("Segoe UI", 16, "bold"),
                     text_color=TEXT_PRI, fg_color=BG_CARD).pack(pady=(16, 2))
        ctk.CTkLabel(self, text="APA Format Calculator", font=FONT_TINY,
                     text_color=TEXT_SEC, fg_color=BG_CARD).pack(pady=(0, 10))

        divider(self)

        section_label(self, "TEST TYPE")
        self.test_type_var = ctk.StringVar(value="independent")
        self.test_menu = ctk.CTkOptionMenu(
            self, variable=self.test_type_var,
            values=["one-sample", "independent", "paired"],
            fg_color=BG_INPUT, button_color=ACCENT, button_hover_color="#009e82",
            text_color=TEXT_PRI, dropdown_fg_color=BG_PANEL,
            dropdown_text_color=TEXT_PRI, font=FONT_BODY, height=34,
            corner_radius=6)
        self.test_menu.pack(fill="x", padx=14, pady=(0, 6))

        section_label(self, "ALPHA LEVEL")
        self.alpha_entry = styled_entry(self, placeholder="0.05")
        self.alpha_entry.insert(0, "0.05")
        self.alpha_entry.pack(fill="x", padx=14, pady=(0, 6))

        section_label(self, "RESEARCHER NAME")
        self.researcher_entry = styled_entry(self, placeholder="e.g. Dr. John Smith")
        self.researcher_entry.pack(fill="x", padx=14, pady=(0, 6))

        divider(self)

        pad = {"fill": "x", "padx": 14, "pady": 4}

        self.import_btn = sidebar_btn(self, "📁  Import Excel / CSV",
                                      fg=ACCENT2, hover="#3b7ddd")
        self.import_btn.pack(**pad)

        self.preview_btn = sidebar_btn(self, "👁  Preview Data",
                                       fg=BG_PANEL, hover=BORDER,
                                       text_color=TEXT_SEC, state="disabled")
        self.preview_btn.pack(**pad)

        divider(self)

        self.run_btn = sidebar_btn(self, "▶   Run t-Test",
                                   fg=ACCENT, hover="#009e82",
                                   text_color="#0d1117",
                                   font=("Segoe UI", 13, "bold"), height=44)
        self.run_btn.pack(**pad)

        self.save_part_btn = sidebar_btn(self, "🗂  Save Part",
                                         fg="#7c3aed", hover="#5b21b6",
                                         font=("Segoe UI", 12, "bold"), height=40)
        self.save_part_btn.pack(**pad)

        self.pdf_btn = sidebar_btn(self, "📄  Export PDF",
                                   fg="#1d4ed8", hover="#1e3a8a", state="disabled")
        self.pdf_btn.pack(**pad)

        self.docx_btn = sidebar_btn(self, "💾  Export DOCX",
                                    fg=PURPLE, hover="#7e22ce", state="disabled")
        self.docx_btn.pack(**pad)

        self.session_btn = sidebar_btn(self, "📂  Session Manager",
                                       fg="#1e3a2f", hover="#14532d",
                                       font=("Segoe UI", 12, "bold"), height=40)
        self.session_btn.pack(**pad)

        self.clear_btn = sidebar_btn(self, "🗑  Clear All",
                                     fg=DANGER, hover="#b91c1c")
        self.clear_btn.pack(**pad)

        divider(self)
        self.theme_btn = sidebar_btn(self, "☀️  Light Mode",
                                     fg="#374151", hover="#4b5563",
                                     font=FONT_BODY, height=32)
        self.theme_btn.pack(fill="x", padx=14, pady=8)

        divider(self)
        self.settings_btn = sidebar_btn(self, "⚙   Settings",
                                        fg="#374151", hover="#4b5563",
                                        font=FONT_BODY, height=32)
        self.settings_btn.pack(fill="x", padx=14, pady=8)

        self.status_label = ctk.CTkLabel(self, text="", font=FONT_TINY,
                                         text_color=ACCENT, fg_color=BG_CARD,
                                         wraplength=200)
        self.status_label.pack(side="bottom", padx=12, pady=8)


# ─── Main App ─────────────────────────────────────────────────────────────────

class TTestApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("t-Test Analysis")
        self.geometry("1280x820")
        self.minsize(1100, 700)
        self.configure(fg_color=BG_DEEP)

        self.results     = None
        self.imported_data = None
        self.saved_parts = []
        self._session_win = None
        self.dark_mode   = True
        self.alpha       = 0.05

        self._build_ui()
        self._on_test_change("independent")
        self._prefill_next_part_label()

    # ── UI Build ──────────────────────────────────────────────────────────────

    def _build_ui(self):
        self.sidebar = Sidebar(self)
        self.sidebar.pack(side="left", fill="y")

        self.sidebar.test_menu.configure(command=self._on_test_change)
        self.sidebar.import_btn.configure(command=self.import_data)
        self.sidebar.preview_btn.configure(command=self.show_preview)
        self.sidebar.run_btn.configure(command=self.run_analysis)
        self.sidebar.save_part_btn.configure(command=self.save_current_part)
        self.sidebar.pdf_btn.configure(command=self.export_pdf)
        self.sidebar.docx_btn.configure(command=self.export_docx)
        self.sidebar.session_btn.configure(command=self.open_session_manager)
        self.sidebar.clear_btn.configure(command=self.clear_fields)
        self.sidebar.theme_btn.configure(command=self.toggle_theme)
        self.sidebar.settings_btn.configure(command=self.open_settings)

        content = ctk.CTkFrame(self, fg_color=BG_DEEP, corner_radius=0)
        content.pack(side="left", fill="both", expand=True)

        # Header
        header = ctk.CTkFrame(content, fg_color=BG_CARD, corner_radius=0, height=64)
        header.pack(fill="x")
        header.pack_propagate(False)

        meta = ctk.CTkFrame(header, fg_color=BG_CARD)
        meta.pack(side="left", padx=16)

        ctk.CTkLabel(meta, text="Title:", font=("Segoe UI", 13, "bold"),
                     text_color=TEXT_PRI).grid(row=0, column=0, padx=(0, 4))
        self.title_entry = styled_entry(meta, placeholder="Report Title", width=200)
        self.title_entry.grid(row=0, column=1, padx=4)

        ctk.CTkLabel(meta, text="Subtitle:", font=("Segoe UI", 13, "bold"),
                     text_color=TEXT_PRI).grid(row=0, column=2, padx=(12, 4))
        self.subtitle_entry = styled_entry(meta, placeholder="e.g. t-Test", width=170)
        self.subtitle_entry.insert(0, "t-Test")
        self.subtitle_entry.grid(row=0, column=3, padx=4)

        ctk.CTkLabel(meta, text="Part:", font=("Segoe UI", 12, "bold"),
                     text_color=ACCENT).grid(row=0, column=4, padx=(16, 4))
        self.part_label_entry = styled_entry(meta, placeholder="e.g. Part 1-A", width=110)
        self.part_label_entry.grid(row=0, column=5, padx=4)

        ctk.CTkLabel(header, text="t-Test Analysis",
                     font=("Segoe UI", 13), text_color=TEXT_SEC).pack(side="right", padx=24)

        import tkinter as tk
        from tkinter import ttk

        outer = ctk.CTkFrame(content, fg_color=BG_DEEP)
        outer.pack(fill="both", expand=True, padx=16, pady=16)

        style = ttk.Style()
        style.theme_use("default")
        style.configure("Sash", sashthickness=6, sashrelief="flat", background="#30363d")

        pane = ttk.PanedWindow(outer, orient=tk.HORIZONTAL)
        pane.pack(fill="both", expand=True)

        left_wrap = ctk.CTkFrame(pane, fg_color=BG_DEEP)
        left = card(left_wrap, title="📋  Data Input")
        left.pack(fill="both", expand=True)
        self._build_input_panel(left)
        pane.add(left_wrap, weight=1)

        right_wrap = ctk.CTkFrame(pane, fg_color=BG_DEEP)
        right = card(right_wrap, title="📊  Analysis Results")
        right.pack(fill="both", expand=True)
        self._build_results_panel(right)
        pane.add(right_wrap, weight=1)

        def _set_sash(*_):
            total = pane.winfo_width()
            if total > 100:
                pane.sashpos(0, total // 2)
                pane.unbind("<Configure>")
        pane.bind("<Configure>", _set_sash)

        bar = ctk.CTkFrame(content, fg_color=BG_CARD, height=28, corner_radius=0)
        bar.pack(fill="x", side="bottom")
        bar.pack_propagate(False)
        self.file_label = ctk.CTkLabel(bar, text="No file saved yet",
                                       font=FONT_TINY, text_color=TEXT_SEC)
        self.file_label.pack(side="left", padx=12)
        self.stat_label = ctk.CTkLabel(bar, text="", font=("Segoe UI", 9, "bold"),
                                       text_color=ACCENT)
        self.stat_label.pack(side="right", padx=12)

    def _build_input_panel(self, parent):
        self.scroll = ctk.CTkScrollableFrame(parent, fg_color=BG_CARD,
                                             scrollbar_button_color=BORDER)
        self.scroll.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        self.test_val_frame = ctk.CTkFrame(self.scroll, fg_color=BG_PANEL,
                                           corner_radius=8, border_width=1,
                                           border_color=BORDER)
        ctk.CTkLabel(self.test_val_frame, text="TEST VALUE  (μ₀)",
                     font=("Segoe UI", 11, "bold"), text_color=TEXT_SEC).pack(
            anchor="w", padx=12, pady=(10, 3))
        self.test_value_entry = styled_entry(self.test_val_frame,
                                             placeholder="0", width=120)
        self.test_value_entry.insert(0, "0")
        self.test_value_entry.pack(anchor="w", padx=12, pady=(0, 10))

        self.g1_card = GroupCard(self.scroll, badge_text="Group 1",
                                 badge_color=ACCENT, group_key="Group_1",
                                 show_sd_in_summary=True)
        self.g1_card.pack(fill="x", pady=(0, 8))

        self.g2_card = GroupCard(self.scroll, badge_text="Group 2",
                                 badge_color=ACCENT2, group_key="Group_2",
                                 show_sd_in_summary=True)

        self.preview_card = ctk.CTkFrame(self.scroll, fg_color=BG_PANEL,
                                         corner_radius=8, border_width=1,
                                         border_color=BORDER)
        ph = ctk.CTkFrame(self.preview_card, fg_color="transparent")
        ph.pack(fill="x", padx=12, pady=(8, 4))
        ctk.CTkLabel(ph, text="DATA PREVIEW", font=("Segoe UI", 11, "bold"),
                     text_color=TEXT_SEC).pack(side="left")
        ctk.CTkButton(ph, text="✕", width=22, height=22, corner_radius=4,
                      fg_color=BORDER, hover_color=DANGER,
                      command=self.hide_preview).pack(side="right")

        self.preview_text = ctk.CTkTextbox(self.preview_card, height=140,
                                           fg_color=BG_INPUT, text_color=TEXT_PRI,
                                           font=FONT_MONO, border_width=1,
                                           border_color=BORDER, corner_radius=6)
        self.preview_text.pack(fill="x", padx=12, pady=(0, 12))
        self.preview_text.configure(state="disabled")

    def _build_results_panel(self, parent):
        self.results_text = ctk.CTkTextbox(parent, fg_color=BG_INPUT,
                                           text_color=TEXT_PRI, font=FONT_MONO,
                                           wrap="none", border_width=1,
                                           border_color=BORDER, corner_radius=8)
        self.results_text.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        self.results_text.configure(state="disabled")

    # ── Test-type switching ────────────────────────────────────────────────────

    def _on_test_change(self, choice):
        self.test_val_frame.pack_forget()
        self.g2_card.pack_forget()

        if choice == "one-sample":
            self.test_val_frame.pack(fill="x", pady=(0, 8), in_=self.scroll)
            self.g1_card.badge.delete(0, "end")
            self.g1_card.badge.insert(0, "Sample")
            self.g1_card.badge.configure(fg_color=ACCENT, border_color=ACCENT)
        elif choice == "paired":
            self.g1_card.badge.delete(0, "end")
            self.g1_card.badge.insert(0, "Pre")
            self.g1_card.badge.configure(fg_color=ACCENT, border_color=ACCENT)
            self.g2_card.badge.delete(0, "end")
            self.g2_card.badge.insert(0, "Post")
            self.g2_card.badge.configure(fg_color=ACCENT2, border_color=ACCENT2)
            self.g2_card.set_show_paired_note(True)
            self.g2_card.pack(fill="x", pady=(0, 8), in_=self.scroll)
        else:  # independent
            self.g1_card.badge.delete(0, "end")
            self.g1_card.badge.insert(0, "Group 1")
            self.g1_card.badge.configure(fg_color=ACCENT, border_color=ACCENT)
            self.g2_card.badge.delete(0, "end")
            self.g2_card.badge.insert(0, "Group 2")
            self.g2_card.badge.configure(fg_color=ACCENT2, border_color=ACCENT2)
            self.g2_card.set_show_paired_note(False)
            self.g2_card.pack(fill="x", pady=(0, 8), in_=self.scroll)

    # ── Import ────────────────────────────────────────────────────────────────

    def import_data(self):
        fp = filedialog.askopenfilename(
            filetypes=[("Excel", "*.xlsx *.xls"), ("CSV", "*.csv"), ("All", "*.*")],
            title="Import Data File")
        if not fp:
            return
        try:
            df = pd.read_csv(fp) if fp.endswith(".csv") else pd.read_excel(fp)
            self.imported_data = df

            if len(df.columns) >= 1:
                d1 = df.iloc[:, 0].dropna().tolist()
                self.g1_card.data_text.delete("1.0", "end")
                self.g1_card.data_text.insert("1.0", ", ".join(map(str, d1)))
                self.g1_card.name_entry.delete(0, "end")
                self.g1_card.name_entry.insert(0, str(df.columns[0]))
                self.g1_card._set_raw()

            if len(df.columns) >= 2 and self.sidebar.test_type_var.get() != "one-sample":
                d2 = df.iloc[:, 1].dropna().tolist()
                self.g2_card.data_text.delete("1.0", "end")
                self.g2_card.data_text.insert("1.0", ", ".join(map(str, d2)))
                self.g2_card.name_entry.delete(0, "end")
                self.g2_card.name_entry.insert(0, str(df.columns[1]))
                self.g2_card._set_raw()

            self.sidebar.preview_btn.configure(state="normal")
            self.sidebar.status_label.configure(text=f"✓ Imported\n{os.path.basename(fp)}")
            messagebox.showinfo("Imported",
                                f"Data imported!\nColumns: {list(df.columns)}\nRows: {len(df)}")
        except Exception as e:
            messagebox.showerror("Import Error", f"Failed:\n{e}")

    def show_preview(self):
        if self.imported_data is None:
            messagebox.showinfo("No Data", "Import a file first.")
            return
        self.preview_card.pack(fill="x", pady=(0, 8), in_=self.scroll)
        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")
        df = self.imported_data
        txt = (f"Shape: {df.shape[0]} rows × {df.shape[1]} cols\n"
               f"Columns: {list(df.columns)}\n\n"
               f"First 10 rows:\n{'─'*50}\n{df.head(10).to_string()}")
        self.preview_text.insert("1.0", txt)
        self.preview_text.configure(state="disabled")

    def hide_preview(self):
        self.preview_card.pack_forget()

    # ── Formatting helpers ─────────────────────────────────────────────────────

    def _fmt_p(self, p):
        return "< .001" if p < 0.001 else f"= {p:.3f}"

    def _fmt_p_tbl(self, p):
        if p < 0.001:
            return "< .001"
        s = f"{p:.3f}"
        return ("." + s[2:]) if s.startswith("0.") else s

    # ── Analysis dispatch ─────────────────────────────────────────────────────

    def run_analysis(self):
        try:
            alpha = float(self.sidebar.alpha_entry.get())
            if not 0 < alpha < 1:
                messagebox.showerror("Error", "Alpha must be between 0 and 1")
                return
        except ValueError:
            messagebox.showerror("Error", "Invalid alpha value")
            return

        self.alpha = alpha
        ttype = self.sidebar.test_type_var.get()

        try:
            if ttype == "one-sample":
                self.results = self._run_one_sample()
            elif ttype == "independent":
                self.results = self._run_independent()
            else:
                self.results = self._run_paired()
        except Exception as e:
            messagebox.showerror("Input Error", str(e))
            return

        self._display_results()
        self.sidebar.pdf_btn.configure(state="normal")
        self.sidebar.docx_btn.configure(state="normal")
        r = self.results
        self.stat_label.configure(
            text=(f"t({r['df']:.2f}) = {r['t_statistic']:.2f}   "
                  f"p {self._fmt_p(r['p_value'])}   "
                  f"{'✓ Significant' if r['p_value'] < r['alpha'] else '✗ Not Significant'}"))

    def _get_group_data_or_summary(self, card, label="Group"):
        if card.mode == "raw":
            data = card.get_raw_data()
            if len(data) < 2:
                raise ValueError(f"{label}: Need at least 2 values in Raw Data mode.")
            return {"mode": "raw", "data": data}
        else:
            try:
                m, sd, n = card.get_summary()
            except (ValueError, TypeError):
                raise ValueError(
                    f"{label}: Please fill in M, SD, and N (all must be valid numbers).")
            if n < 2:
                raise ValueError(f"{label}: N must be at least 2.")
            if sd is not None and sd < 0:
                raise ValueError(f"{label}: SD cannot be negative.")
            return {"mode": "summary", "mean": m, "sd": sd, "n": int(n)}

    # ── One-sample ────────────────────────────────────────────────────────────

    def _run_one_sample(self):
        try:
            tv = float(self.test_value_entry.get())
        except ValueError:
            tv = 0

        info = self._get_group_data_or_summary(self.g1_card, "Sample")
        name = self.g1_card.get_name()

        if info["mode"] == "raw":
            data = info["data"]
            n = len(data)
            m = float(np.mean(data))
            s = float(np.std(data, ddof=1))
            t, p = stats.ttest_1samp(data, tv)
            t = float(t); p = float(p)
        else:
            m, s, n = info["mean"], info["sd"], info["n"]
            if s is None or s == 0:
                raise ValueError("SD is required (and must be > 0) for one-sample summary input.")
            se = s / np.sqrt(n)
            t = float((m - tv) / se)
            df_val = n - 1
            p = float(2 * stats.t.sf(abs(t), df_val))

        return dict(
            test_type="one-sample",
            test_name="One-Sample t-Test",
            group1_name=name,
            test_value=tv,
            n=n, mean=m, std=s,
            se=s / np.sqrt(n),
            t_statistic=t,
            df=n - 1,
            p_value=p,
            cohens_d=(m - tv) / s,
            alpha=self.alpha,
            input_mode=info["mode"])

    # ── Independent ───────────────────────────────────────────────────────────

    def _run_independent(self):
        info1 = self._get_group_data_or_summary(self.g1_card, "Group 1")
        info2 = self._get_group_data_or_summary(self.g2_card, "Group 2")
        name1 = self.g1_card.get_name()
        name2 = self.g2_card.get_name()

        if info1["mode"] == "raw":
            d1 = info1["data"]
            n1 = len(d1)
            m1 = float(np.mean(d1))
            s1 = float(np.std(d1, ddof=1))
        else:
            m1, s1, n1 = info1["mean"], info1["sd"], info1["n"]
            if s1 is None:
                raise ValueError("Group 1: SD is required for summary input.")

        if info2["mode"] == "raw":
            d2 = info2["data"]
            n2 = len(d2)
            m2 = float(np.mean(d2))
            s2 = float(np.std(d2, ddof=1))
        else:
            m2, s2, n2 = info2["mean"], info2["sd"], info2["n"]
            if s2 is None:
                raise ValueError("Group 2: SD is required for summary input.")

        se1_sq = (s1 ** 2) / n1
        se2_sq = (s2 ** 2) / n2
        t = float((m1 - m2) / np.sqrt(se1_sq + se2_sq))
        df_val = float((se1_sq + se2_sq) ** 2 / (se1_sq ** 2 / (n1 - 1) + se2_sq ** 2 / (n2 - 1)))
        p = float(2 * stats.t.sf(abs(t), df_val))
        ps = float(np.sqrt(((n1 - 1) * s1 ** 2 + (n2 - 1) * s2 ** 2) / (n1 + n2 - 2)))
        cohens_d = (m1 - m2) / ps

        return dict(
            test_type="independent",
            test_name="Independent Samples t-Test (Welch's)",
            group1_name=name1, group2_name=name2,
            n1=n1, n2=n2,
            mean1=m1, mean2=m2,
            std1=s1, std2=s2,
            se1=s1 / np.sqrt(n1), se2=s2 / np.sqrt(n2),
            t_statistic=t, df=df_val, p_value=p,
            cohens_d=cohens_d, alpha=self.alpha,
            input_mode=f"G1:{info1['mode']} / G2:{info2['mode']}")

    # ── Paired ────────────────────────────────────────────────────────────────

    def _run_paired(self):
        info1 = self._get_group_data_or_summary(self.g1_card, "Pre")
        info2 = self._get_group_data_or_summary(self.g2_card, "Post")
        name1 = self.g1_card.get_name()
        name2 = self.g2_card.get_name()

        if info1["mode"] == "raw" and info2["mode"] == "raw":
            d1, d2 = info1["data"], info2["data"]
            if len(d1) != len(d2):
                raise ValueError("Paired samples must have equal size.")
            n = len(d1)
            m1, m2 = float(np.mean(d1)), float(np.mean(d2))
            diff = np.array(d1) - np.array(d2)
            md = float(np.mean(diff))
            sd = float(np.std(diff, ddof=1))
            t, p = stats.ttest_rel(d1, d2)
            t = float(t); p = float(p)
        else:
            raise ValueError(
                "⚠  Paired t-test requires Raw Data mode for both groups.\n\n"
                "Tip: For summary-stats paired input, use One-Sample t-test\n"
                "with the differences as your data and test value = 0.")

        return dict(
            test_type="paired",
            test_name="Paired Samples t-Test",
            group1_name=name1, group2_name=name2,
            n=n, mean1=m1, mean2=m2,
            mean_diff=md, std_diff=sd,
            se_diff=sd / np.sqrt(n),
            t_statistic=t, df=n - 1, p_value=p,
            cohens_d=md / sd, alpha=self.alpha,
            input_mode="raw")

    # ── Display ───────────────────────────────────────────────────────────────

    def _display_results(self):
        self.results_text.configure(state="normal")
        self.results_text.delete("1.0", "end")
        r = self.results
        line = "─" * 54

        ct   = self.title_entry.get().strip()
        cs   = self.subtitle_entry.get().strip()
        auth = self.sidebar.researcher_entry.get().strip()

        out = f"{'═'*54}\n"
        if ct:   out += f"{ct.upper()}\n"
        if cs:   out += f"{cs}\n"
        if auth: out += f"by: {auth}\n"
        out += f"{'═'*54}\n\n"

        out += f"{r['test_name'].upper()}\n{line}\n"
        out += f"α = {r['alpha']}   Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        if "input_mode" in r:
            out += f"Input mode: {r['input_mode']}\n"
        out += "\n"

        out += f"DESCRIPTIVE STATISTICS\n{line}\n"
        if r["test_type"] == "one-sample":
            out += f"{r['group1_name']}:  M = {r['mean']:.2f},  SD = {r['std']:.2f},  N = {r['n']}\n"
            out += f"Test Value:  μ₀ = {r['test_value']}\n\n"
        elif r["test_type"] == "independent":
            out += f"{r['group1_name']}:  M = {r['mean1']:.2f},  SD = {r['std1']:.2f},  N = {r['n1']}\n"
            out += f"{r['group2_name']}:  M = {r['mean2']:.2f},  SD = {r['std2']:.2f},  N = {r['n2']}\n\n"
        else:
            out += f"{r['group1_name']}:  M = {r['mean1']:.2f},  N = {r['n']}\n"
            out += f"{r['group2_name']}:  M = {r['mean2']:.2f},  N = {r['n']}\n"
            out += f"Mean Diff:  {r['mean_diff']:.2f},  SD(diff) = {r['std_diff']:.2f}\n\n"

        out += f"TEST STATISTICS\n{line}\n"
        out += f"t  = {r['t_statistic']:.4f}\n"
        out += f"df = {r['df']:.2f}\n"
        out += f"p  {self._fmt_p(r['p_value'])}\n"
        out += f"d  = {r['cohens_d']:.4f}\n\n"

        out += f"DECISION\n{line}\n"
        if r["p_value"] < r["alpha"]:
            out += f"✓ REJECT H₀  (p {self._fmt_p(r['p_value'])} < α = {r['alpha']})\n\n"
        else:
            out += f"✗ FAIL TO REJECT H₀  (p {self._fmt_p(r['p_value'])} ≥ α = {r['alpha']})\n\n"

        out += f"INTERPRETATION\n{line}\n"
        out += self._interpretation() + "\n\n"
        out += "═" * 54 + "\n"

        self.results_text.insert("1.0", out)
        self.results_text.configure(state="disabled")

    def _interpretation(self):
        r   = self.results
        pf  = self._fmt_p(r["p_value"])
        sig = r["p_value"] < r["alpha"]

        if r["test_type"] == "one-sample":
            w = ("was significantly different from" if sig
                 else "was not significantly different from")
            return (f"A one-sample t-test revealed that the sample mean "
                    f"(M = {r['mean']:.2f}) {w} the test value ({r['test_value']}), "
                    f"t({r['df']:.0f}) = {r['t_statistic']:.2f}, p {pf}, "
                    f"d = {r['cohens_d']:.2f}.")

        elif r["test_type"] == "independent":
            w = ("revealed a statistically significant difference" if sig
                 else "showed no statistically significant difference")
            return (f"An independent samples t-test {w} between "
                    f"{r['group1_name']} (M = {r['mean1']:.2f}) and "
                    f"{r['group2_name']} (M = {r['mean2']:.2f}), "
                    f"t({r['df']:.2f}) = {r['t_statistic']:.2f}, p {pf}, "
                    f"d = {r['cohens_d']:.2f}.")
        else:
            w = ("revealed a statistically significant difference" if sig
                 else "showed no statistically significant difference")
            return (f"A paired samples t-test {w} between "
                    f"{r['group1_name']} (M = {r['mean1']:.2f}) and "
                    f"{r['group2_name']} (M = {r['mean2']:.2f}), "
                    f"mean difference = {r['mean_diff']:.2f}, "
                    f"t({r['df']:.0f}) = {r['t_statistic']:.2f}, p {pf}, "
                    f"d = {r['cohens_d']:.2f}.")

    # ── Theme ─────────────────────────────────────────────────────────────────

    def toggle_theme(self):
        if self.dark_mode:
            ctk.set_appearance_mode("light")
            self.sidebar.theme_btn.configure(text="🌙  Dark Mode")
            self.dark_mode = False
        else:
            ctk.set_appearance_mode("dark")
            self.sidebar.theme_btn.configure(text="☀️  Light Mode")
            self.dark_mode = True

    # ── Export PDF ────────────────────────────────────────────────────────────

    def export_pdf(self):
        if not self.results:
            messagebox.showerror("Error", "Run analysis first")
            return
        if not PDF_AVAILABLE:
            messagebox.showerror("Error", "Install reportlab: pip install reportlab")
            return

        fp = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")],
            initialfile=f"ttest_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            title="Save PDF Report")
        if not fp:
            return

        try:
            doc = SimpleDocTemplate(fp, pagesize=letter)
            story, styles = [], getSampleStyleSheet()
            r  = self.results
            ct = self.title_entry.get().strip()
            cs = self.subtitle_entry.get().strip()

            if ct:
                story.append(Paragraph(ct, styles["Title"]))
                story.append(Spacer(1, 0.1 * inch))
            if cs:
                sub_style = ParagraphStyle("Sub", parent=styles["Heading2"],
                                           fontSize=13, textColor=colors.HexColor("#555555"))
                story.append(Paragraph(cs, sub_style))
                story.append(Spacer(1, 0.15 * inch))

            story.append(Paragraph(r["test_name"], styles["Title"]))
            story.append(Spacer(1, 0.15 * inch))
            story.append(Paragraph(f"Alpha: α = {r['alpha']}", styles["Normal"]))
            if "input_mode" in r:
                story.append(Paragraph(f"Input mode: {r['input_mode']}", styles["Normal"]))
            story.append(Spacer(1, 0.25 * inch))

            tbl_style_cmd = TableStyle([
                ("ALIGN",    (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1,  0), "Helvetica-Bold"),
                ("LINEABOVE",(0, 0), (-1,  0), 1, colors.black),
                ("LINEBELOW",(0, 0), (-1,  0), 1, colors.black),
                ("LINEBELOW",(0,-1), (-1, -1), 1, colors.black),
            ])

            story.append(Paragraph("Descriptive Statistics", styles["Heading2"]))
            if r["test_type"] == "one-sample":
                d = [["Group", "Mean", "SD", "N"],
                     [r["group1_name"], f"{r['mean']:.2f}", f"{r['std']:.2f}", str(r["n"])]]
            elif r["test_type"] == "independent":
                d = [["Group", "Mean", "SD", "N"],
                     [r["group1_name"], f"{r['mean1']:.2f}", f"{r['std1']:.2f}", str(r["n1"])],
                     [r["group2_name"], f"{r['mean2']:.2f}", f"{r['std2']:.2f}", str(r["n2"])]]
            else:
                d = [["Measurement", "Mean", "N"],
                     [r["group1_name"], f"{r['mean1']:.2f}", str(r["n"])],
                     [r["group2_name"], f"{r['mean2']:.2f}", str(r["n"])]]
            t1 = Table(d)
            t1.setStyle(tbl_style_cmd)
            story.append(t1)
            story.append(Spacer(1, 0.2 * inch))

            story.append(Paragraph("Test Statistics", styles["Heading2"]))
            sd_data = [["Statistic", "Value"],
                       ["t",         f"{r['t_statistic']:.4f}"],
                       ["df",        f"{r['df']:.2f}"],
                       ["p",         self._fmt_p_tbl(r["p_value"])],
                       ["Cohen's d", f"{r['cohens_d']:.4f}"]]
            t2 = Table(sd_data)
            t2.setStyle(tbl_style_cmd)
            story.append(t2)
            story.append(Spacer(1, 0.2 * inch))

            story.append(Paragraph("Interpretation", styles["Heading2"]))
            story.append(Paragraph(self._interpretation(), styles["Normal"]))
            story.append(Spacer(1, 0.3 * inch))

            fs = ParagraphStyle("Foot", parent=styles["Normal"], fontSize=7,
                                textColor=colors.grey, fontName="Helvetica-Oblique")
            story.append(Paragraph(f"Saved: {fp}", fs))
            story.append(Paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", fs))

            doc.build(story)
            self.file_label.configure(text=f"Last saved: {fp}")
            self.sidebar.status_label.configure(text=f"✓ PDF Saved\n{os.path.basename(fp)}")
            messagebox.showinfo("Saved", f"PDF exported:\n{fp}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed:\n{e}")

    # ── Export DOCX ───────────────────────────────────────────────────────────

    def export_docx(self):
        if not self.results:
            messagebox.showerror("Error", "Run analysis first")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "Install python-docx: pip install python-docx")
            return

        fp = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx"), ("All", "*.*")],
            initialfile=f"ttest_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            title="Save Word Report")
        if not fp:
            return

        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin    = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin   = Inches(1.0)
                section.right_margin  = Inches(1.0)

            r    = self.results
            ct   = self.title_entry.get().strip()
            cs   = self.subtitle_entry.get().strip()
            auth = self.sidebar.researcher_entry.get().strip()

            def apa_borders(table):
                tbl = table._tbl
                tblPr = tbl.tblPr
                tb = OxmlElement("w:tblBorders")
                for bn in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    b = OxmlElement(f"w:{bn}")
                    b.set(qn("w:val"), "none")
                    tb.append(b)
                for bn, sz in [("top", "12"), ("bottom", "12")]:
                    b = OxmlElement(f"w:{bn}")
                    b.set(qn("w:val"), "single")
                    b.set(qn("w:sz"), sz)
                    tb.append(b)
                tblPr.append(tb)

            def hdr_sep(table):
                for cell in table.rows[0].cells:
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcB  = OxmlElement("w:tcBorders")
                    bot  = OxmlElement("w:bottom")
                    bot.set(qn("w:val"), "single")
                    bot.set(qn("w:sz"), "6")
                    tcB.append(bot)
                    tcPr.append(tcB)

            def cfmt(cell, text, bold=False, size=10, align="left"):
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                               else WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(str(text))
                run.font.size = Pt(size)
                run.bold = bold

            if ct:
                h = doc.add_heading(ct.upper(), 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cs:
                sp = doc.add_paragraph(cs)
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp.runs[0].italic = True
                sp.runs[0].font.size = Pt(12)
            if auth:
                ap = doc.add_paragraph(f"by: {auth}")
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ap.runs[0].italic = True
                ap.runs[0].font.size = Pt(10)

            doc.add_heading(r["test_name"], 1)
            doc.add_paragraph(f"Alpha: α = {r['alpha']}")
            if "input_mode" in r:
                doc.add_paragraph(f"Input mode: {r['input_mode']}")
            doc.add_paragraph()

            doc.add_heading("Descriptive Statistics", 2)
            if r["test_type"] == "one-sample":
                dt = doc.add_table(rows=2, cols=4)
                apa_borders(dt); hdr_sep(dt)
                for i, h in enumerate(["Group", "M", "SD", "N"]):
                    cfmt(dt.rows[0].cells[i], h, bold=True, align="center")
                cfmt(dt.rows[1].cells[0], r["group1_name"])
                cfmt(dt.rows[1].cells[1], f"{r['mean']:.2f}", align="center")
                cfmt(dt.rows[1].cells[2], f"{r['std']:.2f}",  align="center")
                cfmt(dt.rows[1].cells[3], str(r["n"]),         align="center")
            elif r["test_type"] == "independent":
                dt = doc.add_table(rows=3, cols=4)
                apa_borders(dt); hdr_sep(dt)
                for i, h in enumerate(["Group", "M", "SD", "N"]):
                    cfmt(dt.rows[0].cells[i], h, bold=True, align="center")
                cfmt(dt.rows[1].cells[0], r["group1_name"])
                cfmt(dt.rows[1].cells[1], f"{r['mean1']:.2f}", align="center")
                cfmt(dt.rows[1].cells[2], f"{r['std1']:.2f}",  align="center")
                cfmt(dt.rows[1].cells[3], str(r["n1"]),         align="center")
                cfmt(dt.rows[2].cells[0], r["group2_name"])
                cfmt(dt.rows[2].cells[1], f"{r['mean2']:.2f}", align="center")
                cfmt(dt.rows[2].cells[2], f"{r['std2']:.2f}",  align="center")
                cfmt(dt.rows[2].cells[3], str(r["n2"]),         align="center")
            else:
                dt = doc.add_table(rows=4, cols=3)
                apa_borders(dt); hdr_sep(dt)
                for i, h in enumerate(["Measurement", "M", "N"]):
                    cfmt(dt.rows[0].cells[i], h, bold=True, align="center")
                cfmt(dt.rows[1].cells[0], r["group1_name"])
                cfmt(dt.rows[1].cells[1], f"{r['mean1']:.2f}", align="center")
                cfmt(dt.rows[1].cells[2], str(r["n"]),          align="center")
                cfmt(dt.rows[2].cells[0], r["group2_name"])
                cfmt(dt.rows[2].cells[1], f"{r['mean2']:.2f}", align="center")
                cfmt(dt.rows[2].cells[2], str(r["n"]),          align="center")
                cfmt(dt.rows[3].cells[0], "Mean Difference")
                cfmt(dt.rows[3].cells[1], f"{r['mean_diff']:.2f}", align="center")
                cfmt(dt.rows[3].cells[2], f"SD(diff)={r['std_diff']:.2f}", align="center")

            doc.add_paragraph()

            doc.add_heading("Test Statistics", 2)
            st = doc.add_table(rows=5, cols=2)
            apa_borders(st); hdr_sep(st)
            for i, h in enumerate(["Statistic", "Value"]):
                cfmt(st.rows[0].cells[i], h, bold=True, align="center")
            for row_i, (stat, val) in enumerate([
                ("t",         f"{r['t_statistic']:.4f}"),
                ("df",        f"{r['df']:.2f}"),
                ("p",         self._fmt_p_tbl(r["p_value"])),
                ("Cohen's d", f"{r['cohens_d']:.4f}")
            ], 1):
                cfmt(st.rows[row_i].cells[0], stat)
                cfmt(st.rows[row_i].cells[1], val, align="center")

            doc.add_paragraph()

            doc.add_heading("Decision", 2)
            if r["p_value"] < r["alpha"]:
                doc.add_paragraph(
                    f"Reject the null hypothesis "
                    f"(p {self._fmt_p(r['p_value'])} < α = {r['alpha']}).")
            else:
                doc.add_paragraph(
                    f"Fail to reject the null hypothesis "
                    f"(p {self._fmt_p(r['p_value'])} ≥ α = {r['alpha']}).")

            doc.add_heading("Interpretation", 2)
            doc.add_paragraph(self._interpretation())

            doc.add_paragraph()
            fp_p = doc.add_paragraph(f"Saved: {fp}")
            fp_p.runs[0].italic = True
            fp_p.runs[0].font.size = Pt(7)
            gp = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            gp.runs[0].italic = True
            gp.runs[0].font.size = Pt(7)

            doc.save(fp)
            self.file_label.configure(text=f"Last saved: {fp}")
            self.sidebar.status_label.configure(text=f"✓ DOCX Saved\n{os.path.basename(fp)}")
            messagebox.showinfo("Saved", f"Word document exported:\n{fp}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed:\n{e}")

    # ── Settings & Clear ──────────────────────────────────────────────────────

    def open_settings(self):
        if SETTINGS_AVAILABLE:
            SettingsWindow(self, self)
        else:
            messagebox.showinfo("Settings", "app_settings module not found.")

    def apply_settings(self):
        if not SETTINGS_AVAILABLE:
            return
        sm = SettingsManager()
        fb, fc, fh, fm, fbt, ft = sm.fonts
        ff = sm.font_family
        self.results_text.configure(font=(ff, fm), wrap=sm.wrap_mode)
        self.sidebar.status_label.configure(font=(ff, ft))
        self.stat_label.configure(font=(ff, ft, "bold"))
        self.file_label.configure(font=(ff, ft))
        self.sidebar.run_btn.configure(fg_color=sm.accent, hover_color=sm.accent_hover)
        self.sidebar.configure(width=sm.sidebar_width)

    def clear_fields(self):
        self.g1_card.clear()
        self.g2_card.clear()
        self.g1_card.name_entry.delete(0, "end")
        self.g1_card.name_entry.insert(0, "Group_1")
        self.g2_card.name_entry.delete(0, "end")
        self.g2_card.name_entry.insert(0, "Group_2")
        self.title_entry.delete(0, "end")
        self.subtitle_entry.delete(0, "end")
        self.subtitle_entry.insert(0, "t-Test")
        self.results_text.configure(state="normal")
        self.results_text.delete("1.0", "end")
        self.results_text.configure(state="disabled")
        self.results = None
        self.imported_data = None
        self.stat_label.configure(text="")
        self.sidebar.status_label.configure(text="")
        self.sidebar.pdf_btn.configure(state="disabled")
        self.sidebar.docx_btn.configure(state="disabled")
        self.sidebar.preview_btn.configure(state="disabled")
        self.hide_preview()

    def _prefill_next_part_label(self):
        if SESSION_AVAILABLE:
            next_label = _next_part_label([p.get("label", "") for p in self.saved_parts])
            self.part_label_entry.delete(0, "end")
            self.part_label_entry.insert(0, next_label)

    # ── Session Manager ───────────────────────────────────────────────────────

    def open_session_manager(self):
        if self._session_win and self._session_win.winfo_exists():
            self._session_win.focus()
            return
        if not SESSION_AVAILABLE:
            messagebox.showerror("Error",
                                 "ttest_session_manager.py not found.\n"
                                 "Make sure it is in the same folder.")
            return

        self._session_win = SessionManagerPanel(
            self,
            saved_parts=self.saved_parts,
            on_delete_part=self._delete_part,
            on_load_part=self._load_part,
            on_export_all_docx=self._export_all_docx,
            on_export_all_pdf=self._export_all_pdf,
        )

    def _delete_part(self, label: str):
        self.saved_parts = [p for p in self.saved_parts if p.get("label") != label]
        self._prefill_next_part_label()
        self.sidebar.status_label.configure(
            text=f"✓ Part deleted\n{len(self.saved_parts)} remaining")

    def _load_part(self, part: dict):
        self.title_entry.delete(0, "end")
        self.title_entry.insert(0, part.get("title", ""))
        self.subtitle_entry.delete(0, "end")
        self.subtitle_entry.insert(0, part.get("subtitle", ""))
        self.part_label_entry.delete(0, "end")
        self.part_label_entry.insert(0, part.get("label", ""))
        self.sidebar.researcher_entry.delete(0, "end")
        self.sidebar.researcher_entry.insert(0, part.get("researcher", ""))
        self.sidebar.alpha_entry.delete(0, "end")
        self.sidebar.alpha_entry.insert(0, str(part.get("alpha", 0.05)))

        test_type = part.get("test_type", "independent")
        self.sidebar.test_type_var.set(test_type)
        self._on_test_change(test_type)

        if test_type == "one-sample":
            self.test_value_entry.delete(0, "end")
            self.test_value_entry.insert(0, str(part.get("test_value", 0)))
            self.g1_card.data_text.delete("1.0", "end")
            self.g1_card.data_text.insert("1.0", ", ".join(map(str, part.get("data", []))))
            self.g1_card.name_entry.delete(0, "end")
            self.g1_card.name_entry.insert(0, part.get("group1_name", "Sample"))
        else:
            self.g1_card.data_text.delete("1.0", "end")
            self.g1_card.data_text.insert("1.0", ", ".join(map(str, part.get("data1", []))))
            self.g1_card.name_entry.delete(0, "end")
            self.g1_card.name_entry.insert(0, part.get("group1_name", "Group 1"))
            self.g2_card.data_text.delete("1.0", "end")
            self.g2_card.data_text.insert("1.0", ", ".join(map(str, part.get("data2", []))))
            self.g2_card.name_entry.delete(0, "end")
            self.g2_card.name_entry.insert(0, part.get("group2_name", "Group 2"))

        self.sidebar.status_label.configure(text=f"✓ Loaded\n{part.get('label','')}")

    def _export_all_docx(self):
        if not self.saved_parts:
            messagebox.showwarning("Warning", "No saved parts to export!")
            return
        fp = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            title="Export All Parts to DOCX")
        if not fp:
            return
        try:
            from ttest_session_manager import export_all_to_docx
            export_all_to_docx(self.saved_parts, fp)
            self.sidebar.status_label.configure(
                text=f"✓ {len(self.saved_parts)} parts\nexported")
            messagebox.showinfo("Success",
                                f"All {len(self.saved_parts)} part(s) exported!\n\n{fp}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export:\n{e}")

    def _export_all_pdf(self):
        if not self.saved_parts:
            messagebox.showwarning("Warning", "No saved parts to export!")
            return
        fp = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF File", "*.pdf"), ("All Files", "*.*")],
            title="Export All Parts to PDF")
        if not fp:
            return
        try:
            from ttest_session_manager import export_all_to_pdf
            export_all_to_pdf(self.saved_parts, fp)
            self.sidebar.status_label.configure(
                text=f"✓ {len(self.saved_parts)} parts\nexported")
            messagebox.showinfo("Success",
                                f"All {len(self.saved_parts)} part(s) exported!\n\n{fp}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export:\n{e}")

    def save_current_part(self):
        if not self.results:
            messagebox.showwarning("Warning", "Run analysis first before saving a part!")
            return

        part_label = self.part_label_entry.get().strip()
        if not part_label:
            part_label = (_next_part_label([p.get("label", "") for p in self.saved_parts])
                          if SESSION_AVAILABLE else f"Part {len(self.saved_parts) + 1}-A")

        r         = self.results
        test_type = self.sidebar.test_type_var.get()

        part = {
            "label":      part_label,
            "title":      self.title_entry.get().strip(),
            "subtitle":   self.subtitle_entry.get().strip(),
            "researcher": self.sidebar.researcher_entry.get().strip(),
            "test_type":  test_type,
            "alpha":      float(self.sidebar.alpha_entry.get()),
            "results":    r.copy(),
            # t-test specific fields for display / export
            "t_statistic":    r.get("t_statistic"),
            "df":             r.get("df"),
            "p_value":        r.get("p_value"),
            "cohens_d":       r.get("cohens_d"),
            "is_significant": r.get("p_value", 1) < r.get("alpha", 0.05),
            "decision": ("Reject H₀" if r.get("p_value", 1) < r.get("alpha", 0.05)
                         else "Fail to Reject H₀"),
            "interpretation": self._interpretation(),
            "saved_at":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "test_name":  r.get("test_name", "t-Test"),
        }

        # Descriptive stats
        if test_type == "one-sample":
            part.update({
                "test_value":   float(self.test_value_entry.get()),
                "group1_name":  self.g1_card.get_name(),
                "data":         self.g1_card.get_raw_data(),
                "mean1":        r.get("mean"),
                "std1":         r.get("std"),
                "n1":           r.get("n"),
            })
        else:
            part.update({
                "group1_name": self.g1_card.get_name(),
                "group2_name": self.g2_card.get_name(),
                "data1":       self.g1_card.get_raw_data(),
                "data2":       self.g2_card.get_raw_data(),
                "mean1":       r.get("mean1"),
                "mean2":       r.get("mean2"),
                "std1":        r.get("std1"),
                "std2":        r.get("std2"),
                "n1":          r.get("n1") or r.get("n"),
                "n2":          r.get("n2") or r.get("n"),
            })
            if test_type == "paired":
                part.update({
                    "mean_diff": r.get("mean_diff"),
                    "std_diff":  r.get("std_diff"),
                })

        self.saved_parts.append(part)
        self._prefill_next_part_label()
        self.sidebar.status_label.configure(text=f"✓ Part saved\n{part_label}")
        messagebox.showinfo("Part Saved", f"Analysis saved as:\n{part_label}")


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = TTestApp()
    app.mainloop()